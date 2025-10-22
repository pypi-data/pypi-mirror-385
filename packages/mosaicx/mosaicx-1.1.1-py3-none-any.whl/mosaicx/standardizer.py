"""
MOSAICX Standardizer - Schema-Guided Structured Reporting

================================================================================
MOSAICX: Medical cOmputational Suite for Advanced Intelligent eXtraction
================================================================================

Structure first. Insight follows.

Author: Lalith Kumar Shiyam Sundar, PhD
Lab: DIGIT-X Lab
Department: Department of Radiology
University: LMU University Hospital | LMU Munich

Overview:
---------
Translate a clinical document into a structured report that adheres to a
previously generated Pydantic schema. This module builds on the Docling +
LLM extraction pipeline and adds artifact writers for JSON, plain text,
PDF, and DOCX outputs.
"""

from __future__ import annotations

import json
import re
from collections import OrderedDict
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Tuple, Union

from pydantic import BaseModel

from .constants import APPLICATION_VERSION, DEFAULT_LLM_MODEL, PROJECT_ROOT
from .display import styled_message
from .utils import resolve_openai_config
from .extractor import (
    ExtractionError,
    extract_structured_data,
    extract_text_from_document,
    load_schema_model,
)

try:  # Optional LLM dependency (currently unused but kept for future extensions)
    from openai import OpenAI  # type: ignore[import-not-found]
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    OpenAI = None  # type: ignore[assignment]

try:  # Optional PDF dependency
    from reportlab.lib import colors as rl_colors  # type: ignore[import-not-found]
    from reportlab.lib.pagesizes import A4  # type: ignore[import-not-found]
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-not-found]
    from reportlab.lib.units import cm  # type: ignore[import-not-found]
    from reportlab.platypus import (  # type: ignore[import-not-found]
        Image,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table as RLTable,
        TableStyle,
    )
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    rl_colors = None  # type: ignore[assignment]
    A4 = None  # type: ignore[assignment]
    ParagraphStyle = None  # type: ignore[assignment]
    getSampleStyleSheet = None  # type: ignore[assignment]
    cm = None  # type: ignore[assignment]
    Image = None  # type: ignore[assignment]
    Paragraph = None  # type: ignore[assignment]
    SimpleDocTemplate = None  # type: ignore[assignment]
    Spacer = None  # type: ignore[assignment]
    RLTable = None  # type: ignore[assignment]
    TableStyle = None  # type: ignore[assignment]

try:  # Optional DOCX dependency
    from docx import Document  # type: ignore[import-not-found]
    from docx.shared import Pt  # type: ignore[import-not-found]
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]

try:
    from xml.sax.saxutils import escape
except Exception:  # pragma: no cover - fallback
    def escape(text: str) -> str:
        return text


@dataclass(slots=True)
class ReportSection:
    title: str
    kind: str  # "kv", "bullets", "text"
    data: Any


@dataclass(slots=True)
class StandardizeResult:
    """Structured reporting payload produced by :func:`standardize_document`."""

    record: BaseModel
    schema_path: Path
    document_path: Optional[Path]
    schema_name: str
    model_name: str
    generated_at: datetime
    extracted_text: Optional[str] = None
    narrative: Optional[str] = None
    narrative_error: Optional[str] = None
    snapshot_data: Optional[OrderedDict[str, str]] = None
    structured_sections: Optional[List[ReportSection]] = None
    narrative_sections: Optional[List[ReportSection]] = None

    def to_dict(self) -> Dict[str, Any]:
        return self.record.model_dump()

    def to_json(self, *, indent: int = 2) -> str:
        return self.record.model_dump_json(indent=indent, by_alias=True)

    def default_basename(self) -> str:
        timestamp = self.generated_at.strftime("%Y%m%d_%H%M%S")
        if self.document_path is not None and self.document_path.stem:
            doc_stem = self.document_path.stem
        else:
            doc_stem = self.schema_name.lower()
        return f"standardize_{doc_stem}_{timestamp}"

    def default_path(self, *, extension: str, directory: Optional[Path] = None) -> Path:
        root = Path(directory) if directory is not None else Path("output")
        return root / f"{self.default_basename()}.{extension.lstrip('.')}"

    def write_json(self, path: Path, *, indent: int = 2) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(self.to_json(indent=indent), encoding="utf-8")
        return path


def _is_empty(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip() == ""
    if isinstance(value, (list, tuple, set)):
        return all(_is_empty(v) for v in value)
    if isinstance(value, dict):
        return all(_is_empty(v) for v in value.values())
    return False


def _format_scalar(value: Any) -> str:
    if value is None:
        return "None"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return str(value)
    if isinstance(value, str):
        stripped = value.strip()
        return stripped if stripped else "None"
    return json.dumps(value, ensure_ascii=False)


_LEVEL_PATTERN = re.compile(r"^[A-Za-z]\d+$")


def _humanize_key(key: str) -> str:
    clean = key.replace("_", " ").replace("-", " ").strip()
    if not clean:
        return key
    raw_tokens = re.split(r"\s+", clean)
    words: List[str] = []
    for token in raw_tokens:
        upper = token.upper()
        if upper in {"MRI", "CT", "PET", "MR", "T1", "T2"}:
            words.append(upper)
        else:
            words.append(token.capitalize())
    if len(words) == 2 and _LEVEL_PATTERN.match(words[0]) and _LEVEL_PATTERN.match(words[1]):
        return f"{words[0]}-{words[1]}"
    return " ".join(words)


def _collect_sections_from_payload(payload: Dict[str, Any]) -> Tuple[OrderedDict[str, str], List[ReportSection]]:
    snapshot_tokens = {"patient", "dob", "birth", "age", "sex", "mrn", "id", "name"}
    technique_tokens = {"sequence", "technique", "protocol", "procedure"}

    snapshot: "OrderedDict[str, str]" = OrderedDict()
    sections: List[ReportSection] = []
    misc_kv: "OrderedDict[str, str]" = OrderedDict()

    for key, value in payload.items():
        if _is_empty(value):
            continue
        lower = key.lower()
        if any(token in lower for token in snapshot_tokens) and not isinstance(value, (list, dict)):
            snapshot[_humanize_key(key)] = _format_scalar(value)
            continue
        if isinstance(value, list):
            items = [item for item in value if not _is_empty(item)]
            if not items:
                continue
            title = _humanize_key(key)
            if all(not isinstance(item, (dict, list, tuple, set)) for item in items):
                sections.append(ReportSection(title=title, kind="bullets", data=[_format_scalar(item) for item in items]))
            else:
                sentences: List[str] = []
                for idx, item in enumerate(items, start=1):
                    if isinstance(item, dict):
                        parts = []
                        for sub_key, sub_val in item.items():
                            if _is_empty(sub_val):
                                continue
                            parts.append(f"{_humanize_key(sub_key)}: {_format_scalar(sub_val)}")
                        if parts:
                            sentences.append("; ".join(parts))
                    else:
                        sentences.append(_format_scalar(item))
                if sentences:
                    sections.append(ReportSection(title=title, kind="text", data="\n".join(sentences)))
        elif isinstance(value, dict):
            kv_pairs = OrderedDict()
            for sub_key, sub_val in value.items():
                if _is_empty(sub_val):
                    continue
                kv_pairs[_humanize_key(sub_key)] = _format_scalar(sub_val)
            if kv_pairs:
                sections.append(ReportSection(title=_humanize_key(key), kind="kv", data=kv_pairs))
        else:
            if any(token in lower for token in technique_tokens):
                sections.append(ReportSection(title=_humanize_key(key), kind="text", data=_format_scalar(value)))
            else:
                misc_kv[_humanize_key(key)] = _format_scalar(value)

    if misc_kv:
        sections.insert(0, ReportSection(title="Report Summary", kind="kv", data=misc_kv))

    return snapshot, sections


def _prepare_sections(result: StandardizeResult) -> Tuple[OrderedDict[str, str], List[ReportSection]]:
    snapshot = result.snapshot_data
    sections = result.structured_sections
    if snapshot is None or sections is None:
        snapshot, sections = _collect_sections_from_payload(result.to_dict())
    if result.narrative_sections is not None:
        return snapshot or OrderedDict(), result.narrative_sections
    if result.narrative:
        narrative_text = result.narrative.strip()
        if narrative_text:
            return snapshot or OrderedDict(), [ReportSection(title="Findings", kind="text", data=narrative_text)]
    return snapshot or OrderedDict(), sections


def _fallback_text_for_section(section: ReportSection) -> str:
    if section.kind == "kv":
        return "; ".join(f"{label}: {value}" for label, value in section.data.items())
    if section.kind == "bullets":
        return "; ".join(section.data)
    return str(section.data)


def _render_text_report(result: StandardizeResult) -> str:
    snapshot, sections = _prepare_sections(result)
    doc_label = result.document_path.name if result.document_path is not None else "—"
    lines = [
        "MOSAICX Structured Report",
        f"Schema: {result.schema_name}",
        f"Model: {result.model_name}",
        f"Document: {doc_label}",
        f"Generated: {result.generated_at.isoformat().replace('+00:00', 'Z')}",
        "",
    ]
    if snapshot:
        lines.append("PATIENT SNAPSHOT:")
        for label, value in snapshot.items():
            lines.append(f"  - {label}: {value}")
        lines.append("")
    for section in sections:
        if section.title:
            lines.append(f"{section.title.upper()}:")
        if section.kind == "kv":
            for label, value in section.data.items():
                lines.append(f"  - {label}: {value}")
        elif section.kind == "bullets":
            for item in section.data:
                lines.append(f"  - {item}")
        elif section.kind == "text":
            for line in str(section.data).splitlines():
                lines.append(f"  {line}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _summarize_schema(schema_json: Dict[str, Any]) -> str:
    props = schema_json.get("properties", {}) or {}
    lines: List[str] = []
    for name, spec in props.items():
        spec = spec or {}
        spec_type = spec.get("type")
        if isinstance(spec_type, list):
            spec_type = "/".join(spec_type)
        elif spec_type is None and "$ref" in spec:
            spec_type = "$ref"
        enum_vals = spec.get("enum")
        if enum_vals:
            preview = ", ".join(map(str, enum_vals))
            if len(preview) > 80:
                preview = preview[:77] + "..."
            lines.append(f"- {name}: type={spec_type or 'object'} enum=[{preview}]")
        else:
            lines.append(f"- {name}: type={spec_type or 'object'}")
    return "\n".join(lines)


def _build_narrative_prompt(schema_class: type[BaseModel], record_dict: Dict[str, Any]) -> Tuple[str, str, List[str]]:
    schema_json = schema_class.model_json_schema()
    schema_summary = _summarize_schema(schema_json)
    snapshot, sections = _collect_sections_from_payload(record_dict)
    outline_lines: List[str] = []
    ordered_titles: List[str] = []
    if snapshot:
        outline_lines.append("PATIENT SNAPSHOT:")
        for label, value in snapshot.items():
            outline_lines.append(f"  - {label}: {value}")
        ordered_titles.append("PATIENT SNAPSHOT")
    for section in sections:
        title = section.title.upper()
        outline_lines.append(f"{title}:")
        ordered_titles.append(title)
        if section.kind == "kv":
            for label, value in section.data.items():
                outline_lines.append(f"  - {label}: {value}")
        elif section.kind == "bullets":
            for item in section.data:
                outline_lines.append(f"  - {item}")
        elif section.kind == "text":
            for line in str(section.data).splitlines():
                outline_lines.append(f"  - {line}")
    outline = "\n".join(outline_lines)
    return schema_summary, outline, ordered_titles


ABSENCE_STRINGS = {
    "none",
    "normal",
    "no abnormality",
    "no abnormality detected",
    "not seen",
    "absent",
    "n/a",
    "na",
    "--",
    "-",
}


def _generate_narrative(
    schema_class: type[BaseModel],
    record_dict: Dict[str, Any],
    structured_sections: List[ReportSection],
    *,
    model: str,
    base_url: Optional[str],
    api_key: Optional[str],
    temperature: float,
) -> List[ReportSection]:
    narrative_sections: List[ReportSection] = []
    for section in structured_sections:
        sentences: List[str] = []
        if section.kind == "kv":
            for label, value in section.data.items():
                sentences.append(_sentence_for_field(_humanize_key(label), _format_scalar(value)))
        elif section.kind == "bullets":
            for item in section.data:
                if isinstance(item, str):
                    sentences.append(_sentence_for_field(section.title, item))
                elif isinstance(item, dict):
                    sentences.append(_sentence_for_dict(item))
        elif section.kind == "text":
            sentences.append(str(section.data))
        paragraph = " ".join(sent for sent in sentences if sent)
        if paragraph:
            narrative_sections.append(ReportSection(title=section.title, kind="text", data=paragraph))
    return narrative_sections


def _is_absence_value(text: str) -> bool:
    normalised = text.strip().lower()
    if normalised in ABSENCE_STRINGS:
        return True
    return normalised.startswith("no ") and "abnormal" in normalised


def _sentence_for_field(label: str, value: str) -> str:
    label_clean = _humanize_key(label)
    text = value.strip()
    if not text:
        return ""
    if _is_absence_value(text):
        return f"{label_clean} appears normal."
    sentence = text[0].upper() + text[1:]
    if not sentence.endswith("."):
        sentence += "."
    if sentence.lower().startswith("no "):
        return f"{label_clean}: {sentence}"
    return f"{label_clean}: {sentence}"


def _sentence_for_dict(item: Dict[str, Any]) -> str:
    level = _humanize_key(str(item.get("level", ""))) if "level" in item else ""
    details: List[str] = []
    for key, val in item.items():
        if key == "level":
            continue
        label = _humanize_key(key)
        text = _format_scalar(val)
        if _is_absence_value(text):
            details.append(f"{label.lower()} normal")
        else:
            details.append(f"{label.lower()} {text}")
    if not details:
        summary = "normal appearance"
    else:
        summary = ", ".join(details)
    if not summary.endswith("."):
        summary += "."
    if level:
        return f"{level}: {summary}"
    return summary


def _clean_narrative(text: str) -> str:
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # remove bold markdown
    cleaned = re.sub(r"\*(.*?)\*", r"\1", cleaned)  # remove italic markdown
    cleaned = cleaned.replace("•", "-").replace("·", "-")
    cleaned = cleaned.replace("  ", " ")
    return cleaned.strip()


def _ensure_pdf_dependencies() -> None:
    required = (rl_colors, A4, Paragraph, SimpleDocTemplate, RLTable, TableStyle, cm, Spacer)
    if any(dep is None for dep in required):
        raise RuntimeError("PDF export unavailable – install ReportLab (`pip install reportlab`).")


class _RoundedTable(RLTable):  # type: ignore[misc]
    def __init__(
        self,
        *args,
        corner_radius: float = 6.0,
        border_width: float = 0.8,
        border_color=rl_colors.HexColor("#1b1b1d") if rl_colors else None,
        fill_color=rl_colors.HexColor("#ffffff") if rl_colors else None,
        **kwargs,
    ) -> None:
        super().__init__(*args, **kwargs)
        self._corner_radius = corner_radius
        self._border_width = border_width
        self._border_color = border_color
        self._fill_color = fill_color

    def draw(self) -> None:  # pragma: no cover - visual rendering
        canvas = self.canv
        canvas.saveState()
        stroke = 0
        if self._border_color is not None and self._border_width > 0:
            canvas.setStrokeColor(self._border_color)
            canvas.setLineWidth(self._border_width)
            stroke = 1
        fill = 0
        if self._fill_color is not None:
            canvas.setFillColor(self._fill_color)
            fill = 1
        canvas.roundRect(0, 0, self._width, self._height, self._corner_radius, stroke=stroke, fill=fill)
        canvas.restoreState()
        super().draw()


def _discover_logo_flowable(max_width: float = 4.2 * cm) -> Optional[Any]:
    if Image is None:
        return None
    candidates = [
        PROJECT_ROOT / "assets" / "digitx_logo.png",
        Path("assets") / "digitx_logo.png",
    ]
    for candidate in candidates:
        if candidate.exists():
            logo = Image(str(candidate))  # type: ignore[call-arg]
            if getattr(logo, "imageWidth", max_width) > max_width:
                scale = max_width / float(logo.imageWidth)
                logo.drawWidth = max_width
                logo.drawHeight = float(logo.imageHeight) * scale
            logo.hAlign = "LEFT"
            return logo
    return None


def _wrap_text_card(content: Paragraph, width: float) -> _RoundedTable:
    card = _RoundedTable(
        [[content]],
        colWidths=[width],
        corner_radius=0.25 * cm,
        border_width=0.4,
        border_color=rl_colors.HexColor("#d7d7da") if rl_colors else None,
        fill_color=rl_colors.HexColor("#ffffff") if rl_colors else None,
    )
    card.setStyle(
        TableStyle(
            [
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    return card


def _write_pdf(result: StandardizeResult, path: Path, snapshot: OrderedDict[str, str], sections: List[ReportSection]) -> Path:
    _ensure_pdf_dependencies()

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2.2 * cm,
        bottomMargin=2 * cm,
    )

    palette_primary = rl_colors.HexColor("#0f0f10")  # type: ignore[call-arg]
    palette_muted = rl_colors.HexColor("#6a6a6d")  # type: ignore[call-arg]
    palette_surface = rl_colors.HexColor("#ffffff")  # type: ignore[call-arg]
    palette_surface_alt = rl_colors.HexColor("#f0f1f5")  # type: ignore[call-arg]

    styles = getSampleStyleSheet()  # type: ignore[operator]
    styles.add(ParagraphStyle(name="StdHeaderTitle", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=palette_primary, alignment=2))
    styles.add(ParagraphStyle(name="StdHeaderMeta", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=14, textColor=palette_muted, alignment=2))
    styles.add(ParagraphStyle(name="StdSectionTitle", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=12, leading=16, textColor=palette_primary))
    styles.add(ParagraphStyle(name="StdTableCell", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=13, textColor=palette_primary))
    styles.add(ParagraphStyle(name="StdBodyText", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=14, textColor=palette_primary))
    styles.add(ParagraphStyle(name="StdSnapshot", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=14, textColor=palette_primary))

    story: List[Any] = []

    logo_flowable = _discover_logo_flowable()
    doc_label = result.document_path.name if result.document_path is not None else "—"
    meta_paragraph = Paragraph(
        "<br/>".join(
            [
                f"Schema: {escape(result.schema_name)}",
                f"Model: {escape(result.model_name)}",
                f"Document: {escape(doc_label)}",
                f"Generated: {escape(result.generated_at.isoformat().replace('+00:00', 'Z'))}",
                f"MOSAICX v{APPLICATION_VERSION}",
            ]
        ),
        styles["StdHeaderMeta"],
    )
    left_cell: Any = logo_flowable if logo_flowable is not None else Paragraph("DIGIT-X Lab", styles["StdHeaderMeta"])
    header_rows = [[left_cell, Paragraph("MOSAICX: Structured Report", styles["StdHeaderTitle"])], ["", meta_paragraph]]
    header_table = _RoundedTable(
        header_rows,
        colWidths=[doc.width * 0.35, doc.width * 0.65],
        corner_radius=0.35 * cm,
        border_width=0.8,
        border_color=palette_primary,
        fill_color=palette_surface,
    )
    header_table.setStyle(
        TableStyle(
            [
                ("SPAN", (0, 0), (0, 1)),
                ("VALIGN", (0, 0), (0, 1), "MIDDLE"),
                ("VALIGN", (1, 0), (1, 1), "BOTTOM"),
                ("ALIGN", (1, 0), (1, 1), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 14),
                ("RIGHTPADDING", (0, 0), (-1, -1), 14),
                ("TOPPADDING", (0, 0), (-1, -1), 16),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 16),
            ]
        )
    )
    story.append(header_table)

    band = RLTable([[""]], colWidths=[doc.width])
    band.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), palette_surface_alt), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
    band._argH = [0.1 * cm]  # type: ignore[attr-defined]
    story.append(band)
    story.append(Spacer(1, 0.35 * cm))

    if snapshot:
        snapshot_cells: List[Paragraph] = []
        for label, value in snapshot.items():
            html = (
                f"<font color='{palette_muted}' size='8'>{escape(label)}</font><br/>"
                f"<font color='{palette_primary}' size='11'>{escape(value)}</font>"
            )
            snapshot_cells.append(Paragraph(html, styles["StdSnapshot"]))
        column_count = min(3, max(1, len(snapshot_cells)))
        col_width = doc.width / column_count
        rows: List[List[Any]] = []
        for i in range(0, len(snapshot_cells), column_count):
            row = snapshot_cells[i : i + column_count]
            while len(row) < column_count:
                row.append(Paragraph("&nbsp;", styles["StdSnapshot"]))
            rows.append(row)
        snapshot_table = _RoundedTable(
            rows,
            colWidths=[col_width] * column_count,
            corner_radius=0.25 * cm,
            border_width=0.4,
            border_color=palette_muted,
            fill_color=palette_surface,
        )
        snapshot_table.setStyle(TableStyle([( "LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
        story.append(snapshot_table)
        story.append(Spacer(1, 0.4 * cm))

    for section in sections:
        if section.title:
            story.append(Paragraph(escape(section.title.upper()), styles["StdSectionTitle"]))
            story.append(Spacer(1, 0.12 * cm))

        cards: List[_RoundedTable] = []
        if section.kind == "kv":
            kv_rows = []
            for label, value in section.data.items():
                kv_rows.append([
                    Paragraph(f"<b>{escape(label)}</b>", styles["StdTableCell"]),
                    Paragraph(escape(value).replace("\n", "<br/>"), styles["StdTableCell"]),
                ])
            cards.append(
                _RoundedTable(
                    kv_rows,
                    colWidths=[doc.width * 0.32, doc.width * 0.68],
                    corner_radius=0.25 * cm,
                    border_width=0.4,
                    border_color=palette_muted,
                    fill_color=palette_surface,
                )
            )
        elif section.kind == "bullets":
            bullet_html = "<br/>".join(f"• {escape(item)}" for item in section.data)
            cards.append(_wrap_text_card(Paragraph(bullet_html, styles["StdBodyText"]), doc.width))
        elif section.kind == "text":
            cards.append(_wrap_text_card(Paragraph(escape(str(section.data)).replace("\n", "<br/>"), styles["StdBodyText"]), doc.width))

        for card in cards:
            card.setStyle(TableStyle([( "LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6), ("GRID", (0, 0), (-1, -1), 0.15, palette_surface_alt)]))
            story.append(card)
        story.append(Spacer(1, 0.3 * cm))

    footer_text = (
        f"Generated with MOSAICX v{APPLICATION_VERSION} • Schema: {result.schema_name} • "
        f"Document: {doc_label}"
    )

    def _draw_footer(canvas, doc_obj) -> None:
        canvas.saveState()
        margin = 1.2 * cm
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(palette_muted)
        canvas.drawString(doc_obj.leftMargin, margin, footer_text)
        canvas.restoreState()

    doc.build(story, onFirstPage=_draw_footer, onLaterPages=_draw_footer)
    return path


def _write_text(result: StandardizeResult, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(_render_text_report(result), encoding="utf-8")
    return path


def _write_docx(result: StandardizeResult, path: Path, snapshot: OrderedDict[str, str], sections: List[ReportSection]) -> Path:
    if Document is None or Pt is None:
        raise RuntimeError("DOCX export unavailable – install python-docx (`pip install python-docx`).")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    title = doc.add_heading("MOSAICX Structured Report", level=1)
    title.alignment = 1

    doc_label = result.document_path.name if result.document_path is not None else "—"
    meta_paragraph = doc.add_paragraph()
    meta_lines = [
        f"Schema: {result.schema_name}",
        f"Model: {result.model_name}",
        f"Document: {doc_label}",
        f"Generated: {result.generated_at.isoformat().replace('+00:00', 'Z')}",
        f"MOSAICX v{APPLICATION_VERSION}",
    ]
    meta_paragraph.add_run("\n".join(meta_lines))

    if snapshot:
        items = list(snapshot.items())
        col_count = 2 if len(items) > 1 else 1
        table = doc.add_table(rows=1, cols=col_count)
        table.style = "Light Grid"
        table.autofit = True
        row_idx = 0
        for idx, (label, value) in enumerate(items):
            if idx % col_count == 0 and idx != 0:
                row_idx += 1
                table.add_row()
            cell = table.rows[row_idx].cells[idx % col_count]
            cell.text = f"{label}: {value}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # type: ignore[union-attr]

    for section in sections:
        if section.title:
            doc.add_heading(section.title.upper(), level=2)
        if section.kind == "kv":
            kv_table = doc.add_table(rows=1, cols=2)
            kv_table.style = "Light Grid"
            kv_table.rows[0].cells[0].text = "Field"
            kv_table.rows[0].cells[1].text = "Value"
            for label, value in section.data.items():
                row_cells = kv_table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = value
            for row in kv_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)  # type: ignore[union-attr]
        elif section.kind == "bullets":
            for item in section.data:
                doc.add_paragraph(item, style="List Bullet")
        elif section.kind == "text":
            for line in str(section.data).splitlines():
                doc.add_paragraph(line)

    doc.save(path)
    return path


def standardize_document(
    document_path: Optional[Union[str, Path]],
    schema_path: Union[str, Path],
    *,
    record_data: Optional[Union[BaseModel, Dict[str, Any]]] = None,
    model: str = DEFAULT_LLM_MODEL,
    base_url: Optional[str] = None,
    api_key: Optional[str] = None,
    temperature: float = 0.0,
    generate_narrative: bool = False,
    narrative_temperature: Optional[float] = None,
    status_callback: Optional[Callable[[str], None]] = None,
) -> StandardizeResult:
    schema_path = Path(schema_path)
    schema_class = load_schema_model(str(schema_path))
    doc_path = Path(document_path) if document_path is not None else None

    record: BaseModel
    extracted_text: Optional[str] = None

    if record_data is not None:
        if status_callback:
            status_callback("Validating supplied structured data…")
        if isinstance(record_data, BaseModel):
            if isinstance(record_data, schema_class):
                record = record_data
            else:
                record = schema_class.model_validate(record_data.model_dump(by_alias=True))
        else:
            record = schema_class.model_validate(record_data)
    else:
        if doc_path is None:
            raise ExtractionError("Provide either a document path or pre-extracted record_data.")
        extraction = extract_text_from_document(doc_path, return_details=True, status_callback=status_callback)
        extracted_text = extraction.markdown
        record = extract_structured_data(
            extracted_text,
            schema_class,
            model=model,
            base_url=base_url,
            api_key=api_key,
            temperature=temperature,
        )

    record_dict = record.model_dump(by_alias=True)
    snapshot_data, structured_sections = _collect_sections_from_payload(record_dict)
    narrative_text: Optional[str] = None
    narrative_error: Optional[str] = None
    narrative_sections: Optional[List[ReportSection]] = None
    if generate_narrative:
        if status_callback:
            status_callback("Generating narrative summary…")
        narrative_temp = (
            narrative_temperature
            if narrative_temperature is not None
            else (0.2 if temperature == 0.0 else temperature)
        )
        try:
            narrative_sections = _generate_narrative(
                schema_class,
                record_dict,
                structured_sections,
                model=model,
                base_url=base_url,
                api_key=api_key,
                temperature=narrative_temp,
            )
            narrative_text = "\n\n".join(section.data for section in narrative_sections)
        except Exception as exc:  # pragma: no cover - best-effort
            narrative_error = str(exc)
        finally:
            if status_callback:
                status_callback("Finalizing structured report…")

    generated_at = datetime.now(tz=timezone.utc)
    schema_name = schema_class.__name__
    return StandardizeResult(
        record=record,
        schema_path=schema_path,
        document_path=doc_path,
        schema_name=schema_name,
        model_name=model,
        generated_at=generated_at,
        extracted_text=extracted_text,
        narrative=narrative_text,
        narrative_error=narrative_error,
        snapshot_data=snapshot_data,
        structured_sections=structured_sections,
        narrative_sections=narrative_sections,
    )


def save_standardize_artifacts(
    result: StandardizeResult,
    *,
    artifacts: Iterable[str],
    json_path: Optional[Path],
    pdf_path: Optional[Path],
    text_path: Optional[Path],
    docx_path: Optional[Path],
    out_dir: Optional[Path] = None,
    emit_messages: bool = True,
) -> Dict[str, Path]:
    selection = tuple(dict.fromkeys(a.lower() for a in artifacts if a))
    if not selection:
        return {}

    allowed = {"json", "pdf", "txt", "text", "docx"}
    invalid = set(selection) - allowed
    if invalid:
        raise ValueError(f"Unsupported artifact type(s): {', '.join(sorted(invalid))}.")

    resolved_out_dir = Path(out_dir) if out_dir is not None else None
    saved: Dict[str, Path] = {}
    snapshot, sections = _prepare_sections(result)

    for artifact in selection:
        if artifact == "json":
            target = Path(json_path) if json_path else result.default_path(extension="json", directory=resolved_out_dir)
            result.write_json(target)
            saved["json"] = target
            if emit_messages:
                styled_message(f"Saved JSON: {target}", "accent")
        elif artifact in {"txt", "text"}:
            target = Path(text_path) if text_path else result.default_path(extension="txt", directory=resolved_out_dir)
            _write_text(result, target)
            saved["txt"] = target
            if emit_messages:
                styled_message(f"Saved text: {target}", "accent")
        elif artifact == "pdf":
            target = Path(pdf_path) if pdf_path else result.default_path(extension="pdf", directory=resolved_out_dir)
            try:
                _write_pdf(result, target, snapshot, sections)
            except RuntimeError as exc:
                if emit_messages:
                    styled_message(str(exc), "muted")
            else:
                saved["pdf"] = target
                if emit_messages:
                    styled_message(f"Saved PDF: {target}", "accent")
        elif artifact == "docx":
            target = Path(docx_path) if docx_path else result.default_path(extension="docx", directory=resolved_out_dir)
            try:
                _write_docx(result, target, snapshot, sections)
            except RuntimeError as exc:
                if emit_messages:
                    styled_message(str(exc), "muted")
            else:
                saved["docx"] = target
                if emit_messages:
                    styled_message(f"Saved DOCX: {target}", "accent")

    return saved


__all__ = [
    "StandardizeResult",
    "standardize_document",
    "save_standardize_artifacts",
]
