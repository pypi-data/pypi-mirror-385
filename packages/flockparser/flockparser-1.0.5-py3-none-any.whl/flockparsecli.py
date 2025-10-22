import ollama
from pathlib import Path
from PyPDF2 import PdfReader
import docx
import subprocess
import tempfile
import json
import numpy as np
from datetime import datetime
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import socket
import requests
import chromadb
from sollol.vram_monitor import VRAMMonitor, monitor_distributed_nodes
from gpu_controller import GPUController
from sollol.intelligent_gpu_router import IntelligentGPURouter
from sollol.adaptive_parallelism import AdaptiveParallelismStrategy
from logging_config import setup_logging
from sollol import OllamaPool  # Direct SOLLOL integration
from sollol_compat import add_flockparser_methods  # FlockParser compatibility layer

# Initialize logging
logger = setup_logging()

# 🚀 AVAILABLE COMMANDS:
COMMANDS = """
   📖 open_pdf <file>   → Process a single PDF file
   📂 open_dir <dir>    → Process all PDFs in a directory
   💬 chat              → Chat with processed PDFs
   📊 list_docs         → List all processed documents
   🔍 check_deps        → Check for required dependencies
   🌐 discover_nodes    → Auto-discover Ollama nodes on local network
   ➕ add_node <url>    → Manually add an Ollama node (e.g., http://192.168.1.100:11434)
   ➖ remove_node <url> → Remove an Ollama node from the pool
   📋 list_nodes        → List all configured Ollama nodes
   🔬 verify_models     → Check which models are available on each node
   ⚖️  lb_stats          → Show load balancer statistics
   🎯 set_routing <strategy> → Set routing: adaptive, round_robin, least_loaded, lowest_latency
   🖥️  vram_report       → Show detailed VRAM usage report
   🚀 force_gpu <model> → Force model to GPU on all capable nodes
   🎯 gpu_status        → Show intelligent GPU routing status
   🧠 gpu_route <model> → Show routing decision for a model
   🔧 gpu_optimize      → Trigger intelligent GPU optimization
   ✅ gpu_check <model> → Check which nodes can fit a model
   📚 gpu_models        → List all known models and sizes
   🗑️  unload_model <model> → Unload a specific model from memory
   🧹 cleanup_models    → Unload all non-priority models
   🔀 parallelism_report → Show adaptive parallelism analysis
   🧹 clear_cache       → Clear embedding cache (keeps documents)
   🗑️  clear_db          → Clear ChromaDB vector store (removes all documents)
   ❌ exit              → Quit the program

   🌐 API Server: Automatically starts on port 8000 (http://localhost:8000)
      Configure with: FLOCKPARSER_API=true/false, FLOCKPARSER_API_PORT=8000
"""

# 🔥 AI MODELS
EMBEDDING_MODEL = "mxbai-embed-large"
CHAT_MODEL = "qwen3:8b"  # Fast and fits in available RAM (5.2 GB)

# 🚀 MODEL CACHING CONFIGURATION
# Keep models in VRAM for faster inference (prevents reloading)
EMBEDDING_KEEP_ALIVE = "1h"  # Embedding model used frequently for chunking/search
CHAT_KEEP_ALIVE = "15m"  # Chat model used less frequently

# 📊 RAG CONFIGURATION
# Retrieval settings for chat
RETRIEVAL_TOP_K = 10  # Number of chunks to retrieve (default: 10)
RETRIEVAL_MIN_SIMILARITY = 0.3  # Minimum similarity score (0.0-1.0)
CHUNKS_TO_SHOW = 10  # Number of source chunks to display (show all retrieved)

# Acceptable model variations (allows flexible matching)
ACCEPTABLE_EMBEDDING_MODELS = [
    "mxbai-embed-large",
    "mxbai-embed-large:latest",
    "nomic-embed-text",
    "nomic-embed-text:latest",
    "all-minilm",
    "all-minilm:latest",
    "bge-large",
    "bge-large:latest",
]

ACCEPTABLE_CHAT_MODELS = [
    "llama3.1",
    "llama3.1:8b",
    "llama3.2",
    "llama3.2:latest",
    "llama3.2:3b",
    "llama3",
    "llama3:latest",
    "llama3:8b",
    "mistral",
    "mistral:latest",
    "mixtral",
    "mixtral:latest",
    "qwen",
    "qwen:latest",
    "qwen2.5",
    "qwen3",
    "qwen3:14b",
    "qwen3:8b",
    "qwen3:4b",
    "gemma2:9b",
    "phi3",
    "deepseek-coder-v2",
    "codellama:13b",
]

# 🌐 OLLAMA LOAD BALANCER CONFIGURATION
# SOLLOL auto-discovers all Ollama nodes on the network

# 📁 Directory setup
_SCRIPT_DIR = Path(__file__).parent
PROCESSED_DIR = _SCRIPT_DIR / "converted_files"
PROCESSED_DIR.mkdir(exist_ok=True)

# 📚 Knowledge Base (legacy JSON storage - kept for backwards compatibility)
KB_DIR = _SCRIPT_DIR / "knowledge_base"
KB_DIR.mkdir(exist_ok=True)

# 🗄️ ChromaDB Vector Store (production storage)
CHROMA_DB_DIR = _SCRIPT_DIR / "chroma_db_cli"
CHROMA_DB_DIR.mkdir(exist_ok=True)

# Initialize ChromaDB client and collection
chroma_client = chromadb.PersistentClient(path=str(CHROMA_DB_DIR))
chroma_collection = chroma_client.get_or_create_collection(
    name="documents", metadata={"hnsw:space": "cosine"}  # Use cosine similarity for better semantic search
)

# ============================================================
# SOLLOL Direct Integration (replaces ~1100 lines of custom code)
# Pure SOLLOL OllamaPool - no adapter layer
# Original implementation backed up to /tmp/old_loadbalancer_backup.py
# ============================================================

# Global reference (will be initialized in setup_load_balancer())
load_balancer = None

def setup_load_balancer():
    """Initialize SOLLOL pool with auto-discovery and dashboard.

    Must be called from within if __name__ == '__main__': to avoid
    multiprocessing issues with Dask worker spawning.
    """
    global load_balancer

    # Initialize SOLLOL pool with auto-discovery FIRST
    load_balancer = OllamaPool(
        nodes=None,  # Auto-discover all Ollama nodes on network
        enable_intelligent_routing=True,
        exclude_localhost=True,  # Use real IP instead of localhost
        discover_all_nodes=True,  # Scan full network for all nodes
        app_name="FlockParser",  # Identify as FlockParser in dashboard
        enable_ray=True,  # Start Ray cluster for multi-app coordination
        register_with_dashboard=False  # Don't auto-register yet - dashboard not running
    )

    # Add FlockParser compatibility methods
    load_balancer = add_flockparser_methods(load_balancer, KB_DIR)

    return load_balancer

# Dashboard configuration (read from environment)
import os
_dashboard_enabled = os.getenv("FLOCKPARSER_DASHBOARD", "true").lower() in ("true", "1", "yes", "on")
_dashboard_port = int(os.getenv("FLOCKPARSER_DASHBOARD_PORT", "8080"))

def setup_dashboard():
    """Start SOLLOL unified dashboard after pool creation.

    Must be called from within if __name__ == '__main__': to avoid
    multiprocessing issues with Dask worker spawning.
    """
    global load_balancer

    if not _dashboard_enabled:
        logger.info("📊 Dashboard disabled (set FLOCKPARSER_DASHBOARD=true to enable)")
        return

    # Install Redis log publisher for distributed log streaming
    from sollol.dashboard_service import install_redis_log_publisher
    try:
        install_redis_log_publisher()
        logger.info("📡 Redis log publisher installed - logs streaming to dashboard")
    except Exception as e:
        logger.warning(f"⚠️  Redis log publisher failed: {e}")

    # Check if dashboard already running, if not start dashboard_service
    import requests
    import subprocess
    import threading
    import time

    dashboard_url = f"http://localhost:{_dashboard_port}"
    dashboard_running = False

    try:
        response = requests.get(f"{dashboard_url}/api/applications", timeout=1)
        if response.status_code == 200:
            dashboard_running = True
            logger.info(f"✅ Dashboard already running at {dashboard_url}")
    except requests.exceptions.RequestException:
        logger.info(f"🚀 Starting dashboard service at {dashboard_url}")

    if not dashboard_running:
        # Start dashboard_service as subprocess (not daemon thread)
        dashboard_proc = subprocess.Popen([
            "python3", "-m", "sollol.dashboard_service",
            "--port", str(_dashboard_port),
            "--redis-url", "redis://localhost:6379",
            "--ray-dashboard-port", "8265",
            "--dask-dashboard-port", "8787",
        ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        # Wait for dashboard to start
        for attempt in range(10):
            time.sleep(0.5)
            try:
                response = requests.get(f"{dashboard_url}/api/applications", timeout=1)
                if response.status_code == 200:
                    logger.info(f"✅ Dashboard service started at {dashboard_url}")
                    dashboard_running = True
                    break
            except requests.exceptions.RequestException:
                continue

        if not dashboard_running:
            logger.error("❌ Dashboard failed to start - check logs for errors")

    # Register FlockParser with dashboard
    if dashboard_running:
        try:
            from sollol.dashboard_client import DashboardClient
            load_balancer._dashboard_client = DashboardClient(
                app_name="FlockParser",
                router_type="OllamaPool",
                dashboard_url=dashboard_url,
                auto_register=True
            )
            logger.info("✅ FlockParser registered with dashboard")
            logger.info(f"📊 Dashboard: {dashboard_url}")
            logger.info(f"   - Ray Dashboard: http://localhost:8265")
            logger.info(f"   - Dask Dashboard: http://localhost:8787")
        except Exception as e:
            logger.warning(f"⚠️  Dashboard registration failed: {e}")

# 💾 Index file for tracking processed documents
INDEX_FILE = KB_DIR / "document_index.json"

# 🔄 Cache for embeddings to avoid regenerating
EMBEDDING_CACHE_FILE = KB_DIR / "embedding_cache.json"


def load_embedding_cache():
    """Load the embedding cache from disk."""
    if not EMBEDDING_CACHE_FILE.exists():
        return {}
    try:
        with open(EMBEDDING_CACHE_FILE, "r") as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return {}


def save_embedding_cache(cache):
    """Save the embedding cache to disk."""
    with open(EMBEDDING_CACHE_FILE, "w") as f:
        json.dump(cache, f)


def get_cached_embedding(text, use_load_balancer=True):
    """Get embedding from cache or generate new one."""
    import hashlib

    cache = load_embedding_cache()

    # Create hash of text for cache key
    text_hash = hashlib.md5(text.encode()).hexdigest()

    if text_hash in cache:
        return cache[text_hash]

    # Generate new embedding using load balancer
    if use_load_balancer:
        embedding_result = load_balancer.embed(EMBEDDING_MODEL, text, keep_alive=EMBEDDING_KEEP_ALIVE, priority=7)
    else:
        embedding_result = ollama.embed(model=EMBEDDING_MODEL, input=text, keep_alive=EMBEDDING_KEEP_ALIVE)

    embeddings = embedding_result.get("embeddings", [])
    embedding = embeddings[0] if embeddings else []

    # Cache it
    cache[text_hash] = embedding
    save_embedding_cache(cache)

    return embedding


def load_document_index():
    """Load the document index or create it if it doesn't exist."""
    if not INDEX_FILE.exists():
        return {"documents": []}

    try:
        with open(INDEX_FILE, "r") as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        logger.error("⚠️ Error loading index file. Creating a new one.")
        return {"documents": []}


def save_document_index(index_data):
    """Save the document index to disk."""
    with open(INDEX_FILE, "w") as f:
        json.dump(index_data, f, indent=4)
    logger.info(f"✅ Document index updated with {len(index_data['documents'])} documents")


def register_document(pdf_path, txt_path, content, chunks=None):
    """Register a processed document in the knowledge base index."""
    # Load existing index
    index_data = load_document_index()

    # Create document record
    document_id = f"doc_{len(index_data['documents']) + 1}"

    # Get PDF filename for better logging (especially in parallel mode)
    from pathlib import Path
    pdf_name = Path(pdf_path).stem if pdf_path else "unknown"

    # Generate embeddings and chunks for search
    chunks = chunks or chunk_text(content)
    chunk_embeddings = []

    # Batch process embeddings for better performance
    logger.info(f"🔄 [{pdf_name}] Processing {len(chunks)} chunks in batches...")

    import hashlib

    cache = load_embedding_cache()
    uncached_chunks = []
    uncached_indices = []

    # Check cache first
    for i, chunk in enumerate(chunks):
        text_hash = hashlib.md5(chunk.encode()).hexdigest()
        if text_hash not in cache:
            uncached_chunks.append(chunk)
            uncached_indices.append(i)

    # Batch embed uncached chunks using load balancer
    if uncached_chunks:
        logger.info(f"🚀 [{pdf_name}] Embedding {len(uncached_chunks)} new chunks in parallel...")
        logger.info(f"   Using {len(load_balancer.instances)} Ollama nodes")

        # Process in batches of 100 to save cache periodically
        batch_size = 100
        for batch_start in range(0, len(uncached_chunks), batch_size):
            batch_end = min(batch_start + batch_size, len(uncached_chunks))
            batch = uncached_chunks[batch_start:batch_end]

            batch_num = batch_start // batch_size + 1
            total_batches = (len(uncached_chunks) + batch_size - 1) // batch_size
            logger.info(f"   [{pdf_name}] Batch {batch_num}/{total_batches}...")

            batch_results = load_balancer.embed_batch(EMBEDDING_MODEL, batch)

            # Cache the batch embeddings
            cached_count = 0
            for chunk, result in zip(batch, batch_results):
                if result:
                    text_hash = hashlib.md5(chunk.encode()).hexdigest()
                    embeddings = result.get("embeddings", [])
                    embedding = embeddings[0] if embeddings else []
                    cache[text_hash] = embedding
                    cached_count += 1

            # Save cache after each batch
            save_embedding_cache(cache)
            logger.info(f"   [{pdf_name}] ✅ Cached {cached_count} embeddings from this batch")

        logger.info(f"✅ [{pdf_name}] All {len(uncached_chunks)} new embeddings cached")
    else:
        logger.info(f"✅ [{pdf_name}] All chunks found in cache!")

    # Now process all chunks
    for i, chunk in enumerate(chunks):
        try:
            # Show progress every 50 chunks
            if i % 50 == 0 and i > 0:
                logger.info(f"🔄 Processed {i}/{len(chunks)} chunks...")

            # Get embedding from cache
            text_hash = hashlib.md5(chunk.encode()).hexdigest()
            embedding = cache.get(text_hash, [])

            # Store chunk with its embedding
            chunk_file = KB_DIR / f"{document_id}_chunk_{i}.json"
            chunk_data = {"text": chunk, "embedding": embedding}

            with open(chunk_file, "w") as f:
                json.dump(chunk_data, f)

            # Remember the chunk reference
            chunk_embeddings.append({"chunk_id": f"{document_id}_chunk_{i}", "file": str(chunk_file)})
        except Exception as e:
            logger.error(f"⚠️ Error embedding chunk {i}: {e}")

    # Add document to index
    doc_entry = {
        "id": document_id,
        "original": str(pdf_path),
        "text_path": str(txt_path),
        "processed_date": datetime.now().isoformat(),
        "chunks": chunk_embeddings,
    }

    index_data["documents"].append(doc_entry)
    save_document_index(index_data)
    return document_id


def chunk_text(text, chunk_size=512, overlap=100):
    """
    Split text into overlapping chunks with intelligent token-aware splitting.

    Args:
        chunk_size: Target chunk size in tokens (approximate via chars * 0.25)
        overlap: Number of characters to overlap between chunks
    """
    # Token limits for mxbai-embed-large: 512 tokens max
    # Rough estimate: 1 token ≈ 4 characters
    MAX_TOKENS = 480  # Leave buffer for model
    MAX_CHARS = MAX_TOKENS * 4  # ~1920 chars
    TARGET_CHARS = chunk_size * 4  # ~2048 chars for chunk_size=512

    def split_large_text(text, max_size):
        """Recursively split text that's too large."""
        if len(text) <= max_size:
            return [text]

        # Try splitting by sentences first
        sentences = text.replace("! ", "!|").replace("? ", "?|").replace(". ", ".|").split("|")

        chunks = []
        current = []
        current_len = 0

        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue

            # If single sentence exceeds limit, split by words
            if len(sent) > max_size:
                words = sent.split()
                # Calculate words per chunk (with safety margin)
                words_per_chunk = int((max_size / len(sent)) * len(words) * 0.9)
                words_per_chunk = max(50, words_per_chunk)  # At least 50 words

                for i in range(0, len(words), words_per_chunk):
                    word_chunk = " ".join(words[i: i + words_per_chunk])
                    if word_chunk:
                        chunks.append(word_chunk)
                continue

            # Add sentence to current chunk
            if current_len + len(sent) > max_size and current:
                chunks.append(" ".join(current))
                current = [sent]
                current_len = len(sent)
            else:
                current.append(sent)
                current_len += len(sent)

        if current:
            chunks.append(" ".join(current))

        return chunks

    # Split into paragraphs first
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    chunks = []
    current_chunk = []
    current_length = 0

    for para in paragraphs:
        para_len = len(para)

        # If paragraph is too large, split it first
        if para_len > MAX_CHARS:
            # Finalize current chunk if any
            if current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = []
                current_length = 0

            # Split the large paragraph
            para_chunks = split_large_text(para, MAX_CHARS)
            chunks.extend(para_chunks)
            continue

        # Check if adding this paragraph exceeds target size
        if current_length + para_len > TARGET_CHARS and current_chunk:
            # Finalize current chunk
            chunks.append("\n\n".join(current_chunk))

            # Start new chunk with overlap (keep last paragraph if small enough)
            if overlap > 0 and current_chunk and len(current_chunk[-1]) < overlap:
                current_chunk = [current_chunk[-1], para]
                current_length = len(current_chunk[-1]) + para_len
            else:
                current_chunk = [para]
                current_length = para_len
        else:
            current_chunk.append(para)
            current_length += para_len

    # Add final chunk
    if current_chunk:
        final_chunk = "\n\n".join(current_chunk)
        # Safety check
        if len(final_chunk) > MAX_CHARS:
            chunks.extend(split_large_text(final_chunk, MAX_CHARS))
        else:
            chunks.append(final_chunk)

    # Final validation: ensure no chunk exceeds MAX_CHARS
    validated_chunks = []
    for chunk in chunks:
        if len(chunk) > MAX_CHARS:
            validated_chunks.extend(split_large_text(chunk, MAX_CHARS))
        else:
            validated_chunks.append(chunk)

    return validated_chunks


def list_documents():
    """List all processed documents in the knowledge base."""
    index_data = load_document_index()
    if not index_data["documents"]:
        logger.info("📚 No documents have been processed yet.")
        return

    logger.info(f"\n📚 Knowledge Base: {len(index_data['documents'])} documents")
    logger.info("-" * 60)
    for i, doc in enumerate(index_data["documents"]):
        logger.info(f"{i+1}. {Path(doc['original']).name}")
        logger.info(f"   ID: {doc['id']} | Processed: {doc['processed_date'][:10]}")
        logger.info(f"   Chunks: {len(doc['chunks'])}")
        logger.info("-" * 60)


def get_similar_chunks(query, top_k=None, min_similarity=None):
    """Find text chunks similar to the query using vector similarity with adaptive top-k."""
    # Use configured defaults if not specified
    if min_similarity is None:
        min_similarity = RETRIEVAL_MIN_SIMILARITY

    try:
        # Get embedding for the query from cache
        query_embedding = get_cached_embedding(query)

        if not query_embedding:
            logger.error("⚠️ Failed to generate query embedding")
            return []

        # Load document index
        index_data = load_document_index()

        # Check if we have documents
        if not index_data["documents"]:
            logger.info("📚 No documents in knowledge base yet")
            return []

        # Adaptive top-k based on total chunks in database
        if top_k is None:
            total_chunks = sum(len(doc["chunks"]) for doc in index_data["documents"])

            # Scale top_k based on database size
            if total_chunks < 50:
                adaptive_k = min(total_chunks, 5)  # Very small DB, use fewer
            elif total_chunks < 200:
                adaptive_k = 10  # Small-medium DB, use default
            elif total_chunks < 1000:
                adaptive_k = 20  # Medium DB, retrieve more context
            else:
                adaptive_k = 30  # Large DB, need more chunks for good coverage

            top_k = adaptive_k
            logger.info(f"   📊 Adaptive top-k: {top_k} (from {total_chunks} total chunks)")
        else:
            logger.info(f"   📊 Using fixed top-k: {top_k}")

        # Collect all chunks with their embeddings
        chunks_with_similarity = []

        for doc in index_data["documents"]:
            for chunk_ref in doc["chunks"]:
                try:
                    # Load chunk data
                    chunk_file = Path(chunk_ref["file"])
                    if chunk_file.exists():
                        with open(chunk_file, "r") as f:
                            chunk_data = json.load(f)

                        # Calculate cosine similarity
                        chunk_embedding = chunk_data.get("embedding", [])
                        if chunk_embedding:
                            similarity = cosine_similarity(query_embedding, chunk_embedding)
                            if similarity >= min_similarity:
                                chunks_with_similarity.append(
                                    {
                                        "doc_id": doc["id"],
                                        "doc_name": Path(doc["original"]).name,
                                        "text": chunk_data["text"],
                                        "similarity": similarity,
                                    }
                                )
                except Exception as e:
                    logger.error(f"⚠️ Error processing chunk {chunk_ref['chunk_id']}: {e}")

        # Sort by similarity (highest first) and get top k
        chunks_with_similarity.sort(key=lambda x: x["similarity"], reverse=True)

        # Return top k results
        results = chunks_with_similarity[:top_k]

        # Print retrieval stats
        logger.info(f"   Found {len(results)} relevant chunks (similarity >= {min_similarity:.2f})")

        return results

    except Exception as e:
        logger.error(f"⚠️ Error searching knowledge base: {e}")
        return []


def sanitize_for_xml(text):
    """Remove null bytes and control characters that break XML/DOCX."""
    import re

    # Remove NULL bytes
    text = text.replace("\x00", "")
    # Remove other control characters except newline, carriage return, and tab
    text = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]", "", text)
    return text


def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors."""
    if not vec1 or not vec2:
        return 0

    vec1 = np.array(vec1)
    vec2 = np.array(vec2)

    dot_product = np.dot(vec1, vec2)
    norm_a = np.linalg.norm(vec1)
    norm_b = np.linalg.norm(vec2)

    if norm_a == 0 or norm_b == 0:
        return 0

    return dot_product / (norm_a * norm_b)


def embed_text(text):
    """Embeds text using Ollama without storing vector data in files."""
    try:
        # Using 'input' instead of 'prompt'
        _ = ollama.embed(model=EMBEDDING_MODEL, input=text)
        return text  # Return the original text for saving to files
    except Exception as e:
        logger.error(f"❌ Embedding error: {e}")
        return None


def clean_extracted_text(text):
    """Clean extracted text by normalizing Unicode and fixing common LaTeX/PDF extraction issues."""
    import re
    import unicodedata

    if not text:
        return text

    # Step 1: Normalize Unicode (convert composed chars to decomposed and back)
    text = unicodedata.normalize('NFKC', text)

    # Step 2: Fix common Unicode escape sequences that appear as literal text
    # Replace \uXXXX patterns with actual Unicode characters
    def replace_unicode_escapes(match):
        try:
            code = match.group(1)
            return chr(int(code, 16))
        except:
            return match.group(0)

    text = re.sub(r'\\u([0-9a-fA-F]{4})', replace_unicode_escapes, text)
    text = re.sub(r'\\x([0-9a-fA-F]{2})', replace_unicode_escapes, text)

    # Step 3: Clean up common LaTeX remnants that get corrupted
    # Replace common Greek letter codes with their actual Unicode
    greek_map = {
        r'\\alpha': 'α', r'\\beta': 'β', r'\\gamma': 'γ', r'\\delta': 'δ',
        r'\\epsilon': 'ε', r'\\zeta': 'ζ', r'\\eta': 'η', r'\\theta': 'θ',
        r'\\iota': 'ι', r'\\kappa': 'κ', r'\\lambda': 'λ', r'\\mu': 'μ',
        r'\\nu': 'ν', r'\\xi': 'ξ', r'\\pi': 'π', r'\\rho': 'ρ',
        r'\\sigma': 'σ', r'\\tau': 'τ', r'\\upsilon': 'υ', r'\\phi': 'φ',
        r'\\chi': 'χ', r'\\psi': 'ψ', r'\\omega': 'ω',
        # Capital letters
        r'\\Gamma': 'Γ', r'\\Delta': 'Δ', r'\\Theta': 'Θ', r'\\Lambda': 'Λ',
        r'\\Xi': 'Ξ', r'\\Pi': 'Π', r'\\Sigma': 'Σ', r'\\Phi': 'Φ',
        r'\\Psi': 'Ψ', r'\\Omega': 'Ω'
    }

    for latex, unicode_char in greek_map.items():
        text = text.replace(latex, unicode_char)

    # Step 4: Fix spacing issues - add space after periods if missing
    text = re.sub(r'\.([A-Z])', r'. \1', text)

    # Step 5: Remove excessive whitespace
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces to single space
    text = re.sub(r'\n{3,}', '\n\n', text)  # Multiple newlines to double newline

    return text.strip()


def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file using multiple methods for better reliability."""
    pdf_path_str = str(pdf_path)
    extracted_text = ""

    # Method 1: Try PyMuPDF (fitz) first - better word spacing preservation
    try:
        import fitz  # PyMuPDF
        logger.info("🔍 Attempting extraction with PyMuPDF (better word spacing)...")

        doc = fitz.open(pdf_path_str)
        pymupdf_text = ""

        for page_num, page in enumerate(doc):
            # extract_text() with "text" mode preserves word spacing better
            page_text = page.get_text("text")
            if page_text:
                # Clean the text immediately after extraction
                page_text = clean_extracted_text(page_text)
                pymupdf_text += f"{page_text}\n\n"
            else:
                logger.warning(f"⚠️ PyMuPDF: No text extracted from page {page_num + 1}")

        doc.close()

        if pymupdf_text.strip():
            logger.info(f"✅ PyMuPDF successfully extracted {len(pymupdf_text)} characters")
            extracted_text = pymupdf_text
        else:
            logger.warning("⚠️ PyMuPDF extraction yielded no text, trying alternative method...")
    except ImportError:
        logger.warning("⚠️ PyMuPDF not installed (pip install pymupdf), falling back to PyPDF2...")
    except Exception as e:
        logger.warning(f"⚠️ PyMuPDF extraction error: {e}")

    # Method 2: Try PyPDF2 as fallback (has known word spacing issues)
    if not extracted_text:
        try:
            logger.info("🔍 Attempting extraction with PyPDF2...")
            reader = PdfReader(pdf_path_str)
            pypdf_text = ""

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    pypdf_text += f"{page_text}\n\n"
                else:
                    logger.warning(f"⚠️ PyPDF2: No text extracted from page {page_num + 1}")

            if pypdf_text.strip():
                logger.info(f"✅ PyPDF2 successfully extracted {len(pypdf_text)} characters")
                logger.warning("⚠️ PyPDF2 may have word spacing issues - install pymupdf for better quality")
                extracted_text = pypdf_text
            else:
                logger.warning("⚠️ PyPDF2 extraction yielded no text, trying alternative method...")
        except Exception as e:
            logger.warning(f"⚠️ PyPDF2 extraction error: {e}")

    # Method 3: If PyMuPDF and PyPDF2 failed, try pdftotext if available
    if not extracted_text:
        try:
            logger.info("🔍 Attempting extraction with pdftotext (if installed)...")
            with tempfile.NamedTemporaryFile(suffix=".txt") as temp:
                # Try to use pdftotext (from poppler-utils) if installed
                result = subprocess.run(
                    ["pdftotext", "-layout", pdf_path_str, temp.name], capture_output=True, text=True
                )

                if result.returncode == 0:
                    with open(temp.name, "r", encoding="utf-8") as f:
                        pdftotext_text = f.read()

                    if pdftotext_text.strip():
                        logger.info(f"✅ pdftotext successfully extracted {len(pdftotext_text)} characters")
                        extracted_text = pdftotext_text
                    else:
                        logger.warning("⚠️ pdftotext extraction yielded no text")
                else:
                    logger.warning(f"⚠️ pdftotext error: {result.stderr}")
        except FileNotFoundError:
            logger.warning("⚠️ pdftotext not found on system, skipping alternative extraction")
        except Exception as e:
            logger.warning(f"⚠️ Alternative extraction error: {e}")

    # Method 4: If still no text, try OCR (for scanned documents)
    if not extracted_text or len(extracted_text.strip()) < 100:
        try:
            logger.info("🔍 Attempting OCR extraction (for scanned/image-based PDFs)...")
            from pdf2image import convert_from_path
            import pytesseract

            # Convert PDF to images
            images = convert_from_path(pdf_path_str, dpi=300)
            logger.info(f"📄 Converted PDF to {len(images)} image(s)")

            ocr_text = ""
            for i, image in enumerate(images, 1):
                logger.info(f"   OCR processing page {i}/{len(images)}...")
                page_text = pytesseract.image_to_string(image, lang="eng")
                if page_text.strip():
                    ocr_text += f"--- Page {i} ---\n\n{page_text.strip()}\n\n"

            if ocr_text.strip():
                logger.info(f"✅ OCR successfully extracted {len(ocr_text)} characters")
                extracted_text = ocr_text
            else:
                logger.warning("⚠️ OCR extraction yielded no text")

        except ImportError:
            logger.warning("⚠️ OCR libraries not available (pdf2image, pytesseract)")
            logger.info("   Install with: pip install pdf2image pytesseract")
            logger.info("   Also need: sudo apt-get install tesseract-ocr poppler-utils")
        except Exception as e:
            logger.warning(f"⚠️ OCR extraction error: {e}")

    # Check if we have any text after trying all methods
    if not extracted_text:
        logger.error("❌ Failed to extract text with all available methods")
        return ""

    # Process the text to make it more readable
    processed_text = ""
    pages = extracted_text.split("\f")  # Form feed character often separates PDF pages

    for page_num, page_content in enumerate(pages):
        if page_content.strip():
            processed_text += f"--- Page {page_num + 1} ---\n\n{page_content.strip()}\n\n"

    return processed_text.strip()


def process_pdf(pdf_path):
    """Extracts text from PDF, embeds it, and saves clean conversions."""
    start_time = time.time()

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        logger.error(f"❌ Error: File not found → {pdf_path}")
        return

    logger.info(f"📄 Processing '{pdf_path.name}'...")

    # Extract text from PDF using multiple methods
    extracted_text = extract_text_from_pdf(pdf_path)

    if not extracted_text:
        logger.error(f"❌ Failed to extract text from {pdf_path.name}")
        logger.info("💡 This PDF might be:")
        logger.info("   - Scanned (image-based) without OCR")
        logger.info("   - Protected/encrypted")
        logger.info("   - Using non-standard fonts")
        logger.info("   - Corrupted or malformed")
        return

    # Debug: Show a sample of the extracted text
    sample_length = min(200, len(extracted_text))
    logger.info(f"📊 Extracted {len(extracted_text)} characters from {pdf_path.name}")
    logger.info(f"📃 Sample of extracted text: \n{extracted_text[:sample_length]}...")

    # Prepare a clean version of text without the page markers for better readability
    clean_text = extracted_text
    if "--- Page" in clean_text:
        # Remove page markers if present but preserve content
        clean_text = "\n\n".join([line for line in clean_text.split("\n") if not line.strip().startswith("--- Page")])

    # Get the PDF filename without extension for use in the document title and headings
    pdf_filename = pdf_path.stem

    # Save TXT - preserving the original filename
    txt_path = PROCESSED_DIR / f"{pdf_filename}.txt"
    with open(txt_path, "w", encoding="utf-8") as txt_file:
        # Add the PDF filename as the first line of the text file
        txt_file.write(f"# {pdf_filename}\n\n")
        txt_file.write(clean_text)
    logger.info(f"✅ Saved TXT → {txt_path}")

    # Save Markdown - preserving the original filename
    md_path = PROCESSED_DIR / f"{pdf_filename}.md"
    with open(md_path, "w", encoding="utf-8") as md_file:
        # Creating proper markdown with the PDF filename as the title
        md_content = f"# {pdf_filename}\n\n{clean_text}"
        md_file.write(md_content)
    logger.info(f"✅ Saved Markdown → {md_path}")

    # Save DOCX - preserving the original filename
    docx_path = PROCESSED_DIR / f"{pdf_filename}.docx"
    doc = docx.Document()

    # Add the PDF filename as the document title/heading
    doc.add_heading(sanitize_for_xml(pdf_filename), level=1)

    # Split text into paragraphs for better DOCX formatting
    paragraphs = clean_text.split("\n\n")
    for para in paragraphs:
        if para.strip():
            # Sanitize text to remove control characters that break XML
            sanitized_para = sanitize_for_xml(para.strip())
            if sanitized_para:
                doc.add_paragraph(sanitized_para)

    doc.save(docx_path)
    logger.info(f"✅ Saved DOCX → {docx_path}")

    # Save JSON - preserving the original filename with metadata
    json_path = PROCESSED_DIR / f"{pdf_filename}.json"
    json_data = {
        "filename": pdf_filename,
        "original_path": str(pdf_path),
        "processed_date": datetime.now().isoformat(),
        "character_count": len(clean_text),
        "word_count": len(clean_text.split()),
        "title": pdf_filename,
        "content": clean_text,
        "metadata": {
            "extraction_method": "PyPDF2/pdftotext",
            "file_size_bytes": pdf_path.stat().st_size,
            "formats_generated": ["txt", "md", "docx", "json"],
        },
    }
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(json_data, json_file, indent=2, ensure_ascii=False)
    logger.info(f"✅ Saved JSON → {json_path}")

    # Add to knowledge base for chat capability
    logger.info(f"🧠 [{pdf_path.stem}] Adding document to knowledge base...")
    chunks = chunk_text(clean_text)
    logger.info(f"📊 [{pdf_path.stem}] Document divided into {len(chunks)} semantic chunks")

    doc_id = register_document(pdf_path, txt_path, clean_text, chunks)
    logger.info(f"✅ Document registered with ID: {doc_id}")

    elapsed_time = time.time() - start_time
    logger.info(f"🎯 Completed processing {pdf_path.name}")
    logger.info(f"⏱️  Total time: {elapsed_time:.2f}s ({elapsed_time/60:.1f} minutes)")


def process_directory(dir_path):
    """Processes all PDFs in a given directory."""
    start_time = time.time()

    dir_path = Path(dir_path).expanduser().resolve()

    logger.info(f"🔍 Debug: Checking path → {dir_path}")

    if not dir_path.exists() or not dir_path.is_dir():
        logger.error(f"❌ Error: Directory not found → {dir_path}")
        return

    pdf_files = list(dir_path.glob("*.pdf"))
    if not pdf_files:
        logger.warning(f"⚠️ No PDFs found in {dir_path}")
        return

    logger.info(f"📂 Found {len(pdf_files)} PDFs. Processing...")

    # Determine if we should use parallel batch processing
    num_nodes = len(load_balancer.nodes) if load_balancer and load_balancer.nodes else 1
    num_pdfs = len(pdf_files)

    # Use SOLLOL's locality detection to decide if parallel batch processing is beneficial
    use_parallel_batch = False
    if load_balancer and num_nodes > 1 and num_pdfs > 1:
        # Check if nodes are on different physical machines
        unique_hosts = load_balancer.count_unique_physical_hosts()
        use_parallel_batch = unique_hosts >= 2

        if use_parallel_batch:
            logger.info(
                f"⚡ Parallel batch processing ENABLED: {num_pdfs} PDFs across "
                f"{num_nodes} nodes on {unique_hosts} physical machines"
            )
            logger.info(
                f"   Expected speedup: ~{unique_hosts}x faster than sequential\n"
                f"   (Each PDF will also use parallel embedding within itself)"
            )
        else:
            logger.info(
                f"ℹ️  Sequential batch processing: all {num_nodes} nodes on same machine\n"
                f"   (SOLLOL will parallelize embeddings within each PDF)"
            )
    else:
        logger.info(
            f"ℹ️  Sequential batch processing: {num_nodes} node(s) available\n"
            f"   (SOLLOL will parallelize embeddings within each PDF)"
        )

    # Process PDFs: parallel if beneficial, sequential otherwise
    if use_parallel_batch:
        # Parallel batch processing: process multiple PDFs concurrently
        # Each PDF processing will internally use SOLLOL's embedding parallelization
        max_workers = min(unique_hosts, num_pdfs)

        logger.info(f"🚀 Starting {max_workers} parallel workers...\n")

        completed = 0
        failed = 0

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all PDF processing tasks
            future_to_pdf = {executor.submit(process_pdf, pdf): pdf for pdf in pdf_files}

            # Process results as they complete
            for future in as_completed(future_to_pdf):
                pdf = future_to_pdf[future]
                try:
                    future.result()
                    completed += 1
                    logger.info(f"✅ Progress: {completed}/{num_pdfs} PDFs completed")
                except Exception as e:
                    failed += 1
                    logger.error(f"❌ Failed to process {pdf.name}: {e}")

        if failed > 0:
            logger.warning(f"⚠️  {failed}/{num_pdfs} PDFs failed to process")
    else:
        # Sequential batch processing (original behavior)
        # SOLLOL's intelligent load balancing handles parallelization during embedding
        for i, pdf in enumerate(pdf_files, 1):
            logger.info(f"📄 Processing {i}/{num_pdfs}: {pdf.name}")
            process_pdf(pdf)

    elapsed_time = time.time() - start_time
    logger.info("✅ All PDFs processed!")
    logger.info(f"⏱️  Total batch time: {elapsed_time:.2f}s ({elapsed_time/60:.1f} minutes)")
    logger.info(f"📊 Average: {elapsed_time/len(pdf_files):.2f}s per PDF")


def chat():
    """Starts an interactive chat with embedded documents."""
    index_data = load_document_index()
    if not index_data["documents"]:
        logger.info("📚 No documents in knowledge base yet. Process a PDF first.")
        return

    logger.info("\n💬 Chat with your Documents")
    logger.info("Type 'exit' to return to main menu")
    logger.info(f"Knowledge base contains {len(index_data['documents'])} documents")

    chat_history = []

    while True:
        user_query = input("\n🙋 You: ").strip()

        if user_query.lower() == "exit":
            logger.info("Returning to main menu...")
            break

        if not user_query:
            continue

        # Start timing
        response_start_time = time.time()

        # Find relevant document chunks
        logger.info("🔍 Searching knowledge base...")
        retrieval_start = time.time()
        relevant_chunks = get_similar_chunks(user_query)
        retrieval_time = time.time() - retrieval_start

        if not relevant_chunks:
            logger.error("❌ No relevant information found in the knowledge base.")
            continue

        logger.info(f"   ⏱️  Retrieval: {retrieval_time:.2f}s")

        # Document-aware intelligent context fitting
        # Conservative token limits for 2048-4096 context window models
        # Token estimation: 1 token ≈ 3.5 chars (conservative)
        # Reserve: system prompt (~100) + query (~150) + response (~1500) + history (~250)
        # Available for context: ~1500 tokens base, adjusted dynamically
        BASE_CONTEXT_TOKENS = 1500

        def estimate_tokens(text):
            """Conservative token estimation: 1 token ≈ 3.5 chars."""
            return int(len(text) / 3.5)

        # Group chunks by document
        from collections import defaultdict

        doc_chunks = defaultdict(list)
        for chunk in relevant_chunks:
            doc_chunks[chunk["doc_name"]].append(chunk)

        num_docs = len(doc_chunks)
        logger.info(f"   📚 Chunks span {num_docs} document(s)")

        # Dynamic strategy based on document count
        if num_docs == 1:
            # Single document: prioritize depth - use more chunks from same doc
            MAX_CONTEXT_TOKENS = BASE_CONTEXT_TOKENS * 1.3  # 1950 tokens
            min_chunks_per_doc = 3
            logger.info("   🎯 Strategy: Deep dive (single document)")
        elif num_docs <= 3:
            # Few documents: balanced approach
            MAX_CONTEXT_TOKENS = BASE_CONTEXT_TOKENS * 1.1  # 1650 tokens
            min_chunks_per_doc = 2
            logger.info(f"   🎯 Strategy: Balanced coverage ({num_docs} documents)")
        else:
            # Many documents: prioritize breadth - sample from each doc
            MAX_CONTEXT_TOKENS = BASE_CONTEXT_TOKENS  # 1500 tokens
            min_chunks_per_doc = 1
            logger.info(f"   🎯 Strategy: Broad coverage ({num_docs} documents)")

        # Build context with document-aware selection
        context_parts = []
        current_tokens = 0
        chunks_used = 0
        docs_included = set()

        # Phase 1: Ensure minimum representation from each document
        for doc_name, chunks in sorted(
            doc_chunks.items(), key=lambda x: max(c["similarity"] for c in x[1]), reverse=True
        ):
            doc_chunks_added = 0
            for chunk in sorted(chunks, key=lambda x: x["similarity"], reverse=True):
                if doc_chunks_added >= min_chunks_per_doc:
                    break

                chunk_text = chunk["text"]
                similarity = chunk["similarity"]
                formatted_chunk = f"[Doc: {doc_name}, Relevance: {similarity:.2f}]\n{chunk_text}"
                chunk_tokens = estimate_tokens(formatted_chunk)

                if current_tokens + chunk_tokens <= MAX_CONTEXT_TOKENS:
                    context_parts.append((similarity, formatted_chunk))
                    current_tokens += chunk_tokens
                    chunks_used += 1
                    doc_chunks_added += 1
                    docs_included.add(doc_name)

        # Phase 2: Fill remaining space with highest relevance chunks
        remaining_chunks = [chunk for doc_name, chunks in doc_chunks.items() for chunk in chunks]
        remaining_chunks.sort(key=lambda x: x["similarity"], reverse=True)

        for chunk in remaining_chunks:
            if chunks_used >= len(relevant_chunks):
                break

            chunk_text = chunk["text"]
            doc_name = chunk["doc_name"]
            similarity = chunk["similarity"]
            formatted_chunk = f"[Doc: {doc_name}, Relevance: {similarity:.2f}]\n{chunk_text}"

            # Skip if already included
            if any(formatted_chunk in part[1] for part in context_parts):
                continue

            chunk_tokens = estimate_tokens(formatted_chunk)

            if current_tokens + chunk_tokens <= MAX_CONTEXT_TOKENS:
                context_parts.append((similarity, formatted_chunk))
                current_tokens += chunk_tokens
                chunks_used += 1
                docs_included.add(doc_name)
            else:
                break

        # Sort final context by relevance
        context_parts.sort(key=lambda x: x[0], reverse=True)
        context = "\n\n".join([part[1] for part in context_parts])

        # Show fitting summary
        logger.info(
            f"   📄 Selected {chunks_used}/{len(relevant_chunks)} chunks "
            f"from {len(docs_included)} document(s) (~{current_tokens} tokens)"
        )

        # Build prompt with context and chat history
        system_prompt = (
            "You are FlockParser AI, a helpful assistant that answers questions based on the user's documents. "
            "Only use information from the provided document context. "
            "If you don't know or the answer isn't in the context, say so."
        )

        # Build user message with context and optional history
        user_message_parts = []

        if chat_history:
            history_text = "\n".join(
                [f"Previous Q: {q}\nPrevious A: {a}" for q, a in chat_history[-2:]]  # Last 2 exchanges only
            )
            user_message_parts.append(f"CHAT HISTORY:\n{history_text}\n")

        user_message_parts.append(f"CONTEXT FROM DOCUMENTS:\n{context}")
        user_message_parts.append(f"\nQUESTION: {user_query}")
        user_message = "\n".join(user_message_parts)

        # Generate response using LLM with load balancing
        logger.info("🤖 Generating response...")
        generation_start = time.time()
        try:
            response = load_balancer.chat(
                model=CHAT_MODEL,
                messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_message}],
                keep_alive=CHAT_KEEP_ALIVE,
                priority=5,
            )

            generation_time = time.time() - generation_start
            answer = response["message"]["content"]

            # Display response
            logger.info(f"\n🤖 AI: {answer}")

            # Update chat history
            chat_history.append((user_query, answer))

            # Show source documents
            logger.info("\n📚 Sources:")
            for i, chunk in enumerate(relevant_chunks[:CHUNKS_TO_SHOW]):
                logger.info(f"  {i+1}. {chunk['doc_name']} (relevance: {chunk['similarity']:.2f})")

            # Show timing breakdown
            total_time = time.time() - response_start_time
            logger.info("\n⏱️  Response timing:")
            logger.info(f"   Retrieval: {retrieval_time:.2f}s")
            logger.info(f"   Generation: {generation_time:.2f}s")
            logger.info(f"   Total: {total_time:.2f}s")

        except Exception as e:
            logger.error(f"❌ Error generating response: {e}")


def check_dependencies():
    """Checks for the presence of external tools that might help with PDF processing."""
    logger.info("🔍 Checking for helpful dependencies...")

    # Check for pdftotext (from Poppler utils)
    try:
        result = subprocess.run(["pdftotext", "-v"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode == 0:
            version_info = result.stderr.strip() if result.stderr else result.stdout.strip()
            logger.info(f"✅ pdftotext found: {version_info}")
        else:
            logger.error("❌ pdftotext is not working properly")
    except FileNotFoundError:
        logger.error("❌ pdftotext not found. For better PDF extraction, consider installing:")
        logger.info("   - Linux: sudo apt-get install poppler-utils")
        logger.info("   - macOS: brew install poppler")
        logger.info("   - Windows: Install from http://blog.alivate.com.au/poppler-windows/")

    # Check PyPDF2 version
    import pkg_resources

    try:
        pypdf_version = pkg_resources.get_distribution("PyPDF2").version
        logger.info(f"✅ PyPDF2 version: {pypdf_version}")
    except pkg_resources.DistributionNotFound:
        logger.error("❌ PyPDF2 not found in installed packages")

    # Check for OCRmyPDF for potential enhancement
    try:
        result = subprocess.run(["ocrmypd", "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode == 0:
            version_info = result.stdout.strip()
            logger.info(f"✅ OCRmyPDF found: {version_info}")
            logger.info("   This can be used to add OCR to scanned PDFs if needed")
        else:
            logger.error("❌ OCRmyPDF is not working properly")
    except FileNotFoundError:
        logger.info("ℹ️ OCRmyPDF not found (optional for OCR capability)")

    # Check Ollama availability
    try:
        result = ollama.list()
        # Display available models - handle both dict and object response formats
        if hasattr(result, "models"):
            # New API returns object with models attribute
            models = result.models
            model_names = [model.model if hasattr(model, "model") else str(model) for model in models]
        else:
            # Fallback for dict format
            models = result.get("models", [])
            model_names = [model.get("name", "unknown") for model in models]

        # Check if embedding model exists (with or without :latest suffix)
        embedding_found = any(EMBEDDING_MODEL in name for name in model_names)
        if embedding_found:
            logger.info(f"✅ Embedding model '{EMBEDDING_MODEL}' is available")
        else:
            logger.warning(f"⚠️ Embedding model '{EMBEDDING_MODEL}' not found in Ollama")
            logger.info(f"   Run 'ollama pull {EMBEDDING_MODEL}' to download it")

        # Check if chat model exists (with or without :latest suffix)
        chat_found = any(CHAT_MODEL in name for name in model_names)
        if chat_found:
            logger.info(f"✅ Chat model '{CHAT_MODEL}' is available")
        else:
            logger.warning(f"⚠️ Chat model '{CHAT_MODEL}' not found in Ollama")
            logger.info(f"   Run 'ollama pull {CHAT_MODEL}' to download it")

    except Exception as e:
        logger.error(f"❌ Ollama not available or error connecting: {e}")
        logger.info("   Make sure Ollama is installed and running")

    logger.info("\n💡 Missing tools can be installed to improve PDF processing capabilities")


def clear_cache():
    """Clear the embedding cache."""
    try:
        if EMBEDDING_CACHE_FILE.exists():
            confirm = input("⚠️  This will delete the embedding cache. Continue? (yes/no): ").strip().lower()
            if confirm == "yes":
                EMBEDDING_CACHE_FILE.unlink()
                logger.info("✅ Embedding cache cleared successfully")
                logger.info("   Next PDF processing will regenerate embeddings")
            else:
                logger.error("❌ Operation cancelled")
        else:
            logger.info("ℹ️ No embedding cache found")
    except Exception as e:
        logger.error(f"❌ Error clearing cache: {e}")


def gpu_status():
    """Show intelligent GPU routing status."""
    logger.info("\n" + "=" * 70)
    logger.info("🎯 INTELLIGENT GPU ROUTING STATUS")
    logger.info("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)
    router.print_cluster_report()


def gpu_route_model(model_name: str):
    """Show routing decision for a specific model."""
    logger.info("\n" + "=" * 70)
    logger.info(f"🧠 ROUTING DECISION FOR: {model_name}")
    logger.info("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)
    decision = router.route_model(model_name)

    logger.info("\n📍 Recommended routing:")
    logger.info(f"   Node: {decision['node']}")
    logger.info(f"   Target: {decision['target']}")
    logger.info(f"   Reason: {decision['reason']}")


def gpu_optimize():
    """Trigger intelligent GPU optimization."""
    priority_models = [EMBEDDING_MODEL, CHAT_MODEL]

    logger.info("\n" + "=" * 70)
    logger.info(f"🔧 OPTIMIZING {len(priority_models)} PRIORITY MODELS")
    logger.info("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)
    router.optimize_cluster(priority_models)


def gpu_check_fit(model_name: str):
    """Check which nodes can fit a specific model."""
    logger.info("\n" + "=" * 70)
    logger.info(f"✅ CHECKING FIT FOR: {model_name}")
    logger.info("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)
    model_size = router.get_model_size(model_name)

    logger.info(f"\n📦 Model size: {model_size} MB")
    logger.info("\n📍 Node compatibility:")

    for node_url in load_balancer.instances:
        can_fit, reason = router.can_fit_on_gpu(node_url, model_name)
        if can_fit:
            logger.info(f"   ✅ {node_url}: {reason}")
        else:
            logger.error(f"   ❌ {node_url}: {reason}")


def gpu_list_models():
    """List all known models and their sizes."""
    logger.info("\n" + "=" * 70)
    logger.info("📚 KNOWN MODELS DATABASE")
    logger.info("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)

    logger.info("\n📦 Model sizes:")
    for model, size_mb in sorted(router.known_model_sizes.items(), key=lambda x: x[1]):
        size_gb = size_mb / 1024
        logger.info(f"   {model:30s} {size_mb:6.0f} MB ({size_gb:.2f} GB)")


def unload_model(model_name: str):
    """Unload a specific model from all nodes."""
    logger.info(f"\n🗑️  Unloading {model_name} from all nodes...")

    for node_url in load_balancer.instances:
        try:
            # Use keep_alive=0 to unload immediately
            if "embed" in model_name.lower():
                response = requests.post(
                    f"{node_url}/api/embed", json={"model": model_name, "input": "unload", "keep_alive": 0}, timeout=10
                )
            else:
                response = requests.post(
                    f"{node_url}/api/generate",
                    json={"model": model_name, "prompt": "unload", "keep_alive": 0},
                    timeout=10,
                )

            if response.status_code == 200:
                logger.info(f"   ✅ {node_url}: Unloaded {model_name}")
            else:
                logger.warning(f"   ⚠️  {node_url}: Status {response.status_code}")

        except Exception as e:
            logger.warning(f"   ⚠️  {node_url}: {str(e)}")

    logger.info(f"\n✅ Unload requests sent for {model_name}")


def cleanup_models():
    """Unload all non-priority models from all nodes."""
    priority_models = {EMBEDDING_MODEL, CHAT_MODEL}

    logger.info("\n🧹 Cleaning up non-priority models...")
    logger.info(f"   Priority models: {', '.join(priority_models)}")

    models_to_unload = set()

    # Check what's loaded on each node
    for node_url in load_balancer.instances:
        try:
            response = requests.get(f"{node_url}/api/ps", timeout=5)
            if response.status_code == 200:
                data = response.json()
                for model in data.get("models", []):
                    model_name = model.get("name", "")
                    # Check if this model is NOT a priority model
                    is_priority = any(priority in model_name for priority in priority_models)
                    if not is_priority and model_name:
                        models_to_unload.add(model_name)
        except Exception as e:
            logger.error(f"   ⚠️  Error checking {node_url}: {e}")

    if not models_to_unload:
        logger.info("\n✅ No non-priority models to unload")
        return

    logger.info(f"\n📋 Found {len(models_to_unload)} non-priority models:")
    for model in models_to_unload:
        logger.info(f"   - {model}")

    confirm = input("\n⚠️  Unload these models? (yes/no): ").strip().lower()
    if confirm != "yes":
        logger.error("❌ Operation cancelled")
        return

    # Unload each model
    for model_name in models_to_unload:
        unload_model(model_name)

    logger.info("\n✅ Cleanup complete!")


def clear_db():
    """Clear the ChromaDB vector store (removes all documents)."""
    try:
        confirm = (
            input("⚠️  This will DELETE ALL DOCUMENTS from the vector database. Continue? (yes/no): ").strip().lower()
        )
        if confirm != "yes":
            logger.error("❌ Operation cancelled")
            return

        global chroma_collection

        # Delete the collection
        try:
            chroma_client.delete_collection(name="documents")
            logger.info("✅ ChromaDB collection deleted")
        except Exception:
            pass  # Collection might not exist

        # Recreate the collection
        chroma_collection = chroma_client.get_or_create_collection(name="documents", metadata={"hnsw:space": "cosine"})
        logger.info("✅ ChromaDB vector store cleared successfully")
        logger.info("   All documents removed from database")

        # Optionally clear the document index too
        clear_index = input("Also clear document index? (yes/no): ").strip().lower()
        if clear_index == "yes":
            if INDEX_FILE.exists():
                INDEX_FILE.unlink()
            logger.info("✅ Document index cleared")

        # Optionally clear JSON knowledge base
        clear_json = input("Also clear legacy JSON knowledge base? (yes/no): ").strip().lower()
        if clear_json == "yes":
            json_files = list(KB_DIR.glob("*.json"))
            for f in json_files:
                f.unlink()
            logger.info(f"✅ Cleared {len(json_files)} JSON files from knowledge base")

    except Exception as e:
        logger.error(f"❌ Error clearing database: {e}")


def vram_report():
    """Show detailed VRAM usage report for all nodes."""
    monitor = VRAMMonitor()

    logger.info(
        f"\n🔍 Detected GPU type: {monitor.gpu_type.upper() if monitor.gpu_type != 'none' else 'None (CPU only)'}"
    )

    # Get local report
    logger.info("\n📊 Local Node Report:")
    local_report = monitor.get_comprehensive_report("http://localhost:11434")
    monitor.print_report(local_report)

    # Get distributed nodes report
    if len(load_balancer.instances) > 1:
        logger.info("\n🌐 Distributed Nodes Report:")
        node_results = monitor_distributed_nodes(load_balancer.instances)

        for node_url, info in node_results.items():
            if info["status"] == "online":
                gpu_status = "🚀 GPU" if info["gpu_accelerated"] else "🐢 CPU"
                vram_gb = info["vram_mb"] / 1024
                ram_gb = info["ram_mb"] / 1024

                logger.info(f"\n   {gpu_status} {node_url}:")
                if info["gpu_accelerated"]:
                    logger.info(f"      VRAM Usage: {vram_gb:.2f} GB")
                else:
                    logger.info(f"      RAM Usage: {ram_gb:.2f} GB (CPU fallback)")

                if info["models"]:
                    logger.info("      Loaded Models:")
                    for model in info["models"]:
                        logger.info(f"         - {model['name']} ({model['location']})")
            else:
                logger.error(f"   ❌ {node_url}: {info['error']}")

        logger.info("\n" + "=" * 70)


def setup_api_server():
    """Start the FlockParse API server in background thread."""
    api_port = int(os.getenv("FLOCKPARSER_API_PORT", "8000"))
    api_enabled = os.getenv("FLOCKPARSER_API", "true").lower() in ("true", "1", "yes", "on")

    if not api_enabled:
        logger.info("🔌 API server disabled (set FLOCKPARSER_API=true to enable)")
        return

    # Check if API already running
    try:
        response = requests.get(f"http://localhost:{api_port}/", timeout=1)
        if response.status_code == 200:
            logger.info(f"✅ API already running at http://localhost:{api_port}")
            return
    except requests.exceptions.RequestException:
        pass

    # Start API server in background thread
    def run_api():
        import sys
        # Suppress uvicorn startup messages
        original_stdout = sys.stdout
        sys.stdout = open(os.devnull, 'w')

        from flock_ai_api import app
        import uvicorn

        os.makedirs("./uploads", exist_ok=True)
        uvicorn.run(app, host="0.0.0.0", port=api_port, log_level="error")

        sys.stdout = original_stdout

    api_thread = threading.Thread(target=run_api, daemon=True, name="FlockParserAPI")
    api_thread.start()

    # Wait for API to start
    for attempt in range(10):
        time.sleep(0.3)
        try:
            response = requests.get(f"http://localhost:{api_port}/", timeout=1)
            if response.status_code == 200:
                logger.info(f"✅ API server started at http://localhost:{api_port}")
                logger.info(f"   📡 Endpoints: /upload, /search, /summarize")
                logger.info(f"   🔑 API Key: Set via FLOCKPARSE_API_KEY environment variable")
                return
        except requests.exceptions.RequestException:
            continue

    logger.warning(f"⚠️  API server may not have started properly on port {api_port}")


def main():
    """Command-line interface."""
    global load_balancer

    # Initialize load balancer and dashboard
    setup_load_balancer()
    setup_dashboard()
    setup_api_server()

    logger.info("🚀 Welcome to FlockParser")
    logger.info(COMMANDS)

    # Show dashboard URL if enabled
    if _dashboard_enabled:
        logger.info(f"\n📊 Observability Dashboard: http://localhost:{_dashboard_port}")

    # Quick dependency check on startup
    logger.info("\nℹ️ Run 'check_deps' for detailed dependency information")

    while True:
        try:
            command = input("\n⚡ Enter command: ").strip().split()
        except EOFError:
            # Running in background or stdin closed - keep dashboard alive
            if _dashboard_enabled:
                logger.info("📊 Dashboard running in background mode - press Ctrl+C to exit")
                import time
                try:
                    while True:
                        time.sleep(60)
                except KeyboardInterrupt:
                    logger.info("Shutting down...")
                    break
            else:
                break

        if not command:
            continue

        action = command[0]
        arg = " ".join(command[1:]) if len(command) > 1 else None

        if action == "open_pd" and arg:
            process_pdf(arg)
        elif action == "open_dir" and arg:
            process_directory(arg)
        elif action == "chat":
            chat()
        elif action == "list_docs":
            list_documents()
        elif action == "check_deps":
            check_dependencies()
        elif action == "discover_nodes":
            load_balancer.discover_nodes()
        elif action == "add_node" and arg:
            load_balancer.add_node(arg)
        elif action == "remove_node" and arg:
            load_balancer.remove_node(arg)
        elif action == "list_nodes":
            load_balancer.list_nodes()
        elif action == "verify_models":
            load_balancer.verify_models_on_nodes()
        elif action == "lb_stats":
            load_balancer.print_stats()
        elif action == "set_routing" and arg:
            load_balancer.set_routing_strategy(arg)
        elif action == "vram_report":
            vram_report()
        elif action == "force_gpu" and arg:
            load_balancer.force_gpu_all_nodes(arg)
        elif action == "gpu_status":
            gpu_status()
        elif action == "gpu_route" and arg:
            gpu_route_model(arg)
        elif action == "gpu_optimize":
            gpu_optimize()
        elif action == "gpu_check" and arg:
            gpu_check_fit(arg)
        elif action == "gpu_models":
            gpu_list_models()
        elif action == "unload_model" and arg:
            unload_model(arg)
        elif action == "cleanup_models":
            cleanup_models()
        elif action == "parallelism_report":
            from sollol.adaptive_parallelism import print_parallelism_report

            print_parallelism_report(load_balancer)
        elif action == "clear_cache":
            clear_cache()
        elif action == "clear_db":
            clear_db()
        elif action == "exit":
            logger.info("👋 Exiting. See you next time!")
            load_balancer.print_stats()  # Show stats on exit
            break
        else:
            logger.warning("⚠️ Invalid command. Try again.")
            logger.info(COMMANDS)


if __name__ == "__main__":
    # Required for multiprocessing (Dask, Ray) to work correctly
    import multiprocessing
    multiprocessing.freeze_support()

    main()
