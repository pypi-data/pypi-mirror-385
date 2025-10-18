"""Document processor for JSON to Google Docs conversion"""

import json
import os
import re
import tempfile
from typing import Dict, Any, List, Optional, Union

from docx import Document
import markdown
from bs4 import BeautifulSoup

from .config import JsonToGoogleDocsConfig


class DocumentProcessor:
    """Processes JSON data and converts to DOCX with template substitution"""

    def __init__(self, config: JsonToGoogleDocsConfig):
        self.config = config

    def flatten_json(
        self, json_obj: Union[Dict, List, Any], prefix: str = ""
    ) -> Dict[str, Any]:
        """
        Flattens nested JSON into dot notation keys

        Args:
            json_obj: JSON object to flatten
            prefix: Key prefix for nested objects

        Returns:
            Flattened dictionary with dot notation keys
        """
        result = {}

        if isinstance(json_obj, dict):
            for key, value in json_obj.items():
                new_key = f"{prefix}.{key}" if prefix else key

                if isinstance(value, (dict, list)):
                    # Recursively flatten nested structures
                    result.update(self.flatten_json(value, new_key))

                    # Also store stringified version for direct replacement
                    if isinstance(value, dict) and value:
                        result[new_key] = json.dumps(
                            value, ensure_ascii=False, indent=2
                        )
                else:
                    # Store leaf values directly
                    result[new_key] = value

        elif isinstance(json_obj, list):
            # For lists, create indexed keys
            for i, item in enumerate(json_obj):
                new_key = f"{prefix}[{i}]" if prefix else f"item{i}"

                if isinstance(item, (dict, list)):
                    result.update(self.flatten_json(item, new_key))
                else:
                    result[new_key] = item

            # Also store the whole list as a joined string
            if all(isinstance(item, (str, int, float, bool)) for item in json_obj):
                result[prefix] = ", ".join([str(item) for item in json_obj])
            else:
                result[prefix] = json.dumps(json_obj, ensure_ascii=False, indent=2)

        return result

    def read_json_data(self, json_input: Union[str, Dict, List]) -> Dict[str, Any]:
        """
        Read JSON data from string, file path, or object

        Args:
            json_input: JSON data as string, file path, or object

        Returns:
            Parsed JSON data
        """
        try:
            if isinstance(json_input, str):
                # Check if it's a file path
                if os.path.exists(json_input) or json_input.startswith("/"):
                    with open(json_input, "r", encoding="utf-8") as f:
                        return json.load(f)
                else:
                    # Try to parse as JSON string
                    try:
                        return json.loads(json_input)
                    except json.JSONDecodeError:
                        # Treat as plain text
                        return {"content": json_input}

            elif isinstance(json_input, (dict, list)):
                return json_input

            else:
                # Convert other types to string
                return {"content": str(json_input)}

        except Exception as e:
            raise ValueError(f"Failed to read JSON data: {str(e)}")

    def extract_template_placeholders(self, template_path: str) -> List[str]:
        """
        Extract {{placeholder}} keys from a DOCX template

        Args:
            template_path: Path to DOCX template file

        Returns:
            List of placeholder keys found in template
        """
        try:
            doc = Document(template_path)
            placeholders = set()

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                matches = re.findall(r"{{([^}]+)}}", paragraph.text)
                for match in matches:
                    clean_key = match.strip()
                    if ":" in clean_key:
                        clean_key = clean_key.split(":", 1)[0].strip()
                    placeholders.add(clean_key)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(r"{{([^}]+)}}", paragraph.text)
                            for match in matches:
                                clean_key = match.strip()
                                if ":" in clean_key:
                                    clean_key = clean_key.split(":", 1)[0].strip()
                                placeholders.add(clean_key)

            return sorted(list(placeholders))

        except Exception as e:
            raise Exception(f"Failed to extract placeholders: {str(e)}")

    def normalize_text(self, text: str) -> str:
        """
        Normalize text by removing soft returns and extra whitespace

        Args:
            text: Text to normalize

        Returns:
            Normalized text
        """
        if not text:
            return ""

        if isinstance(text, str):
            text = text.strip()

            if "\n" not in text:
                return text

            lines = text.split("\n")
            normalized = []
            for line in lines:
                if line.strip():
                    normalized.append(line.strip())

            return " ".join(normalized)

        return text

    def convert_markdown_to_html(self, text: str) -> str:
        """Convert markdown text to HTML"""
        if not self.config.enable_markdown_conversion:
            return text

        try:
            return markdown.markdown(text)
        except Exception:
            return text

    def replace_placeholders_in_document(
        self, doc: Document, placeholders: Dict[str, Any]
    ) -> Document:
        """
        Replace {{placeholder}} text in DOCX document

        Args:
            doc: python-docx Document object
            placeholders: Dictionary of placeholder values

        Returns:
            Modified Document object
        """
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if "{{" in paragraph.text and "}}" in paragraph.text:
                for key, value in placeholders.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in paragraph.text:
                        # Handle markdown conversion if enabled
                        if self.config.enable_markdown_conversion and isinstance(
                            value, str
                        ):
                            if any(
                                marker in str(value)
                                for marker in ["**", "*", "_", "`", "[", "]("]
                            ):
                                html = self.convert_markdown_to_html(str(value))
                                soup = BeautifulSoup(html, "html.parser")
                                text_value = soup.get_text()
                            else:
                                text_value = str(value) if value is not None else ""
                        else:
                            text_value = str(value) if value is not None else ""

                        # Replace placeholder with value
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, text_value)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "{{" in paragraph.text and "}}" in paragraph.text:
                            for key, value in placeholders.items():
                                placeholder = "{{" + key + "}}"
                                if placeholder in paragraph.text:
                                    text_value = str(value) if value is not None else ""

                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(
                                                placeholder, text_value
                                            )

        return doc

    def validate_json_data(self, json_data: Union[str, Dict, List]) -> Dict[str, Any]:
        """
        Validate JSON data and return analysis

        Args:
            json_data: JSON data to validate

        Returns:
            Dictionary with validation results
        """
        try:
            # Parse JSON data
            parsed_data = self.read_json_data(json_data)

            # Flatten for analysis
            flattened_data = self.flatten_json(parsed_data)

            return {
                "valid": True,
                "data_type": type(parsed_data).__name__,
                "keys_count": len(flattened_data),
                "sample_keys": list(flattened_data.keys())[:10],
                "flattened_data": flattened_data,
            }

        except Exception as e:
            return {
                "valid": False,
                "error": str(e),
                "data_type": None,
                "keys_count": 0,
                "sample_keys": [],
                "flattened_data": {},
            }

    async def convert_json_to_docx(
        self,
        json_data: Union[str, Dict, List],
        template_id: str,
        output_filename: str = "output.docx",
        folder_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Convert JSON data to DOCX using template

        Args:
            json_data: JSON data to convert
            template_id: Google Docs template ID
            output_filename: Output filename
            folder_id: Target folder ID

        Returns:
            Dictionary with conversion results
        """
        try:
            # Parse and flatten JSON data
            parsed_data = self.read_json_data(json_data)
            flattened_data = self.flatten_json(parsed_data)

            # Create temporary files
            template_path = tempfile.mktemp(suffix=".docx")
            output_path = tempfile.mktemp(suffix=".docx")

            try:
                # Download template (would need client integration)
                # For now, create a simple document
                doc = Document()
                doc.add_heading("JSON Data Conversion", 0)

                # Add data as key-value pairs
                for key, value in flattened_data.items():
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(f"{key}: ").bold = True
                    paragraph.add_run(str(value))

                # Save document
                doc.save(output_path)

                return {
                    "status": "success",
                    "output_path": output_path,
                    "template_id": template_id,
                    "keys_processed": len(flattened_data),
                    "output_filename": output_filename,
                }

            finally:
                # Clean up temporary files
                for temp_file in [template_path, output_path]:
                    if os.path.exists(temp_file):
                        try:
                            os.unlink(temp_file)
                        except Exception:
                            pass

        except Exception as e:
            return {"status": "error", "error": str(e), "template_id": template_id}
