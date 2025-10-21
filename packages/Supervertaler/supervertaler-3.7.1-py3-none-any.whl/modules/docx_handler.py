"""
DOCX Handler
Import and export DOCX files with formatting preservation
"""

import os
from typing import List, Dict, Any
from dataclasses import dataclass

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("ERROR: python-docx not installed. Run: pip install python-docx")

# Import tag manager for inline formatting
try:
    from .tag_manager import TagManager
except ImportError:
    try:
        from tag_manager import TagManager
    except ImportError:
        print("WARNING: tag_manager not found. Inline formatting will not be preserved.")
        TagManager = None


@dataclass
class ParagraphInfo:
    """Information about a paragraph for reconstruction"""
    text: str
    style: str = None
    alignment: str = None
    paragraph_index: int = 0
    document_position: int = 0  # Position in original document structure
    is_table_cell: bool = False
    table_index: int = None
    row_index: int = None
    cell_index: int = None


class DOCXHandler:
    """Handle DOCX import and export operations"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is required. Install with: pip install python-docx")
        
        self.original_document = None
        self.original_path = None
        self.paragraphs_info: List[ParagraphInfo] = []
        self.tag_manager = TagManager() if TagManager else None
    
    def import_docx(self, file_path: str, extract_formatting: bool = True) -> List[str]:
        """
        Import DOCX file and extract paragraphs with formatting tags
        
        Args:
            file_path: Path to DOCX file
            extract_formatting: If True, convert formatting to inline tags
        
        Returns: List of paragraph texts (with tags if extract_formatting=True)
                 Includes both regular paragraphs AND table cells
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        print(f"[DOCX Handler] Importing: {file_path}")
        if extract_formatting and self.tag_manager:
            print("[DOCX Handler] Extracting inline formatting as tags")
        
        # Load document
        self.original_document = Document(file_path)
        self.original_path = file_path
        self.paragraphs_info = []
        
        paragraphs = []
        
        # Track position in document structure
        para_counter = 0
        doc_position = 0  # Track actual position in document for proper ordering
        
        # Build mapping of paragraph objects to their positions for tables
        para_to_table_info = {}
        for table_idx, table in enumerate(self.original_document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        para_to_table_info[id(para)] = (table_idx, row_idx, cell_idx)
        
        # Process document elements in order
        # Use document.element.body to get elements in document order
        for elem in self.original_document.element.body:
            # Check if it's a paragraph
            if elem.tag.endswith('}p'):
                # Find corresponding paragraph object
                for para in self.original_document.paragraphs:
                    if para._element == elem:
                        text = para.text.strip()
                        
                        # Check if this paragraph is inside a table
                        if id(para) in para_to_table_info:
                            # This paragraph is in a table, skip it here
                            # (tables are handled separately below)
                            break
                        
                        if text:  # Only include non-empty paragraphs
                            # Extract formatting if requested
                            if extract_formatting and self.tag_manager:
                                runs = self.tag_manager.extract_runs(para)
                                text_with_tags = self.tag_manager.runs_to_tagged_text(runs)
                                paragraphs.append(text_with_tags)
                            else:
                                paragraphs.append(text)
                            
                            # Store paragraph info for reconstruction
                            para_info = ParagraphInfo(
                                text=text,
                                style=para.style.name if para.style else None,
                                alignment=str(para.alignment) if para.alignment else None,
                                paragraph_index=para_counter,
                                document_position=doc_position,
                                is_table_cell=False
                            )
                            self.paragraphs_info.append(para_info)
                            para_counter += 1
                        
                        doc_position += 1
                        break
            
            # Check if it's a table
            elif elem.tag.endswith('}tbl'):
                # Find corresponding table object
                for table_idx, table in enumerate(self.original_document.tables):
                    if table._element == elem:
                        # Process this table
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                # Each cell may contain multiple paragraphs
                                for para in cell.paragraphs:
                                    text = para.text.strip()
                                    
                                    if text:  # Only include non-empty cells
                                        # Extract formatting if requested
                                        if extract_formatting and self.tag_manager:
                                            runs = self.tag_manager.extract_runs(para)
                                            text_with_tags = self.tag_manager.runs_to_tagged_text(runs)
                                            paragraphs.append(text_with_tags)
                                        else:
                                            paragraphs.append(text)
                                        
                                        # Store table cell info
                                        para_info = ParagraphInfo(
                                            text=text,
                                            style=para.style.name if para.style else None,
                                            alignment=str(para.alignment) if para.alignment else None,
                                            paragraph_index=para_counter,
                                            document_position=doc_position,
                                            is_table_cell=True,
                                            table_index=table_idx,
                                            row_index=row_idx,
                                            cell_index=cell_idx
                                        )
                                        self.paragraphs_info.append(para_info)
                                        para_counter += 1
                        
                        doc_position += 1  # Table counts as one position
                        break
        
        table_cell_count = sum(1 for p in self.paragraphs_info if p.is_table_cell)
        print(f"[DOCX Handler] Extracted {len(paragraphs)} total items:")
        print(f"  - Regular paragraphs: {len(paragraphs) - table_cell_count}")
        print(f"  - Table cells: {table_cell_count} (from {len(self.original_document.tables)} tables)")
        return paragraphs
    
    def export_docx(self, segments: List[Dict[str, Any]], output_path: str, 
                    preserve_formatting: bool = True):
        """
        Export translated segments back to DOCX
        
        Args:
            segments: List of segment dictionaries with 'paragraph_id', 'source', 'target'
            output_path: Path to save the translated document
            preserve_formatting: Whether to preserve original formatting (default True)
        """
        print(f"[DOCX Handler] Exporting to: {output_path}")
        
        if not self.original_document:
            raise ValueError("No original document loaded. Import a DOCX first.")
        
        # Create a new document based on the original
        if preserve_formatting and self.original_path:
            # Copy the original document structure
            doc = Document(self.original_path)
        else:
            # Create new blank document
            doc = Document()
        
        # Group segments by paragraph index
        para_segments = {}
        for seg in segments:
            para_id = seg.get('paragraph_id', 0)
            if para_id not in para_segments:
                para_segments[para_id] = []
            para_segments[para_id].append(seg)
        
        # Track which paragraphs we've processed
        processed_paras = set()
        
        print(f"[DOCX Export] Starting export with {len(segments)} segments")
        print(f"[DOCX Export] Paragraph segments grouped into {len(para_segments)} paragraph indices")
        print(f"[DOCX Export] Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        
        # Build a mapping of paragraph objects in tables
        table_paras = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        table_paras.add(id(para))
        
        print(f"[DOCX Export] Found {len(table_paras)} paragraphs inside tables")
        
        # First, process regular paragraphs (excluding those in tables)
        non_empty_para_index = 0
        for para_idx, para in enumerate(doc.paragraphs):
            # Skip paragraphs that are inside tables
            if id(para) in table_paras:
                print(f"[DOCX Export] Skipping doc.paragraphs[{para_idx}] - it's inside a table")
                continue
            
            # Only process non-empty paragraphs (same logic as import)
            if not para.text.strip():
                print(f"[DOCX Export] Skipping doc.paragraphs[{para_idx}] - empty paragraph")
                continue
            
            # Check if this paragraph has corresponding segments
            if non_empty_para_index in para_segments:
                para_info = self._get_para_info(non_empty_para_index)
                
                # Double-check it's not a table cell (should already be filtered)
                if para_info and para_info.is_table_cell:
                    print(f"[DOCX Export] ERROR: Para {non_empty_para_index} marked as table cell but found in regular paragraphs!")
                    non_empty_para_index += 1
                    continue
                
                # Combine all segments from this paragraph
                translations = [s['target'] for s in para_segments[non_empty_para_index] 
                              if s['target'].strip()]
                
                if translations:
                    # Join segments back into paragraph (single space, no extra newlines)
                    new_text = ' '.join(translations)
                    
                    print(f"[DOCX Export] Para {non_empty_para_index}: Replacing with {len(translations)} segment(s)")
                    print(f"[DOCX Export]   Original: {para.text[:50]}...")
                    print(f"[DOCX Export]   New: {new_text[:50]}...")
                    
                    # Replace text while preserving formatting AND style
                    self._replace_paragraph_text(para, new_text, para_info.style if para_info else None)
                    processed_paras.add(non_empty_para_index)
                else:
                    print(f"[DOCX Export] Para {non_empty_para_index}: No translations found")
            else:
                print(f"[DOCX Export] Para {non_empty_para_index}: No segments for this paragraph")
            
            non_empty_para_index += 1
        
        # Then, process table cells
        print(f"[DOCX Export] Processing {len(doc.tables)} tables...")
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Each cell may contain multiple paragraphs
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        
                        # Find the paragraph info for this table cell
                        para_info = self._find_table_cell_info(table_idx, row_idx, cell_idx)
                        
                        if para_info and para_info.paragraph_index in para_segments:
                            # Get translations for this cell
                            translations = [s['target'] for s in para_segments[para_info.paragraph_index] 
                                          if s['target'].strip()]
                            
                            if translations:
                                new_text = ' '.join(translations)
                                print(f"[DOCX Export] Table[{table_idx}][{row_idx}][{cell_idx}] Para {para_info.paragraph_index}: Replacing")
                                print(f"[DOCX Export]   Original: {para.text[:50]}...")
                                print(f"[DOCX Export]   New: {new_text[:50]}...")
                                # Table cells can also have styles - preserve them
                                self._replace_paragraph_text(para, new_text, para_info.style)
                                processed_paras.add(para_info.paragraph_index)
                        else:
                            if para_info:
                                print(f"[DOCX Export] Table[{table_idx}][{row_idx}][{cell_idx}] Para {para_info.paragraph_index}: No translations")
                            else:
                                print(f"[DOCX Export] Table[{table_idx}][{row_idx}][{cell_idx}]: No para_info found")
        
        # Save the document
        doc.save(output_path)
        print(f"[DOCX Handler] Export complete: {output_path}")
        print(f"[DOCX Handler] Translated {len(processed_paras)} items (paragraphs + table cells)")
    
    def _get_para_info(self, paragraph_index: int):
        """Get ParagraphInfo by paragraph index"""
        for info in self.paragraphs_info:
            if info.paragraph_index == paragraph_index:
                return info
        return None
    
    def _find_table_cell_info(self, table_idx: int, row_idx: int, cell_idx: int):
        """Find ParagraphInfo for a specific table cell"""
        for info in self.paragraphs_info:
            if (info.is_table_cell and 
                info.table_index == table_idx and 
                info.row_index == row_idx and 
                info.cell_index == cell_idx):
                return info
        return None
    
    def _replace_paragraph_text(self, paragraph, new_text: str, original_style: str = None):
        """
        Replace paragraph text while preserving or applying formatting
        
        If new_text contains inline tags (e.g., <b>text</b>), they will be
        converted to proper formatting runs.
        
        Args:
            paragraph: The paragraph object to modify
            new_text: The new text content
            original_style: Optional original style name to preserve
        """
        # Check if text contains formatting tags
        if self.tag_manager and ('<b>' in new_text or '<i>' in new_text or '<u>' in new_text or '<bi>' in new_text):
            self._replace_paragraph_with_formatting(paragraph, new_text, original_style)
            return
        
        # Simple replacement (no tags) - preserve original formatting
        # Store original formatting from first run (if any)
        original_font_name = None
        original_font_size = None
        original_bold = False
        original_italic = False
        
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.font:
                original_font_name = first_run.font.name
                original_font_size = first_run.font.size
                original_bold = first_run.font.bold or False
                original_italic = first_run.font.italic or False
        
        # Clear paragraph - delete all runs except first
        while len(paragraph.runs) > 1:
            paragraph._element.remove(paragraph.runs[-1]._element)
        
        # If no runs exist, create one
        if not paragraph.runs:
            run = paragraph.add_run()
        else:
            run = paragraph.runs[0]
        
        # Set the new text (strip any trailing/leading whitespace to avoid extra newlines)
        run.text = new_text.strip()
        
        # Restore run-level formatting
        if original_font_name:
            run.font.name = original_font_name
        if original_font_size:
            run.font.size = original_font_size
        if original_bold:
            run.font.bold = True
        if original_italic:
            run.font.italic = True
        
        # Preserve paragraph style if provided
        if original_style:
            try:
                paragraph.style = original_style
            except KeyError:
                # Style doesn't exist in document - keep original
                print(f"[DOCX Handler] Warning: Style '{original_style}' not found, keeping original style")
                pass
    
    def _replace_paragraph_with_formatting(self, paragraph, tagged_text: str, original_style: str = None):
        """
        Replace paragraph text with formatted runs based on inline tags
        
        Example: "Hello <b>world</b>!" creates runs with proper bold formatting
        
        Args:
            paragraph: The paragraph object to modify
            tagged_text: Text with inline formatting tags
            original_style: Optional original style name to preserve
        """
        if not self.tag_manager:
            # Fallback: strip tags and use simple replacement
            clean_text = tagged_text.replace('<b>', '').replace('</b>', '')
            clean_text = clean_text.replace('<i>', '').replace('</i>', '')
            clean_text = clean_text.replace('<u>', '').replace('</u>', '')
            clean_text = clean_text.replace('<bi>', '').replace('</bi>', '')
            self._replace_paragraph_text(paragraph, clean_text, original_style)
            return
        
        # Store original font properties
        original_font_name = None
        original_font_size = None
        
        if paragraph.runs and paragraph.runs[0].font:
            first_run = paragraph.runs[0]
            original_font_name = first_run.font.name
            original_font_size = first_run.font.size
        
        # Clear all runs
        for run in paragraph.runs:
            paragraph._element.remove(run._element)
        
        # Convert tagged text to run specifications
        run_specs = self.tag_manager.tagged_text_to_runs(tagged_text)
        
        # Create runs with proper formatting
        for spec in run_specs:
            run = paragraph.add_run(spec['text'])
            
            # Apply formatting
            if spec.get('bold'):
                run.font.bold = True
            if spec.get('italic'):
                run.font.italic = True
            if spec.get('underline'):
                run.font.underline = True
            
            # Restore original font properties
            if original_font_name:
                run.font.name = original_font_name
            if original_font_size:
                run.font.size = original_font_size
        
        # Preserve paragraph style if provided
        if original_style:
            try:
                paragraph.style = original_style
            except KeyError:
                # Style doesn't exist in document - keep original
                print(f"[DOCX Handler] Warning: Style '{original_style}' not found, keeping original style")
                pass
    
    def export_bilingual_docx(self, segments: List[Dict[str, Any]], output_path: str):
        """
        Export as bilingual document (source | target in table)
        Useful for review purposes
        """
        print(f"[DOCX Handler] Exporting bilingual document: {output_path}")
        
        doc = Document()
        doc.add_heading('Bilingual Translation Document', 0)
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = '#'
        header_cells[1].text = 'Source'
        header_cells[2].text = 'Target'
        
        # Add segments
        for seg in segments:
            row_cells = table.add_row().cells
            row_cells[0].text = str(seg.get('id', ''))
            row_cells[1].text = seg.get('source', '')
            row_cells[2].text = seg.get('target', '')
        
        doc.save(output_path)
        print(f"[DOCX Handler] Bilingual export complete")
    
    def get_document_info(self) -> Dict[str, Any]:
        """Get information about the loaded document"""
        if not self.original_document:
            return {}
        
        # Count table cells
        table_cells = sum(1 for info in self.paragraphs_info if info.is_table_cell)
        regular_paras = sum(1 for info in self.paragraphs_info if not info.is_table_cell)
        
        return {
            'paragraphs': len(self.original_document.paragraphs),
            'sections': len(self.original_document.sections),
            'tables': len(self.original_document.tables),
            'table_cells': table_cells,
            'regular_paragraphs': regular_paras,
            'total_items': len(self.paragraphs_info),
            'path': self.original_path
        }


# Quick test
if __name__ == "__main__":
    print("DOCX Handler Test")
    print("To test, you need a sample DOCX file.")
    
    if DOCX_AVAILABLE:
        print("✓ python-docx is installed")
    else:
        print("✗ python-docx is NOT installed")
        print("  Run: pip install python-docx")
