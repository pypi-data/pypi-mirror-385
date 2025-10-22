from __future__ import annotations 
from ftplib import FTP as FTPClient
import os
import qrcode
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'     # 0 = all logs, 1 = INFO, 2 = WARNING, 3 = ERROR
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'
import torch
import tensorflow as tf
import types
import threading
import base64
import binascii
import urllib.parse
import codecs
import numpy as np
import pandas as pd
import joblib
from math import *
from sklearn import (
    datasets, preprocessing, model_selection, metrics,
    linear_model, tree, neighbors, svm, ensemble, cluster,
    decomposition
)
from flask import *
import numpy as np
import sqlite3
import webbrowser
import pyttsx3
import matplotlib.pyplot as plt
import speech_recognition as sr
import pandas as pd
import pygame 
from PIL import Image
import cv2
import math
import calendar
import datetime
import requests
from docx import Document
import json
import pdfplumber
from bs4 import BeautifulSoup
import pdf2docx
import docx2pdf
import fitz
from fpdf import FPDF
import markdown
from sympy import symbols, Function, diff, integrate, limit, series, summation, Eq, dsolve, sympify, Matrix
from sympy.vector import CoordSys3D, divergence, curl

#__________________________Web searches___________________________

def search_chrome(thing):
    webbrowser.open(rf"https://www.google.com/search?q={thing}")
def search_youtube(thing):
    webbrowser.open(fr"https://www.youtube.com/results?search_query={thing}")
def open_whatsapp():
    webbrowser.open("https://wa.me/")
def open_whatsapp_chat(phone_number:str,message:str="Hi"):
    webbrowser.open(fr"https://wa.me/{phone_number}?text={message}")
def open_other(link):
    webbrowser.open(f"https://www.google.com/search?q={link}")

#________________________Graph Plot_______________________________
class Visualizer2D():
    def __init__(self,x:list=[],y:list=[],title="NeuraPy Graphs",x_label="X-Axis",y_label="Y-Axis"):
        self.x=x
        self.y=y
        self.title=title
        self.x_label=x_label
        self.y_label=y_label
    def bar_graph(self):
        plt.bar(self.x,self.y)
        plt.title(self.title)
        plt.xlabel(self.x_label)
        plt.ylabel(self.y_label)
    def pie_chart(self):
        plt.pie(self.y,labels=self.x,radius=1)
        plt.title(self.title)
        plt.xlabel(self.x_label)
        plt.ylabel(self.y_label)
    def line_graph(self):
        plt.plot(self.x,self.y)
        plt.title(self.title)
        plt.xlabel(self.x_label)
        plt.ylabel(self.y_label)
    def scatter_graph(self):
        plt.scatter(self.x, self.y)
        plt.title(self.title)
        plt.xlabel(self.x_label)
        plt.ylabel(self.y_label)
    def HorizontalBar_chart(self):
        plt.barh(self.x, self.y)
        plt.title(self.title)
        plt.xlabel(self.x_label)
        plt.ylabel(self.y_label)
        
        
    def polar_graph(self):
        if len(self.x) != len(self.y):
            raise ValueError("x and y must be of the same length for polar plotting")

        plt.figure(figsize=(6, 6))
        ax = plt.subplot(111, polar=True)
        ax.plot(self.x, self.y, marker='o', linestyle='-', color='b')
        ax.fill(self.x, self.y, alpha=0.2)
        ax.set_title(self.title, va='bottom')
    
    
    
    
    def show(self):
        plt.show()


#____________________________3D Graphs_____________________________
class Visualizer3D():
    def __init__(self,x=[],y=[],z=[],title="NeuraPy 3D",x_label="X-Axis",y_label="Y-Axis",z_label="Z-Axis"):
        self.x=x
        self.y=y
        self.z=z
        self.title=title
        self.x_label=x_label
        self.y_label=y_label
        self.z_label=z_label
    def plot(self):
        fig = plt.figure()
        ax = fig.add_subplot(111, projection='3d')
        ax.plot(self.x, self.y, self.z)
        ax.set_title(self.title)
        ax.set_xlabel(self.x_label)
        ax.set_ylabel(self.y_label)
        ax.set_zlabel(self.z_label)
    
    
    
    
    def show(self):
        plt.show()
#_______________________Speech____________________________________
def speak(text="Hi I am Neura Python"):
    engine=pyttsx3.init()
    engine.say(text)
    engine.runAndWait()
    engine.stop()
    
    
def voice_input(message="Please say Something: "):
    r = sr.Recognizer()

    
    with sr.Microphone() as source:
        print(message)
        r.adjust_for_ambient_noise(source) 
        audio = r.listen(source)

    try:
        
        text = r.recognize_google(audio)
        return text
    except sr.UnknownValueError:
        print("Sorry, could not understand your voice.")
    except sr.RequestError:
        print("Could not request results; check your internet connection.")




import os
import sqlite3
import json
import uuid
import logging
from flask import Flask, request, jsonify, send_file, render_template_string
import matplotlib.pyplot as plt
import io
import base64

class WebServer:
    def __init__(self, app_name="NeuraPy Web Server"):
        self.app = Flask(app_name)
        self.registered_routes = set()
        logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# -------------------------------------------------------------------
# Helper to avoid duplicate route names
# -------------------------------------------------------------------
    def _unique_route_name(self, route):
        """Generate a unique internal endpoint name for each route."""
        base = route.strip("/").replace("/", "_") or "root"
        unique_name = f"{base}_{uuid.uuid4().hex[:6]}"
        return unique_name

# -------------------------------------------------------------------
# Basic Route Handler
# -------------------------------------------------------------------
    def simple_route(self, route='/', code="NeuraPy web app", html_file_path=""):
        endpoint_name = self._unique_route_name(route)

        @self.app.route(route, endpoint=endpoint_name)
        def webpage():
            if html_file_path:
                if os.path.exists(html_file_path):
                    with open(html_file_path, "r", encoding="utf-8") as f:
                        return f.read()
                return "No file found", 404
            return code or "No content provided"

        logging.info(f"Registered route: {route}")

# -------------------------------------------------------------------
# Error Handler
# -------------------------------------------------------------------
    def error_handler(self, error_code, code="NeuraPy Error Page", error_page_html=""):
        @self.app.errorhandler(error_code)
        def error(e):
            if error_page_html and os.path.exists(error_page_html):
                with open(error_page_html, "r", encoding="utf-8") as f:
                    return f.read(), error_code
            return code or f"Error {error_code}", error_code

# -------------------------------------------------------------------
# JSON Verification (for APIs)
# -------------------------------------------------------------------
    def verify_details(self, route='/', user_data=[], verify_data_from=[]):
        endpoint_name = self._unique_route_name(route)

        @self.app.route(route, methods=['POST'], endpoint=endpoint_name)
        def verify():
            data = request.get_json(force=True)
            if not data:
                return jsonify({"response": False, "error": "No JSON received"})
            if len(user_data) != len(verify_data_from):
                return jsonify({"response": False, "error": "Mismatch in data length"})

            validation = [data.get(k) == v for k, v in zip(user_data, verify_data_from)]
            return jsonify({"response": all(validation)})

# -------------------------------------------------------------------
# SQLite Database Setup & Operations
# -------------------------------------------------------------------
    def DataBase(self, db_path="neurapy.db", query=None):
        query = query or """
            CREATE TABLE IF NOT EXISTS NeuraPy (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                age INTEGER,
                email TEXT UNIQUE
            );
        """
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute(query)
        conn.commit()
        conn.close()
        logging.info(f"Database initialized at {db_path}")

    def insert_data(self, db_path, table, data: dict):
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cols = ', '.join(data.keys())
            placeholders = ', '.join(['?'] * len(data))
            vals = tuple(data.values())
            cursor.execute(f"INSERT INTO {table} ({cols}) VALUES ({placeholders})", vals)
            conn.commit()
        return jsonify({"status": "Data inserted successfully"})

    def retrieve_data(self, db_path, query):
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cursor.execute(query)
            data = cursor.fetchall()
        return jsonify(data)

    def delete_data(self, db_path, table, condition):
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cursor.execute(f"DELETE FROM {table} WHERE {condition}")
            conn.commit()
        return jsonify({"status": "Record deleted"})

# -------------------------------------------------------------------
# File Upload Endpoint
# -------------------------------------------------------------------
    def file_upload(self, route='/upload', upload_folder='uploads'):
        os.makedirs(upload_folder, exist_ok=True)
        endpoint_name = self._unique_route_name(route)

        @self.app.route(route, methods=['POST'], endpoint=endpoint_name)
        def upload():
            if 'file' not in request.files:
                return jsonify({"error": "No file part"})
            file = request.files['file']
            if file.filename == '':
                return jsonify({"error": "No selected file"})
            path = os.path.join(upload_folder, file.filename)
            file.save(path)
            return jsonify({"success": True, "path": path})

# -------------------------------------------------------------------
# Serve Static or Dynamic Files
# -------------------------------------------------------------------
    def serve_file(self, route='/file', file_path=""):
        endpoint_name = self._unique_route_name(route)

        @self.app.route(route, endpoint=endpoint_name)
        def serve():
            if os.path.exists(file_path):
                return send_file(file_path)
            return "File not found", 404

# -------------------------------------------------------------------
# Vector Plotting API
# -------------------------------------------------------------------
    def vector_plot(self, route='/plot_vector'):
        endpoint_name = self._unique_route_name(route)

        @self.app.route(route, methods=['POST'], endpoint=endpoint_name)
        def plot_vector():
            data = request.get_json(force=True)
            vector = data.get("vector", [0, 0])
            origin = data.get("origin", [0, 0])
            color = data.get("color", "r")

            fig, ax = plt.subplots()
            ax.quiver(origin[0], origin[1], vector[0], vector[1],
                      angles='xy', scale_units='xy', scale=1, color=color)
            ax.set_xlim(-10, 10)
            ax.set_ylim(-10, 10)
            ax.grid(True)
            ax.axhline(0, color='black', linewidth=0.5)
            ax.axvline(0, color='black', linewidth=0.5)

            buf = io.BytesIO()
            plt.savefig(buf, format='png')
            buf.seek(0)
            encoded = base64.b64encode(buf.read()).decode('utf-8')
            plt.close(fig)
            return jsonify({"image_base64": encoded})

# -------------------------------------------------------------------
# Health Check Endpoint
# -------------------------------------------------------------------
    def health(self, route="/health"):
        endpoint_name = self._unique_route_name(route)

        @self.app.route(route, endpoint=endpoint_name)
        def health_check():
            return jsonify({"status": "online", "message": "Server is healthy"})

# -------------------------------------------------------------------
# Run Server
# -------------------------------------------------------------------
    def run(self, live_refresh=True, port=5200):
        logging.info(f"Server starting on port {port} ...")
        self.app.run(debug=live_refresh, host='0.0.0.0', port=port)

    
#_____________________________Databases_______________________________
class Database():
    def create(self,path_of_db,name_of_table="Default_Table",columns=[{"name":"ID","datatype":"INT","constraint":""},{"name":"Name","datatype":"TEXT","constraint":""}]):
        conn=sqlite3.connect(fr"{path_of_db}")
        cursor=conn.cursor()
        cursor.execute(f"""
                CREATE TABLE IF NOT EXISTS {name_of_table} ({columns[0].get("name")} {columns[0].get("datatype")} {columns[0].get("constraint")});
                """)
        for i in range(1,len(columns)):
            conn.execute(f"""
                        
                        ALTER TABLE {name_of_table}
                        ADD {columns[i].get("name")} {columns[i].get("datatype")} {columns[i].get("constraint")};
                        """)
        conn.commit()
        conn.close()
    def retrieve_data(self,path_of_db,name_of_table):
        conn=sqlite3.connect(path_of_db)
        cursor=conn.cursor()
        cursor.execute(f"""
                    SELECT * FROM {name_of_table}
                    """)
        rows = cursor.fetchall()
        column_names = [description[0] for description in cursor.description]
        result = [dict(zip(column_names, row)) for row in rows]
        conn.close()
        return result
    

        
    def run_query(self,path_of_db,name_of_table,query):
        conn=sqlite3.connect(path_of_db)
        cursor=conn.cursor()
        cursor.execute(f"""
                    
                    {query.format(table=name_of_table)}
                    """)
        conn.commit()
        conn.close()
    
    
    def insert_data(self, path_of_db, name_of_table,data={"ID": [1, 2, 3, 4, 5],"Name": ["Name1", "Name2", "Name3", "Name4", "Name5"]}):
    
        conn = sqlite3.connect(path_of_db)
        cursor = conn.cursor()
        
        columns = list(data.keys())
        rows = len(data[columns[0]])  # Number of entries
        
        for j in range(rows):
            # Collect row values
            values = []
            for col in columns:
                val = data[col][j]
                if isinstance(val, str):
                    val = f'"{val}"'  # Add quotes for strings
                values.append(str(val))
            
            cursor.execute(f"""
                INSERT INTO {name_of_table} ({', '.join(columns)})
                VALUES ({', '.join(values)});
            """)
        
        conn.commit()
        conn.close()
        
#________________________________Media_____________________________________

class Media():
    def image(self,path_of_image,width,height,title="NeuraPy Image"):
        img = Image.open(path_of_image)
        img = img.resize((width, height))
        plt.imshow(img)
        plt.title(title)
        plt.axis('off')
        plt.show()
        
    def audio(self,path_of_audio):
        try:
            pygame.mixer.init()
            pygame.mixer.music.load(path_of_audio)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
        except Exception as e:
            print(f"Error playing audio: {e}")
    def video(self,path_of_video,width,height):
        cap = cv2.VideoCapture(path_of_video)
        if not cap.isOpened():
            print("Error: Cannot open video.")
            return
        
        while cap.isOpened():
            ret, frame = cap.read()
            if not ret:
                break
            frame = cv2.resize(frame, (width, height))
            cv2.imshow('Video', frame)
            
            # 'q' to quit
            if cv2.waitKey(25) & 0xFF == ord('q'):
                break
        
        cap.release()
        cv2.destroyAllWindows()




#________________________________Machine Learning_________________________

#__________________________________AI_________________________________________
class AI:
    def google_gemini(self, api_key, model, prompt):
        """
        Send a prompt to Google Gemini API and return the reply.
        """
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
            headers = {"Content-Type": "application/json"}
            data = {"contents": [{"parts": [{"text": prompt}]}]}

            response = requests.post(url, headers=headers, json=data)
            response.raise_for_status()

            result = response.json()
            return result["candidates"][0]["content"]["parts"][0]["text"]
        except Exception as e:
            return f"Error (Gemini): {e}"

    def openai_chatgpt(self, api_key, model, prompt):
        """
        Send a prompt to OpenAI ChatGPT API and return the reply.
        """
        try:
            url = "https://api.openai.com/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            data = {
                "model": model,
                "messages": [{"role": "user", "content": prompt}]
            }

            response = requests.post(url, headers=headers, json=data)
            response.raise_for_status()

            result = response.json()
            return result["choices"][0]["message"]["content"].strip()
        except Exception as e:
            return f"Error (ChatGPT): {e}"

#___________________________________Vectors______________________________________



class Vector2D:
    def __init__(self, x: float = 0, y: float = 0):
        self.x = x
        self.y = y

    def __repr__(self):
        return f"Vector2D({self.x}, {self.y})"
    def to_list(self):
        return [self.x,self.y]

    def add(self, *vectors):
        new_x, new_y = self.x, self.y
        for v in vectors:
            new_x += v.x
            new_y += v.y
        return Vector2D(new_x, new_y)

    def subtract(self, *vectors):
        new_x, new_y = self.x, self.y
        for v in vectors:
            new_x -= v.x
            new_y -= v.y
        return Vector2D(new_x, new_y)


    @staticmethod
    def dot_product(v1, v2):
        return v1.x * v2.x + v1.y * v2.y


    @staticmethod
    def cross_product(v1, v2):
        return v1.x * v2.y - v1.y * v2.x

    def magnitude(self):
        return math.sqrt(self.x**2 + self.y**2)

    def unit_vector(self):
        mag = self.magnitude()
        if mag == 0:
            raise ValueError("Cannot compute unit vector of zero vector.")
        return Vector2D(self.x / mag, self.y / mag)

    @staticmethod
    def angle_between(v1, v2 ,angle:bool=True):
        dot = Vector2D.dot_product(v1, v2)
        mag1 = v1.magnitude()
        mag2 = v2.magnitude()
        if mag1 == 0 or mag2 == 0:
            raise ValueError("Cannot compute angle with zero-length vector.")
        cos_theta = dot / (mag1 * mag2)
        cos_theta = max(-1, min(1, cos_theta))
        if angle:
            return math.degrees(math.acos(cos_theta))
        else:
            return math.acos(cos_theta)
    
    @classmethod
    def from_list(cls, data):
        if len(data) != 2:
            raise ValueError("List must contain exactly 2 values.")
        return cls(data[0], data[1])

    def plot(self, *others, color='r', show=True, title="2D Vector Plot"):
        """
        Plot this vector (and optionally others) from the origin.
        """
        plt.figure()
        ax = plt.gca()
        vectors = [self] + list(others)
        colors = ['r', 'b', 'g', 'm', 'c', 'y', 'k']

        for i, v in enumerate(vectors):
            col = colors[i % len(colors)] if len(vectors) > 1 else color
            ax.quiver(0, 0, v.x, v.y, angles='xy', scale_units='xy', scale=1, color=col, label=str(v))

        limit = max(max(abs(v.x), abs(v.y)) for v in vectors) + 1
        ax.set_xlim(-limit, limit)
        ax.set_ylim(-limit, limit)
        ax.axhline(0, color='black', linewidth=0.5)
        ax.axvline(0, color='black', linewidth=0.5)
        plt.grid(True)
        plt.title(title)
        plt.legend()
        if show:
            plt.show()


class Vector3D:
    def __init__(self, x: float = 0, y: float = 0, z: float = 0):
        self.x = x
        self.y = y
        self.z = z

    def __repr__(self):
        return f"Vector3D({self.x}, {self.y}, {self.z})"


    def add(self, *vectors):
        new_x, new_y, new_z = self.x, self.y, self.z
        for v in vectors:
            new_x += v.x
            new_y += v.y
            new_z += v.z
        return Vector3D(new_x, new_y, new_z)


    def subtract(self, *vectors):
        new_x, new_y, new_z = self.x, self.y, self.z
        for v in vectors:
            new_x -= v.x
            new_y -= v.y
            new_z -= v.z
        return Vector3D(new_x, new_y, new_z)

    @staticmethod
    def dot_product(v1, v2):
        return v1.x * v2.x + v1.y * v2.y + v1.z * v2.z

    @staticmethod
    def cross_product(v1, v2):
        cx = v1.y * v2.z - v1.z * v2.y
        cy = v1.z * v2.x - v1.x * v2.z
        cz = v1.x * v2.y - v1.y * v2.x
        return Vector3D(cx, cy, cz)

    def magnitude(self):
        return math.sqrt(self.x**2 + self.y**2 + self.z**2)

    def unit_vector(self):
        mag = self.magnitude()
        if mag == 0:
            raise ValueError("Cannot compute unit vector of zero vector.")
        return Vector3D(self.x / mag, self.y / mag, self.z / mag)

    @staticmethod
    def angle_between(v1, v2, degrees=True):
        dot = Vector3D.dot_product(v1, v2)
        mag1 = v1.magnitude()
        mag2 = v2.magnitude()
        if mag1 == 0 or mag2 == 0:
            raise ValueError("Cannot compute angle with zero-length vector.")
        cos_theta = dot / (mag1 * mag2)
        cos_theta = max(-1, min(1, cos_theta))
        angle = math.acos(cos_theta)
        return math.degrees(angle) if degrees else angle

    def to_list(self):
        return [self.x, self.y, self.z]

    @classmethod
    def from_list(cls, data):
        if len(data) != 3:
            raise ValueError("List must contain exactly 3 values.")
        return cls(data[0], data[1], data[2])
    
    def plot(self, *others, color='r', show=True, title="3D Vector Plot"):
        """
        Plot this vector (and optionally others) in 3D space.
        """
        fig = plt.figure()
        ax = fig.add_subplot(111, projection='3d')
        vectors = [self] + list(others)
        colors = ['r', 'b', 'g', 'm', 'c', 'y', 'k']

        for i, v in enumerate(vectors):
            col = colors[i % len(colors)] if len(vectors) > 1 else color
            ax.quiver(0, 0, 0, v.x, v.y, v.z, color=col, label=str(v))

        limit = max(max(abs(v.x), abs(v.y), abs(v.z)) for v in vectors) + 1
        ax.set_xlim(-limit, limit)
        ax.set_ylim(-limit, limit)
        ax.set_zlim(-limit, limit)
        ax.set_xlabel('X')
        ax.set_ylabel('Y')
        ax.set_zlabel('Z')
        plt.title(title)
        plt.legend()
        if show:
            plt.show()
        
#_____________________________Calender_________________________________
def Calender(year=None, month=None):
    now = datetime.datetime.now()
    if year is None:
        year = now.year
    if month is None:
        month = now.month
    cal = calendar.month(year, month)
    print(cal)


#_________________________Readers_____________________________________
class Reader():
    def html_reader(self,path):
        with open(rf"{path}","r",encoding="utf-8") as html:
            content=html.read()
        return content
    
    def excel_reader(self,path):
        data=pd.read_excel(rf"{path}")
        return data.to_string()
    
    def docx_reader(self,path):
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    
    def json_reader(self,path):
        with open(path, 'r', encoding='utf-8') as file:
            data = json.load(file)  
        return data
    
    def csv_reader(self,path):
        data=pd.read_csv(rf"{path}")
        return data.to_string()
    
    def text_reader(self,path):
        with open(path, 'r', encoding='utf-8') as file:
            content = file.read()  
        return content
    
    def pdf_reader(self,path):
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    def markdown_reader(self,path):
        with open(path, 'r', encoding='utf-8') as file:
            content = file.read()  # Read as plain text
        return content
    def xml_reader(self,path):
        with open(path, 'r', encoding='utf-8') as file:
            content = file.read()
        soup = BeautifulSoup(content, 'xml')
        return soup.prettify()
#___________________________Converters_____________________________
class Converter():
    def pdf_to_docx(self,pdf_path,docx_path):
        cv = pdf2docx.Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
    
    def docx_to_pdf(self,docx_path,pdf_path):
        docx2pdf.convert(docx_path, pdf_path)

    def pdf_to_text(self,pdf_path,txt_path):
        pdf = fitz.open(pdf_path)
        text = ""

        for page in pdf:
            text += page.get_text("text") + "\n"

        pdf.close()

        if txt_path:
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(text)
    
    def text_to_pdf(self,txt_path,pdf_path):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        with open(txt_path, "r", encoding="utf-8") as f:
            text = f.read()

        pdf.multi_cell(0, 10, text)
        pdf.output(pdf_path)
        
    def excel_to_csv(self,excel_path,csv_path):
        df = pd.read_excel(excel_path)
        df.to_csv(csv_path, index=False)
        
    def csv_to_excel(self,csv_path,xlsx_path):
        df = pd.read_csv(csv_path)
        df.to_excel(xlsx_path, index=False, engine='openpyxl')
        
    def json_to_csv(self,json_path,csv_path):
        df = pd.read_json(json_path)
        df.to_csv(csv_path, index=False)
        
    def json_to_excel(self,json_path,excel_path):
        data = pd.read_json(json_path)
        data.to_excel(excel_path, index=False)
        
    def csv_to_json(self,csv_path,json_path):
        data = pd.read_csv(csv_path)
        data.to_json(json_path, orient="records", indent=4)
        
    def excel_to_json(self,excel_path,json_path):
        data = pd.read_excel(excel_path)
        data.to_json(json_path, orient="records", indent=4)
    def markdown_to_html(self,markdown_path,html_path):
        with open(markdown_path, "r", encoding="utf-8") as md_file:
            md_content = md_file.read()

        html_content = markdown.markdown(md_content, extensions=[
        "fenced_code",  # ``` code blocks ```
        "tables",       # Markdown tables
        "attr_list"     # Attributes like {: .class }
        ])

        html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{markdown_path}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            margin: 40px;
            line-height: 1.6;
            background-color: #fafafa;
        }}
        pre, code {{
            background: #f5f5f5;
            padding: 5px;
            border-radius: 4px;
        }}
        table {{
            border-collapse: collapse;
            margin-top: 10px;
        }}
        th, td {{
            border: 1px solid #ccc;
            padding: 8px;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

        with open(html_path, "w", encoding="utf-8") as html_file:
            html_file.write(html_template)
            
            
            
            
    def string_to_pdf(self, text: str, pdf_path: str, font="Arial", size=12):
        """
        Create a PDF from a string.
        """
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font(font, size=size)
        pdf.multi_cell(0, 10, text)
        pdf.output(pdf_path)

    # 🔹 NEW METHOD 2: Convert string directly to DOCX
    def string_to_docx(self, text: str, docx_path: str, title: str = None):
        """
        Create a DOCX file from a string.
        """
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        doc.add_paragraph(text)
        doc.save(docx_path)

#__________________________Calculus_______________________________
class Calculus:
    def __init__(self):
        from sympy import symbols, sin, cos

    # 1️⃣ Derivative
    def derivative(self, expr, var='x', order=1):
        x = symbols(var)
        expr = sympify(expr)
        return diff(expr, x, order)

    # 2️⃣ Integral
    def integral(self, expr, var='x', lower=None, upper=None):
        x = symbols(var)
        expr = sympify(expr)
        if lower is not None and upper is not None:
            return integrate(expr, (x, lower, upper))
        return integrate(expr, x)

    # 3️⃣ Limit
    def calc_limit(self, expr, var='x', point=0):
        x = symbols(var)
        expr = sympify(expr)
        return limit(expr, x, point)

    # 4️⃣ Partial Derivative
    def partial_derivative(self, expr, var):
        expr = sympify(expr)
        return diff(expr, symbols(var))

    # 5️⃣ Gradient
    def gradient(self, expr, vars_list):
        vars = symbols(vars_list)
        expr = sympify(expr)
        return Matrix([diff(expr, v) for v in vars])


    # 7️⃣ Curl
    def calc_curl(self, Fx, Fy, Fz):
        N = CoordSys3D('N')
        Fx, Fy, Fz = sympify(Fx), sympify(Fy), sympify(Fz)
        F = Fx * N.i + Fy * N.j + Fz * N.k
        return curl(F)

    # 8️⃣ Taylor / Maclaurin Series
    def taylor_series(self, expr, var='x', point=0, n=5):
        x = symbols(var)
        expr = sympify(expr)
        return series(expr, x, point, n)

    # 9️⃣ Summation
    def summation_func(self, expr, var='n', start=1, end=10):
        n = symbols(var)
        expr = sympify(expr)
        return summation(expr, (n, start, end))

    # 🔟 Differential Equation Solver (fixed)
    def solve_diff_eq(self, equation, var='x'):
        x = symbols(var)
        y = Function('y')
        expr = sympify(equation)  # ✅ FIXED: parse string to symbolic expression
        eq = Eq(y(x).diff(x), expr)
        return dsolve(eq)

    # 11️⃣ Jacobian Matrix
    def jacobian_matrix(self, funcs, vars_list):
        vars = symbols(vars_list)
        f = Matrix([sympify(fn) for fn in funcs])
        return f.jacobian(vars)

    # 12️⃣ Hessian Matrix
    def hessian_matrix(self, expr, vars_list):
        vars = symbols(vars_list)
        f = sympify(expr)
        return Matrix([[diff(f, vi, vj) for vj in vars] for vi in vars])

#__________________________Matrices_____________________________
class Matrices:
    def __init__(self, data: list):
        self.data = data
        self.rows = len(data)
        self.cols = len(data[0]) if data else 0

    def shape(self):
        return (self.rows, self.cols)

    def size(self):
        return self.rows * self.cols

    def copy(self):
        return Matrices([row[:] for row in self.data])

    def add(self, other):
        # Element-wise addition
        return Matrices([
            [self.data[i][j] + other.data[i][j] for j in range(self.cols)]
            for i in range(self.rows)
        ])

    def subtract(self, other):
        return Matrices([
            [self.data[i][j] - other.data[i][j] for j in range(self.cols)]
            for i in range(self.rows)
        ])

    def multiply(self, other):
        # Matrix multiplication (not element-wise)
        result = [
            [sum(self.data[i][k] * other.data[k][j] for k in range(self.cols))
             for j in range(other.cols)]
            for i in range(self.rows)
        ]
        return Matrices(result)

    def scalar_multiply(self, value):
        return Matrices([
            [self.data[i][j] * value for j in range(self.cols)]
            for i in range(self.rows)
        ])

    def transpose(self):
        return Matrices([
            [self.data[j][i] for j in range(self.rows)]
            for i in range(self.cols)
        ])

    def determinant(self):
        if self.rows != self.cols:
            raise ValueError("Determinant only defined for square matrices")
        if self.rows == 1:
            return self.data[0][0]
        if self.rows == 2:
            return self.data[0][0]*self.data[1][1] - self.data[0][1]*self.data[1][0]
        det = 0
        for c in range(self.cols):
            minor = [row[:c] + row[c+1:] for row in self.data[1:]]
            det += ((-1)**c) * self.data[0][c] * Matrices(minor).determinant()
        return det

    def inverse(self):
        det = self.determinant()
        if det == 0:
            raise ValueError("Matrix is singular and cannot be inverted")
        if self.rows == 2:
            a, b = self.data[0]
            c, d = self.data[1]
            inv = [[d, -b], [-c, a]]
            return Matrices(inv).scalar_multiply(1/det)
        cofactors = []
        for r in range(self.rows):
            cofactor_row = []
            for c in range(self.cols):
                minor = [row[:c] + row[c+1:] for i, row in enumerate(self.data) if i != r]
                cofactor_row.append(((-1)**(r+c)) * Matrices(minor).determinant())
            cofactors.append(cofactor_row)
        cofactors = Matrices(cofactors).transpose()
        return cofactors.scalar_multiply(1/det)

    def is_square(self):
        return self.rows == self.cols

    def flatten(self):
        return [x for row in self.data for x in row]

    def trace(self):
        return sum(self.data[i][i] for i in range(min(self.rows, self.cols)))

    def rank(self):
        import numpy as np
        return np.linalg.matrix_rank(self.data)

    def power(self, n):
        if not self.is_square():
            raise ValueError("Matrix must be square for power operation")
        result = Matrices.identity(self.rows)
        for _ in range(n):
            result = result.multiply(self)
        return result

    @staticmethod
    def identity(n):
        return Matrices([[1 if i == j else 0 for j in range(n)] for i in range(n)])

    def __str__(self):
        return '\n'.join(['\t'.join(map(str, row)) for row in self.data])
#____________________________________Machine Learning_________________________________________
# ===========================================================
# ⚙️ NeuraPython_ML — Unified Scikit-learn Wrapper
# Author: Ibrahim Shahid
# ===========================================================

import numpy as np
import pandas as pd
import joblib

from sklearn import (
    datasets, preprocessing, model_selection, metrics,
    linear_model, tree, neighbors, svm, ensemble, cluster,
    decomposition
)

class NeuraPython_ML:
    def __init__(self):
        print("🧠 NeuraPython_ML initialized ")
        self.models = {}
        self.scaler = None
        self.encoder = None

# -----------------------------------------------------------
# 📥 DATASET LOADING
# -----------------------------------------------------------
    def load_builtin_dataset(self, name="iris"):
        name = name.lower()
        if name == "iris": data = datasets.load_iris()
        elif name == "digits": data = datasets.load_digits()
        elif name == "wine": data = datasets.load_wine()
        elif name == "breast_cancer": data = datasets.load_breast_cancer()
        else: raise ValueError("❌ Unknown dataset name.")
        return pd.DataFrame(data.data, columns=data.feature_names), data.target

# -----------------------------------------------------------
# 🔢 DATA PREPROCESSING
# -----------------------------------------------------------
    def normalize(self, X):
        self.scaler = preprocessing.Normalizer()
        return self.scaler.fit_transform(X)

    def standardize(self, X):
        self.scaler = preprocessing.StandardScaler()
        return self.scaler.fit_transform(X)

    def minmax_scale(self, X):
        self.scaler = preprocessing.MinMaxScaler()
        return self.scaler.fit_transform(X)

    def encode_labels(self, y):
        self.encoder = preprocessing.LabelEncoder()
        return self.encoder.fit_transform(y)

    def one_hot_encode(self, X):
        return preprocessing.OneHotEncoder().fit_transform(X).toarray()

# -----------------------------------------------------------
# 📊 TRAIN/TEST SPLIT
# -----------------------------------------------------------
    def split(self, X, y, test_size=0.2, random_state=42):
        return model_selection.train_test_split(X, y, test_size=test_size, random_state=random_state)

# -----------------------------------------------------------
# 🤖 MODEL CREATION (Unified Interface)
# -----------------------------------------------------------
    def create_model(self, model_name, **kwargs):
        name = model_name.lower()
        if name == "linear_regression": model = linear_model.LinearRegression(**kwargs)
        elif name == "logistic_regression": model = linear_model.LogisticRegression(**kwargs)
        elif name == "decision_tree": model = tree.DecisionTreeClassifier(**kwargs)
        elif name == "random_forest": model = ensemble.RandomForestClassifier(**kwargs)
        elif name == "svm": model = svm.SVC(**kwargs)
        elif name == "knn": model = neighbors.KNeighborsClassifier(**kwargs)
        elif name == "gradient_boosting": model = ensemble.GradientBoostingClassifier(**kwargs)
        elif name == "naive_bayes": 
            from sklearn.naive_bayes import GaussianNB
            model = GaussianNB(**kwargs)
        elif name == "kmeans": model = cluster.KMeans(**kwargs)
        elif name == "pca": model = decomposition.PCA(**kwargs)
        else:
            raise ValueError(f"❌ Unsupported model: {model_name}")
        
        self.models[model_name] = model
        print(f"✅ Model '{model_name}' created.")
        return model

# -----------------------------------------------------------
# 🧩 MODEL TRAINING & PREDICTION
# -----------------------------------------------------------
    def train(self, model_name, X_train, y_train):
        model = self.models.get(model_name)
        if model is None:
            raise ValueError(f"❌ Model '{model_name}' not found.")
        model.fit(X_train, y_train)
        print(f"🚀 Model '{model_name}' trained successfully.")
        return model

    def predict(self, model_name, X_test):
        model = self.models.get(model_name)
        if model is None:
            raise ValueError(f"❌ Model '{model_name}' not found.")
        return model.predict(X_test)

# -----------------------------------------------------------
# 📈 MODEL EVALUATION
# -----------------------------------------------------------
    def evaluate(self, y_true, y_pred):
        return {
            "accuracy": metrics.accuracy_score(y_true, y_pred),
            "precision": metrics.precision_score(y_true, y_pred, average='weighted', zero_division=0),
            "recall": metrics.recall_score(y_true, y_pred, average='weighted', zero_division=0),
            "f1_score": metrics.f1_score(y_true, y_pred, average='weighted', zero_division=0)
        }

    def confusion_matrix(self, y_true, y_pred):
        return metrics.confusion_matrix(y_true, y_pred)

    def classification_report(self, y_true, y_pred):
        return metrics.classification_report(y_true, y_pred)

# -----------------------------------------------------------
# 🧮 DIMENSIONALITY REDUCTION
# -----------------------------------------------------------
    def apply_pca(self, X, n_components=2):
        pca = decomposition.PCA(n_components=n_components)
        return pca.fit_transform(X)

# -----------------------------------------------------------
# 💾 MODEL SAVING & LOADING
# -----------------------------------------------------------
    def save_model(self, model_name, path):
        model = self.models.get(model_name)
        if model is None:
            raise ValueError(f"❌ Model '{model_name}' not found.")
        joblib.dump(model, path)
        print(f"💾 Model '{model_name}' saved at {path}")

    def load_model(self, path, model_name="loaded_model"):
        model = joblib.load(path)
        self.models[model_name] = model
        print(f"📂 Model '{model_name}' loaded from {path}")
        return model

# -----------------------------------------------------------
# 🧠 CROSS VALIDATION & GRID SEARCH
# -----------------------------------------------------------
    def cross_validate(self, model, X, y, cv=5):
        scores = model_selection.cross_val_score(model, X, y, cv=cv)
        return {"mean": np.mean(scores), "scores": scores}

    def grid_search(self, model, params, X, y, cv=5):
        search = model_selection.GridSearchCV(model, params, cv=cv)
        search.fit(X, y)
        return search.best_estimator_, search.best_params_, search.best_score_

# -----------------------------------------------------------
# 🔍 CLUSTERING UTILITIES
# -----------------------------------------------------------
    def kmeans_cluster(self, X, n_clusters=3):
        model = cluster.KMeans(n_clusters=n_clusters)
        y_pred = model.fit_predict(X)
        return model, y_pred

# -----------------------------------------------------------
# 🧩 FEATURE SELECTION
# -----------------------------------------------------------
    def feature_importances(self, model_name, feature_names=None):
        model = self.models.get(model_name)
        if model is None:
            raise ValueError(f"❌ Model '{model_name}' not found.")
        if hasattr(model, "feature_importances_"):
            imp = model.feature_importances_
            if feature_names is not None:
                return dict(zip(feature_names, imp))
            return imp
        else:
            raise AttributeError("⚠️ Model has no feature_importances_ attribute.")

# -----------------------------------------------------------
# 🧾 SUMMARY
# -----------------------------------------------------------
    def summary(self):
        print("=== NeuraPython_ML Models ===")
        for name, model in self.models.items():
            print(f"• {name}: {type(model).__name__}")

#_______________________________________Neural Network________________________________________


class NeuralNetwork:
    def __init__(self, backend='torch', device=None):
        self.backend = backend.lower()
        self.device = device or ("cuda" if torch.cuda.is_available() else "cpu")
        self.model = None

        if self.backend not in ['torch', 'tf']:
            raise ValueError("Backend must be 'torch' or 'tf'")

        # dynamically attach all functions and submodules
        self._attach_backend()

    # -------------------------------------------------------------------------
    # Attach backend attributes dynamically
    # -------------------------------------------------------------------------
    def _attach_backend(self):
        """Attach TensorFlow or Torch modules as attributes for universal access."""
        if self.backend == 'torch':
            self.tensor_lib = torch
            self.nn = torch.nn
            self.optim = torch.optim
            self.utils = torch.utils
            self.vision = None
        else:
            self.tensor_lib = tf
            self.nn = tf.keras.layers
            self.optim = tf.keras.optimizers
            self.losses = tf.keras.losses
            self.metrics = tf.keras.metrics

    # -------------------------------------------------------------------------
    # Universal tensor creator
    # -------------------------------------------------------------------------
    def tensor(self, data, dtype=None, requires_grad=False):
        if self.backend == 'torch':
            return torch.tensor(data, dtype=dtype or torch.float32, requires_grad=requires_grad).to(self.device)
        elif self.backend == 'tf':
            return tf.convert_to_tensor(data, dtype=dtype or tf.float32)

    # -------------------------------------------------------------------------
    # Universal model creation
    # -------------------------------------------------------------------------
    def Sequential(self, layers_list):
        """Create a simple sequential model compatible with backend."""
        if self.backend == 'torch':
            seq_layers = []
            for i in range(len(layers_list) - 1):
                seq_layers.append(torch.nn.Linear(layers_list[i], layers_list[i + 1]))
                if i < len(layers_list) - 2:
                    seq_layers.append(torch.nn.ReLU())
            self.model = torch.nn.Sequential(*seq_layers).to(self.device)
        else:
            self.model = tf.keras.Sequential()
            for i in range(len(layers_list) - 1):
                self.model.add(tf.keras.layers.Dense(
                    layers_list[i + 1],
                    input_dim=layers_list[i] if i == 0 else None,
                    activation='relu' if i < len(layers_list) - 2 else None
                ))

    # -------------------------------------------------------------------------
    # Universal compile/train/predict
    # -------------------------------------------------------------------------
    def compile(self, optimizer='adam', loss='mse'):
        if self.backend == 'torch':
            self.loss_fn = torch.nn.MSELoss() if loss == 'mse' else loss
            self.optimizer = torch.optim.Adam(self.model.parameters()) if optimizer == 'adam' else optimizer
        else:
            self.model.compile(optimizer=getattr(tf.keras.optimizers, optimizer.capitalize())(),
                               loss=loss)

    def fit(self, X, y, epochs=5):
        if self.backend == 'torch':
            X = self.tensor(X)
            y = self.tensor(y)
            for epoch in range(epochs):
                self.optimizer.zero_grad()
                preds = self.model(X)
                loss = self.loss_fn(preds, y)
                loss.backward()
                self.optimizer.step()
                print(f"[Torch] Epoch {epoch+1}/{epochs}, Loss: {loss.item():.4f}")
        else:
            self.model.fit(X, y, epochs=epochs, verbose=1)

    def predict(self, X):
        if self.backend == 'torch':
            with torch.no_grad():
                preds = self.model(self.tensor(X))
            return preds.cpu().numpy()
        else:
            return self.model.predict(X)

    # -------------------------------------------------------------------------
    # Dynamic universal access
    # -------------------------------------------------------------------------
    def __getattr__(self, name):
        """Dynamic universal access to backend functions and submodules."""
        try:
            if hasattr(self.tensor_lib, name):
                return getattr(self.tensor_lib, name)
            elif self.backend == 'torch' and hasattr(torch.nn, name):
                return getattr(torch.nn, name)
            elif self.backend == 'tf':
                # check keras modules
                if hasattr(tf.keras.layers, name):
                    return getattr(tf.keras.layers, name)
                if hasattr(tf.keras.optimizers, name):
                    return getattr(tf.keras.optimizers, name)
                if hasattr(tf.keras.losses, name):
                    return getattr(tf.keras.losses, name)
        except Exception:
            pass
        raise AttributeError(f"'{self.backend}' backend has no attribute '{name}'")



#___________________________________________QR code_____________________________________


class QR_Code:
    @staticmethod
    def generator(data, file_path="qrcode.png", box_size=10, border=4):
        """
        Generate a QR code image from the given data.
        """
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=box_size,
            border=border,
        )
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        img.save(file_path)
        return f"QR code saved as {file_path}"

    @staticmethod
    def reader(file_path):
        """
        Read and decode a QR code from an image.
        """
        img = cv2.imread(file_path)
        detector = cv2.QRCodeDetector()
        data, points, _ = detector.detectAndDecode(img)
        if data:
            return f"Decoded data: {data}"
        else:
            return "No QR code detected or unable to decode."
        
#_______________________________Advance Maths________________________________
import math
import statistics as stats
import random

class Advanced_Maths:
    def __init__(self):
        print("🧮 Advanced_Maths initialized — arithmetic, algebra, and statistics toolkit.")

# -----------------------------------------------------------
# 🔢 BASIC ARITHMETIC (MULTIPLE ARGUMENTS)
# -----------------------------------------------------------
    def add(self, *args):
        """Add all numbers together."""
        return sum(args)

    def subtract(self, *args):
        """Subtract all numbers in sequence."""
        if len(args) < 2:
            raise ValueError("Need at least two numbers to subtract.")
        result = args[0]
        for num in args[1:]:
            result -= num
        return result

    def multiply(self, *args):
        """Multiply all numbers together."""
        result = 1
        for num in args:
            result *= num
        return result

    def divide(self, *args):
        """Divide numbers in sequence (a / b / c / ...)."""
        if len(args) < 2:
            raise ValueError("Need at least two numbers to divide.")
        result = args[0]
        for num in args[1:]:
            if num == 0:
                raise ZeroDivisionError("Cannot divide by zero.")
            result /= num
        return result

    def power(self, base, exponent): return base ** exponent
    def sqrt(self, x): return math.sqrt(x)
    def mod(self, a, b): return a % b

# -----------------------------------------------------------
# 🔣 NUMBER THEORY
# -----------------------------------------------------------
    def factorial(self, n): return math.factorial(n)
    def combination(self, n, r): return math.comb(n, r)
    def permutation(self, n, r): return math.perm(n, r)
    def gcd(self, a, b): return math.gcd(a, b)
    def lcm(self, a, b): return math.lcm(a, b)
    def is_prime(self, n):
        if n < 2: return False
        for i in range(2, int(n ** 0.5) + 1):
            if n % i == 0:
                return False
        return True

# -----------------------------------------------------------
# 📊 STATISTICS
# -----------------------------------------------------------
    def mean(self, data): return stats.mean(data)
    def median(self, data): return stats.median(data)
    def mode(self, data): return stats.mode(data)
    def variance(self, data): return stats.variance(data)
    def stdev(self, data): return stats.stdev(data)
    def data_range(self, data): return max(data) - min(data)

# -----------------------------------------------------------
# 🎲 PROBABILITY & RANDOM
# -----------------------------------------------------------
    def random_choice(self, data): return random.choice(data)
    def random_sample(self, data, k): return random.sample(data, k)
    def random_int(self, a, b): return random.randint(a, b)
    def random_float(self, a=0.0, b=1.0): return random.uniform(a, b)
    def probability(self, favorable, total):
        return favorable / total if total > 0 else 0

# -----------------------------------------------------------
# 🔢 SEQUENCES AND SERIES
# -----------------------------------------------------------
    def arithmetic_series(self, a1, d, n):
        """Sum of arithmetic progression."""
        return n / 2 * (2 * a1 + (n - 1) * d)

    def geometric_series(self, a1, r, n):
        """Sum of geometric progression."""
        if r == 1:
            return a1 * n
        return a1 * (1 - r ** n) / (1 - r)

# -----------------------------------------------------------
# 📏 CONVERSIONS & MISC
# -----------------------------------------------------------
    def deg_to_rad(self, deg): return math.radians(deg)
    def rad_to_deg(self, rad): return math.degrees(rad)
    def absolute(self, x): return abs(x)

#________________________________Physics (Beginer to Advanced)______________________


class Physics:
    def __init__(self):
        print("⚛️ Physics module initialized — classical, relativity, and quantum toolkit.")

# -----------------------------------------------------------
# ⚙️ CONSTANTS
# -----------------------------------------------------------
    G = 6.67430e-11   # Gravitational constant (N·m²/kg²)
    g = 9.81          # Acceleration due to gravity (m/s²)
    c = 3.0e8         # Speed of light (m/s)
    R = 8.314         # Gas constant (J/mol·K)
    k = 1.380649e-23  # Boltzmann constant (J/K)
    e = 1.602176634e-19  # Elementary charge (C)
    h = 6.62607015e-34   # Planck constant (J·s)
    h_bar = h / (2 * math.pi)  # Reduced Planck’s constant (ħ)

# -----------------------------------------------------------
# 🧭 MECHANICS
# -----------------------------------------------------------
    def velocity(self, distance, time):
        return distance / time

    def acceleration(self, v_final, v_initial, time):
        return (v_final - v_initial) / time

    def force(self, mass, acceleration):
        return mass * acceleration

    def weight(self, mass):
        return mass * self.g

    def momentum(self, mass, velocity):
        return mass * velocity

    def kinetic_energy(self, mass, velocity):
        return 0.5 * mass * velocity ** 2

    def potential_energy(self, mass, height):
        return mass * self.g * height

    def power(self, work, time):
        return work / time

    def work_done(self, force, distance, angle=0):
        return force * distance * math.cos(math.radians(angle))

# -----------------------------------------------------------
# 🌌 RELATIVITY
# -----------------------------------------------------------
    def mass_energy_equivalence(self, mass):
        """E = mc²"""
        return mass * (self.c ** 2)

    def relativistic_mass(self, rest_mass, velocity):
        """m = m₀ / sqrt(1 - v²/c²)"""
        if velocity >= self.c:
            raise ValueError("Velocity cannot reach or exceed speed of light.")
        return rest_mass / math.sqrt(1 - (velocity ** 2 / self.c ** 2))

    def time_dilation(self, time_interval, velocity):
        """t' = t / sqrt(1 - v²/c²)"""
        if velocity >= self.c:
            raise ValueError("Velocity cannot reach or exceed speed of light.")
        return time_interval / math.sqrt(1 - (velocity ** 2 / self.c ** 2))

    def length_contraction(self, proper_length, velocity):
        """L = L₀ * sqrt(1 - v²/c²)"""
        if velocity >= self.c:
            raise ValueError("Velocity cannot reach or exceed speed of light.")
        return proper_length * math.sqrt(1 - (velocity ** 2 / self.c ** 2))

    def relativistic_momentum(self, mass, velocity):
        """p = γ * m * v"""
        gamma = 1 / math.sqrt(1 - (velocity ** 2 / self.c ** 2))
        return gamma * mass * velocity

    def lorentz_factor(self, velocity):
        """γ = 1 / sqrt(1 - v²/c²)"""
        if velocity >= self.c:
            raise ValueError("Velocity cannot reach or exceed speed of light.")
        return 1 / math.sqrt(1 - (velocity ** 2 / self.c ** 2))

# -----------------------------------------------------------
# ⚛️ QUANTUM PHYSICS
# -----------------------------------------------------------
    def photon_energy(self, frequency):
        """E = h * f"""
        return self.h * frequency

    def photon_energy_wavelength(self, wavelength):
        """E = h * c / λ"""
        return (self.h * self.c) / wavelength

    def de_broglie_wavelength(self, mass, velocity):
        """λ = h / (m * v)"""
        return self.h / (mass * velocity)

    def heisenberg_uncertainty(self, delta_x=None, delta_p=None):
        """
        Δx * Δp ≥ ħ / 2
        Provide one to calculate the other.
        """
        if delta_x is None and delta_p is None:
            raise ValueError("Provide at least one value (Δx or Δp).")
        if delta_x is not None:
            return self.h_bar / (2 * delta_x)
        elif delta_p is not None:
            return self.h_bar / (2 * delta_p)

    def energy_level_hydrogen(self, n):
        """Eₙ = -13.6 eV / n²"""
        return -13.6 / (n ** 2)

    def particle_energy(self, mass, velocity):
        """E = (γ - 1)mc²"""
        if velocity >= self.c:
            raise ValueError("Velocity cannot reach or exceed speed of light.")
        gamma = 1 / math.sqrt(1 - (velocity ** 2 / self.c ** 2))
        return (gamma - 1) * mass * (self.c ** 2)

# -----------------------------------------------------------
# ⚡ ELECTRICITY & MAGNETISM
# -----------------------------------------------------------
    def ohms_law(self, voltage=None, current=None, resistance=None):
        if voltage is None:
            return current * resistance
        elif current is None:
            return voltage / resistance
        elif resistance is None:
            return voltage / current
        else:
            raise ValueError("Provide only two parameters to find the third.")

    def electric_power(self, voltage, current):
        return voltage * current

    def charge(self, current, time):
        return current * time

    def coulomb_force(self, q1, q2, r):
        return (8.99e9 * q1 * q2) / (r ** 2)

# -----------------------------------------------------------
# 🌊 WAVES & LIGHT
# -----------------------------------------------------------
    def wave_speed(self, frequency, wavelength):
        return frequency * wavelength

    def frequency(self, wave_speed, wavelength):
        return wave_speed / wavelength

    def period(self, frequency):
        return 1 / frequency

# -----------------------------------------------------------
# 🌍 GRAVITATION
# -----------------------------------------------------------
    def gravitational_force(self, m1, m2, r):
        return self.G * m1 * m2 / (r ** 2)

# -----------------------------------------------------------
# 🔄 CONVERSIONS
# -----------------------------------------------------------
    def joule_to_electronvolt(self, joules):
        return joules / self.e

    def electronvolt_to_joule(self, ev):
        return ev * self.e

    def joule_to_calorie(self, joules):
        return joules / 4.184

    def calorie_to_joule(self, calories):
        return calories * 4.184

#__________________________________Chemistry_____________________________________
class Chemistry:
    def __init__(self):
        print("⚗️ Chemistry module initialized — includes classical and quantum functions.")

    # -----------------------------------------------------------
    # 🌍 CONSTANTS
    # -----------------------------------------------------------
    R = 8.314           # Gas constant (J/mol·K)
    NA = 6.022e23       # Avogadro's number (1/mol)
    h = 6.626e-34       # Planck constant (J·s)
    c = 3.0e8           # Speed of light (m/s)
    e = 1.602e-19       # Charge of electron (C)
    me = 9.109e-31      # Mass of electron (kg)
    kB = 1.38e-23       # Boltzmann constant (J/K)
    _elements = [
    {
        "number": 1,
        "symbol": "H",
        "name": "Hydrogen",
        "mass": 1.008,
        "group": 1,
        "period": 1,
        "type": "Nonmetal"
    },
    {
        "number": 2,
        "symbol": "He",
        "name": "Helium",
        "mass": 4.0026,
        "group": 18,
        "period": 1,
        "type": "Noble gas"
    },
    {
        "number": 3,
        "symbol": "Li",
        "name": "Lithium",
        "mass": 6.94,
        "group": 1,
        "period": 2,
        "type": "Alkali metal"
    },
    {
        "number": 4,
        "symbol": "Be",
        "name": "Beryllium",
        "mass": 9.0122,
        "group": 2,
        "period": 2,
        "type": "Alkaline earth metal"
    },
    {
        "number": 5,
        "symbol": "B",
        "name": "Boron",
        "mass": 10.81,
        "group": 13,
        "period": 2,
        "type": "Metalloid"
    },
    {
        "number": 6,
        "symbol": "C",
        "name": "Carbon",
        "mass": 12.011,
        "group": 14,
        "period": 2,
        "type": "Nonmetal"
    },
    {
        "number": 7,
        "symbol": "N",
        "name": "Nitrogen",
        "mass": 14.007,
        "group": 15,
        "period": 2,
        "type": "Nonmetal"
    },
    {
        "number": 8,
        "symbol": "O",
        "name": "Oxygen",
        "mass": 15.999,
        "group": 16,
        "period": 2,
        "type": "Nonmetal"
    },
    {
        "number": 9,
        "symbol": "F",
        "name": "Fluorine",
        "mass": 18.998,
        "group": 17,
        "period": 2,
        "type": "Halogen"
    },
    {
        "number": 10,
        "symbol": "Ne",
        "name": "Neon",
        "mass": 20.180,
        "group": 18,
        "period": 2,
        "type": "Noble gas"
    },
    {
        "number": 11,
        "symbol": "Na",
        "name": "Sodium",
        "mass": 22.990,
        "group": 1,
        "period": 3,
        "type": "Alkali metal"
    },
    {
        "number": 12,
        "symbol": "Mg",
        "name": "Magnesium",
        "mass": 24.305,
        "group": 2,
        "period": 3,
        "type": "Alkaline earth metal"
    },
    {
        "number": 13,
        "symbol": "Al",
        "name": "Aluminum",
        "mass": 26.982,
        "group": 13,
        "period": 3,
        "type": "Post-transition metal"
    },
    {
        "number": 14,
        "symbol": "Si",
        "name": "Silicon",
        "mass": 28.085,
        "group": 14,
        "period": 3,
        "type": "Metalloid"
    },
    {
        "number": 15,
        "symbol": "P",
        "name": "Phosphorus",
        "mass": 30.974,
        "group": 15,
        "period": 3,
        "type": "Nonmetal"
    },
    {
        "number": 16,
        "symbol": "S",
        "name": "Sulfur",
        "mass": 32.06,
        "group": 16,
        "period": 3,
        "type": "Nonmetal"
    },
    {
        "number": 17,
        "symbol": "Cl",
        "name": "Chlorine",
        "mass": 35.45,
        "group": 17,
        "period": 3,
        "type": "Halogen"
    },
    {
        "number": 18,
        "symbol": "Ar",
        "name": "Argon",
        "mass": 39.948,
        "group": 18,
        "period": 3,
        "type": "Noble gas"
    },
    {
        "number": 19,
        "symbol": "K",
        "name": "Potassium",
        "mass": 39.098,
        "group": 1,
        "period": 4,
        "type": "Alkali metal"
    },
    {
        "number": 20,
        "symbol": "Ca",
        "name": "Calcium",
        "mass": 40.078,
        "group": 2,
        "period": 4,
        "type": "Alkaline earth metal"
    },
    {
        "number": 21,
        "symbol": "Sc",
        "name": "Scandium",
        "mass": 44.956,
        "group": 3,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 22,
        "symbol": "Ti",
        "name": "Titanium",
        "mass": 47.867,
        "group": 4,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 23,
        "symbol": "V",
        "name": "Vanadium",
        "mass": 50.942,
        "group": 5,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 24,
        "symbol": "Cr",
        "name": "Chromium",
        "mass": 51.996,
        "group": 6,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 25,
        "symbol": "Mn",
        "name": "Manganese",
        "mass": 54.938,
        "group": 7,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 26,
        "symbol": "Fe",
        "name": "Iron",
        "mass": 55.845,
        "group": 8,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 27,
        "symbol": "Co",
        "name": "Cobalt",
        "mass": 58.933,
        "group": 9,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 28,
        "symbol": "Ni",
        "name": "Nickel",
        "mass": 58.693,
        "group": 10,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 29,
        "symbol": "Cu",
        "name": "Copper",
        "mass": 63.546,
        "group": 11,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 30,
        "symbol": "Zn",
        "name": "Zinc",
        "mass": 65.38,
        "group": 12,
        "period": 4,
        "type": "Transition metal"
    },
    {
        "number": 31,
        "symbol": "Ga",
        "name": "Gallium",
        "mass": 69.723,
        "group": 13,
        "period": 4,
        "type": "Post-transition metal"
    },
    {
        "number": 32,
        "symbol": "Ge",
        "name": "Germanium",
        "mass": 72.630,
        "group": 14,
        "period": 4,
        "type": "Metalloid"
    },
    {
        "number": 33,
        "symbol": "As",
        "name": "Arsenic",
        "mass": 74.922,
        "group": 15,
        "period": 4,
        "type": "Metalloid"
    },
    {
        "number": 34,
        "symbol": "Se",
        "name": "Selenium",
        "mass": 78.971,
        "group": 16,
        "period": 4,
        "type": "Nonmetal"
    },
    {
        "number": 35,
        "symbol": "Br",
        "name": "Bromine",
        "mass": 79.904,
        "group": 17,
        "period": 4,
        "type": "Halogen"
    },
    {
        "number": 36,
        "symbol": "Kr",
        "name": "Krypton",
        "mass": 83.798,
        "group": 18,
        "period": 4,
        "type": "Noble gas"
    },
    {
        "number": 37,
        "symbol": "Rb",
        "name": "Rubidium",
        "mass": 85.468,
        "group": 1,
        "period": 5,
        "type": "Alkali metal"
    },
    {
        "number": 38,
        "symbol": "Sr",
        "name": "Strontium",
        "mass": 87.62,
        "group": 2,
        "period": 5,
        "type": "Alkaline earth metal"
    },
    {
        "number": 39,
        "symbol": "Y",
        "name": "Yttrium",
        "mass": 88.906,
        "group": 3,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 40,
        "symbol": "Zr",
        "name": "Zirconium",
        "mass": 91.224,
        "group": 4,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 41,
        "symbol": "Nb",
        "name": "Niobium",
        "mass": 92.906,
        "group": 5,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 42,
        "symbol": "Mo",
        "name": "Molybdenum",
        "mass": 95.95,
        "group": 6,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 43,
        "symbol": "Tc",
        "name": "Technetium",
        "mass": 98,
        "group": 7,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 44,
        "symbol": "Ru",
        "name": "Ruthenium",
        "mass": 101.07,
        "group": 8,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 45,
        "symbol": "Rh",
        "name": "Rhodium",
        "mass": 102.91,
        "group": 9,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 46,
        "symbol": "Pd",
        "name": "Palladium",
        "mass": 106.42,
        "group": 10,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 47,
        "symbol": "Ag",
        "name": "Silver",
        "mass": 107.87,
        "group": 11,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 48,
        "symbol": "Cd",
        "name": "Cadmium",
        "mass": 112.41,
        "group": 12,
        "period": 5,
        "type": "Transition metal"
    },
    {
        "number": 49,
        "symbol": "In",
        "name": "Indium",
        "mass": 114.82,
        "group": 13,
        "period": 5,
        "type": "Post-transition metal"
    },
    {
        "number": 50,
        "symbol": "Sn",
        "name": "Tin",
        "mass": 118.71,
        "group": 14,
        "period": 5,
        "type": "Post-transition metal"
    },
    {
        "number": 51,
        "symbol": "Sb",
        "name": "Antimony",
        "mass": 121.76,
        "group": 15,
        "period": 5,
        "type": "Metalloid"
    },
    {
        "number": 52,
        "symbol": "Te",
        "name": "Tellurium",
        "mass": 127.60,
        "group": 16,
        "period": 5,
        "type": "Metalloid"
    },
    {
        "number": 53,
        "symbol": "I",
        "name": "Iodine",
        "mass": 126.90,
        "group": 17,
        "period": 5,
        "type": "Halogen"
    },
    {
        "number": 54,
        "symbol": "Xe",
        "name": "Xenon",
        "mass": 131.29,
        "group": 18,
        "period": 5,
        "type": "Noble gas"
    },
    {
        "number": 55,
        "symbol": "Cs",
        "name": "Cesium",
        "mass": 132.91,
        "group": 1,
        "period": 6,
        "type": "Alkali metal"
    },
    {
        "number": 56,
        "symbol": "Ba",
        "name": "Barium",
        "mass": 137.33,
        "group": 2,
        "period": 6,
        "type": "Alkaline earth metal"
    },
    {
        "number": 57,
        "symbol": "La",
        "name": "Lanthanum",
        "mass": 138.91,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 58,
        "symbol": "Ce",
        "name": "Cerium",
        "mass": 140.12,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 59,
        "symbol": "Pr",
        "name": "Praseodymium",
        "mass": 140.91,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 60,
        "symbol": "Nd",
        "name": "Neodymium",
        "mass": 144.24,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 61,
        "symbol": "Pm",
        "name": "Promethium",
        "mass": 145,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 62,
        "symbol": "Sm",
        "name": "Samarium",
        "mass": 150.36,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 63,
        "symbol": "Eu",
        "name": "Europium",
        "mass": 151.96,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 64,
        "symbol": "Gd",
        "name": "Gadolinium",
        "mass": 157.25,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 65,
        "symbol": "Tb",
        "name": "Terbium",
        "mass": 158.93,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 66,
        "symbol": "Dy",
        "name": "Dysprosium",
        "mass": 162.50,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 67,
        "symbol": "Ho",
        "name": "Holmium",
        "mass": 164.93,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 68,
        "symbol": "Er",
        "name": "Erbium",
        "mass": 167.26,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 69,
        "symbol": "Tm",
        "name": "Thulium",
        "mass": 168.93,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 70,
        "symbol": "Yb",
        "name": "Ytterbium",
        "mass": 173.05,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 71,
        "symbol": "Lu",
        "name": "Lutetium",
        "mass": 174.97,
        "group": 3,
        "period": 6,
        "type": "Lanthanide"
    },
    {
        "number": 72,
        "symbol": "Hf",
        "name": "Hafnium",
        "mass": 178.49,
        "group": 4,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 73,
        "symbol": "Ta",
        "name": "Tantalum",
        "mass": 180.95,
        "group": 5,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 74,
        "symbol": "W",
        "name": "Tungsten",
        "mass": 183.84,
        "group": 6,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 75,
        "symbol": "Re",
        "name": "Rhenium",
        "mass": 186.21,
        "group": 7,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 76,
        "symbol": "Os",
        "name": "Osmium",
        "mass": 190.23,
        "group": 8,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 77,
        "symbol": "Ir",
        "name": "Iridium",
        "mass": 192.22,
        "group": 9,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 78,
        "symbol": "Pt",
        "name": "Platinum",
        "mass": 195.08,
        "group": 10,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 79,
        "symbol": "Au",
        "name": "Gold",
        "mass": 196.97,
        "group": 11,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 80,
        "symbol": "Hg",
        "name": "Mercury",
        "mass": 200.59,
        "group": 12,
        "period": 6,
        "type": "Transition metal"
    },
    {
        "number": 81,
        "symbol": "Tl",
        "name": "Thallium",
        "mass": 204.38,
        "group": 13,
        "period": 6,
        "type": "Post-transition metal"
    },
    {
        "number": 82,
        "symbol": "Pb",
        "name": "Lead",
        "mass": 207.2,
        "group": 14,
        "period": 6,
        "type": "Post-transition metal"
    },
    {
        "number": 83,
        "symbol": "Bi",
        "name": "Bismuth",
        "mass": 208.98,
        "group": 15,
        "period": 6,
        "type": "Post-transition metal"
    },
    {
        "number": 84,
        "symbol": "Po",
        "name": "Polonium",
        "mass": 209,
        "group": 16,
        "period": 6,
        "type": "Metalloid"
    },
    {
        "number": 85,
        "symbol": "At",
        "name": "Astatine",
        "mass": 210,
        "group": 17,
        "period": 6,
        "type": "Halogen"
    },
    {
        "number": 86,
        "symbol": "Rn",
        "name": "Radon",
        "mass": 222,
        "group": 18,
        "period": 6,
        "type": "Noble gas"
    },
    {
        "number": 87,
        "symbol": "Fr",
        "name": "Francium",
        "mass": 223,
        "group": 1,
        "period": 7,
        "type": "Alkali metal"
    },
    {
        "number": 88,
        "symbol": "Ra",
        "name": "Radium",
        "mass": 226,
        "group": 2,
        "period": 7,
        "type": "Alkaline earth metal"
    },
    {
        "number": 89,
        "symbol": "Ac",
        "name": "Actinium",
        "mass": 227,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 90,
        "symbol": "Th",
        "name": "Thorium",
        "mass": 232.04,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 91,
        "symbol": "Pa",
        "name": "Protactinium",
        "mass": 231.04,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 92,
        "symbol": "U",
        "name": "Uranium",
        "mass": 238.03,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 93,
        "symbol": "Np",
        "name": "Neptunium",
        "mass": 237,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 94,
        "symbol": "Pu",
        "name": "Plutonium",
        "mass": 244,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 95,
        "symbol": "Am",
        "name": "Americium",
        "mass": 243,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 96,
        "symbol": "Cm",
        "name": "Curium",
        "mass": 247,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 97,
        "symbol": "Bk",
        "name": "Berkelium",
        "mass": 247,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 98,
        "symbol": "Cf",
        "name": "Californium",
        "mass": 251,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 99,
        "symbol": "Es",
        "name": "Einsteinium",
        "mass": 252,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 100,
        "symbol": "Fm",
        "name": "Fermium",
        "mass": 257,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 101,
        "symbol": "Md",
        "name": "Mendelevium",
        "mass": 258,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 102,
        "symbol": "No",
        "name": "Nobelium",
        "mass": 259,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 103,
        "symbol": "Lr",
        "name": "Lawrencium",
        "mass": 266,
        "group": 3,
        "period": 7,
        "type": "Actinide"
    },
    {
        "number": 104,
        "symbol": "Rf",
        "name": "Rutherfordium",
        "mass": 267,
        "group": 4,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 105,
        "symbol": "Db",
        "name": "Dubnium",
        "mass": 268,
        "group": 5,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 106,
        "symbol": "Sg",
        "name": "Seaborgium",
        "mass": 269,
        "group": 6,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 107,
        "symbol": "Bh",
        "name": "Bohrium",
        "mass": 270,
        "group": 7,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 108,
        "symbol": "Hs",
        "name": "Hassium",
        "mass": 277,
        "group": 8,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 109,
        "symbol": "Mt",
        "name": "Meitnerium",
        "mass": 278,
        "group": 9,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 110,
        "symbol": "Ds",
        "name": "Darmstadtium",
        "mass": 281,
        "group": 10,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 111,
        "symbol": "Rg",
        "name": "Roentgenium",
        "mass": 282,
        "group": 11,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 112,
        "symbol": "Cn",
        "name": "Copernicium",
        "mass": 285,
        "group": 12,
        "period": 7,
        "type": "Transition metal"
    },
    {
        "number": 113,
        "symbol": "Nh",
        "name": "Nihonium",
        "mass": 286,
        "group": 13,
        "period": 7,
        "type": "Post-transition metal"
    },
    {
        "number": 114,
        "symbol": "Fl",
        "name": "Flerovium",
        "mass": 289,
        "group": 14,
        "period": 7,
        "type": "Post-transition metal"
    },
    {
        "number": 115,
        "symbol": "Mc",
        "name": "Moscovium",
        "mass": 290,
        "group": 15,
        "period": 7,
        "type": "Post-transition metal"
    },
    {
        "number": 116,
        "symbol": "Lv",
        "name": "Livermorium",
        "mass": 293,
        "group": 16,
        "period": 7,
        "type": "Post-transition metal"
    },
    {
        "number": 117,
        "symbol": "Ts",
        "name": "Tennessine",
        "mass": 294,
        "group": 17,
        "period": 7,
        "type": "Halogen"
    },
    {
        "number": 118,
        "symbol": "Og",
        "name": "Oganesson",
        "mass": 294,
        "group": 18,
        "period": 7,
        "type": "Noble gas"
    }
]      # Blank element list (can be filled later)

    # -----------------------------------------------------------
    # 🧪 BASIC CHEMISTRY
    # -----------------------------------------------------------
    def molar_mass(self, mass, moles):
        """M = m / n"""
        return mass / moles if moles != 0 else float('inf')

    def moles_from_mass(self, mass, molar_mass):
        """n = m / M"""
        return mass / molar_mass if molar_mass != 0 else float('inf')

    def mass_from_moles(self, moles, molar_mass):
        """m = n * M"""
        return moles * molar_mass

    def concentration(self, moles, volume):
        """C = n / V (mol/L)"""
        return moles / volume if volume != 0 else float('inf')

    def percent_yield(self, actual, theoretical):
        """% yield = (actual / theoretical) * 100"""
        return (actual / theoretical) * 100 if theoretical != 0 else 0

    # -----------------------------------------------------------
    # 🌡️ GAS LAWS & THERMODYNAMICS
    # -----------------------------------------------------------
    def ideal_gas_pressure(self, n, V, T):
        """P = (nRT) / V"""
        return (n * self.R * T) / V if V != 0 else float('inf')

    def ideal_gas_volume(self, n, P, T):
        """V = (nRT) / P"""
        return (n * self.R * T) / P if P != 0 else float('inf')

    def ideal_gas_temperature(self, P, V, n):
        """T = (PV) / (nR)"""
        return (P * V) / (n * self.R) if n != 0 else float('inf')

    def heat_energy(self, mass, specific_heat, delta_T):
        """Q = m * c * ΔT"""
        return mass * specific_heat * delta_T

    def gibbs_free_energy(self, delta_H, delta_S, T):
        """ΔG = ΔH - TΔS"""
        return delta_H - T * delta_S

    def equilibrium_constant(self, delta_G, T):
        """K = e^(-ΔG / RT)"""
        return math.exp(-delta_G / (self.R * T))

    # -----------------------------------------------------------
    # ⚛️ ATOMIC & QUANTUM CHEMISTRY
    # -----------------------------------------------------------
    def energy_level_hydrogen(self, n):
        """Eₙ = -13.6 eV / n²"""
        return -13.6 / (n ** 2)

    def photon_energy(self, wavelength):
        """E = h * c / λ"""
        return (self.h * self.c) / wavelength

    def wavelength_from_energy(self, energy):
        """λ = h * c / E"""
        return (self.h * self.c) / energy

    def de_broglie_wavelength(self, mass, velocity):
        """λ = h / (m * v)"""
        return self.h / (mass * velocity)

    def uncertainty(self, delta_x):
        """Δp ≥ h / (4πΔx)"""
        return self.h / (4 * math.pi * delta_x)

    def particle_energy(self, frequency):
        """E = h * f"""
        return self.h * frequency

    def bohr_radius(self, n=1):
        """rₙ = (ε₀ * h² * n²) / (π * me * e²)"""
        ε0 = 8.854e-12
        return (ε0 * self.h**2 * n**2) / (math.pi * self.me * self.e**2)

    def ionization_energy(self, n1, n2):
        """ΔE (eV) = 13.6 * (1/n1² - 1/n2²)"""
        return 13.6 * ((1 / (n1 ** 2)) - (1 / (n2 ** 2)))

    # -----------------------------------------------------------
    # 🧮 STOICHIOMETRY
    # -----------------------------------------------------------
    def limiting_reactant(self, moles_A, ratio_A, moles_B, ratio_B):
        """Return limiting reactant based on stoichiometric ratios"""
        if (moles_A / ratio_A) < (moles_B / ratio_B):
            return "Reactant A is limiting"
        elif (moles_B / ratio_B) < (moles_A / ratio_A):
            return "Reactant B is limiting"
        else:
            return "Reactants are in perfect ratio"

    def percent_purity(self, pure_mass, impure_mass):
        """% Purity = (pure_mass / impure_mass) * 100"""
        return (pure_mass / impure_mass) * 100 if impure_mass != 0 else 0

    # -----------------------------------------------------------
    # 🔥 SPECTROSCOPY & ENERGY TRANSITIONS
    # -----------------------------------------------------------
    def frequency_from_wavelength(self, wavelength):
        """f = c / λ"""
        return self.c / wavelength

    def wavelength_from_frequency(self, frequency):
        """λ = c / f"""
        return self.c / frequency

    def photon_wavenumber(self, wavelength):
        """ν̃ = 1 / λ (in m⁻¹)"""
        return 1 / wavelength

    # -----------------------------------------------------------
    # 🔄 CONVERSIONS
    # -----------------------------------------------------------
    def joule_to_ev(self, joules):
        """Convert Joules to electronvolts"""
        return joules / self.e

    def ev_to_joule(self, ev):
        """Convert electronvolts to Joules"""
        return ev * self.e

    # -----------------------------------------------------------
    # 🧫 ELEMENT INFORMATION
    # -----------------------------------------------------------
    def atom_info(self, query):
        """
        Get atomic details by number, symbol, or name.
        Example:
            atom_info(8)
            atom_info('O')
            atom_info('Oxygen')
        """
        query_str = str(query).capitalize().strip()
        for element in self._elements:
            if (
                str(element["number"]) == str(query)
                or element["symbol"].lower() == str(query).lower()
                or element["name"].lower() == str(query).lower()
            ):
                return {
                    "Atomic Number": element["number"],
                    "Symbol": element["symbol"],
                    "Name": element["name"],
                    "Atomic Mass (u)": element["mass"],
                    "Group": element["group"],
                    "Period": element["period"],
                    "Type": element["type"]
                }
        return f"No element found for '{query}'"


#___________________________FTP_______________________________________
class FTP:
    def __init__(self, host="", user="", password="", port=21, timeout=30):
        self.host = host
        self.user = user
        self.password = password
        self.port = port
        self.timeout = timeout
        self.ftp = None

    # -----------------------------------------------------------
    # 🔌 CONNECTION HANDLING
    # -----------------------------------------------------------
    def connect(self):
        """Connect to the FTP server."""
        try:
            self.ftp = FTPClient()
            self.ftp.connect(self.host, self.port, timeout=self.timeout)
            self.ftp.login(self.user, self.password)
            print(f"✅ Connected to {self.host}")
        except Exception as e:
            print(f"❌ Connection failed: {e}")

    def disconnect(self):
        """Close FTP connection."""
        if self.ftp:
            self.ftp.quit()
            print("🔌 Disconnected from FTP server.")

    # -----------------------------------------------------------
    # 📁 DIRECTORY MANAGEMENT
    # -----------------------------------------------------------
    def current_dir(self):
        """Return current working directory."""
        if self.ftp:
            return self.ftp.pwd()

    def change_dir(self, path):
        """Change working directory."""
        if self.ftp:
            self.ftp.cwd(path)
            print(f"📂 Changed directory to: {path}")

    def make_dir(self, dirname):
        """Create a new directory."""
        if self.ftp:
            self.ftp.mkd(dirname)
            print(f"📁 Directory created: {dirname}")

    def remove_dir(self, dirname):
        """Remove a directory."""
        if self.ftp:
            self.ftp.rmd(dirname)
            print(f"🗑️ Directory removed: {dirname}")

    # -----------------------------------------------------------
    # 📜 FILE LISTING & INFO
    # -----------------------------------------------------------
    def list_files(self):
        """List files in current directory."""
        if self.ftp:
            files = self.ftp.nlst()
            for f in files:
                print(f"📄 {f}")
            return files

    def file_info(self, filename):
        """Get file size and modification time."""
        if self.ftp:
            size = self.ftp.size(filename)
            print(f"📏 Size of {filename}: {size} bytes")
            return size

    # -----------------------------------------------------------
    # ⬆️ UPLOAD & ⬇️ DOWNLOAD
    # -----------------------------------------------------------
    def upload_file(self, local_path, remote_path=None):
        """Upload a local file to the FTP server."""
        if not remote_path:
            remote_path = os.path.basename(local_path)
        if self.ftp and os.path.exists(local_path):
            with open(local_path, "rb") as f:
                self.ftp.storbinary(f"STOR {remote_path}", f)
            print(f"⬆️ Uploaded: {local_path} → {remote_path}")
        else:
            print("⚠️ Local file not found or not connected.")

    def download_file(self, remote_path, local_path=None):
        """Download a file from the FTP server."""
        if not local_path:
            local_path = os.path.basename(remote_path)
        if self.ftp:
            with open(local_path, "wb") as f:
                self.ftp.retrbinary(f"RETR {remote_path}", f.write)
            print(f"⬇️ Downloaded: {remote_path} → {local_path}")

    # -----------------------------------------------------------
    # ⚙️ FILE OPERATIONS
    # -----------------------------------------------------------
    def rename_file(self, old_name, new_name):
        """Rename file on the FTP server."""
        if self.ftp:
            self.ftp.rename(old_name, new_name)
            print(f"✏️ Renamed: {old_name} → {new_name}")

    def delete_file(self, filename):
        """Delete a file from the FTP server."""
        if self.ftp:
            self.ftp.delete(filename)
            print(f"🗑️ Deleted: {filename}")

    # -----------------------------------------------------------
    # 🧩 UTILITY
    # -----------------------------------------------------------
    def is_connected(self):
        """Check if the FTP connection is active."""
        try:
            self.ftp.voidcmd("NOOP")
            return True
        except:
            return False
#_____________________________Encoder______________________________
class Encoding_Decodin:
    def __init__(self):
        print("🔐 Encoding-Decoding module initialized — supports Base64, Hex, Binary, URL, and Unicode operations.")

    # ---------------------------------------------------------
    # 🧩 BASE64 ENCODING / DECODING
    # ---------------------------------------------------------
    def base64_encode(self, text: str) -> str:
        """Encodes text into Base64."""
        return base64.b64encode(text.encode('utf-8')).decode('utf-8')

    def base64_decode(self, encoded_text: str) -> str:
        """Decodes Base64 text into normal string."""
        return base64.b64decode(encoded_text.encode('utf-8')).decode('utf-8')

    # ---------------------------------------------------------
    # 🧱 HEX ENCODING / DECODING
    # ---------------------------------------------------------
    def hex_encode(self, text: str) -> str:
        """Encodes text into hexadecimal representation."""
        return text.encode('utf-8').hex()

    def hex_decode(self, hex_text: str) -> str:
        """Decodes hexadecimal text into normal string."""
        return bytes.fromhex(hex_text).decode('utf-8')

    # ---------------------------------------------------------
    # 💾 BINARY ENCODING / DECODING
    # ---------------------------------------------------------
    def binary_encode(self, text: str) -> str:
        """Encodes text into binary representation."""
        return ' '.join(format(ord(char), '08b') for char in text)

    def binary_decode(self, binary_text: str) -> str:
        """Decodes binary string into normal text."""
        return ''.join(chr(int(b, 2)) for b in binary_text.split())

    # ---------------------------------------------------------
    # 🌐 URL ENCODING / DECODING
    # ---------------------------------------------------------
    def url_encode(self, text: str) -> str:
        """Encodes URL or text safely for web transmission."""
        return urllib.parse.quote(text)

    def url_decode(self, encoded_text: str) -> str:
        """Decodes URL-encoded text back to original."""
        return urllib.parse.unquote(encoded_text)

    # ---------------------------------------------------------
    # 🔡 UTF ENCODING / DECODING
    # ---------------------------------------------------------
    def utf_encode(self, text: str, encoding='utf-8') -> bytes:
        """Encodes text using given UTF format (utf-8, utf-16, utf-32)."""
        return text.encode(encoding)

    def utf_decode(self, data: bytes, encoding='utf-8') -> str:
        """Decodes bytes using given UTF format."""
        return data.decode(encoding)

    # ---------------------------------------------------------
    # 📁 FILE ENCODING / DECODING
    # ---------------------------------------------------------
    def file_to_base64(self, file_path: str) -> str:
        """Encodes a file into Base64 string."""
        with open(file_path, 'rb') as f:
            return base64.b64encode(f.read()).decode('utf-8')

    def base64_to_file(self, encoded_data: str, output_path: str):
        """Decodes Base64 string and writes binary file."""
        with open(output_path, 'wb') as f:
            f.write(base64.b64decode(encoded_data))

    # ---------------------------------------------------------
    # 🧮 ASCII CONVERSION
    # ---------------------------------------------------------
    def text_to_ascii(self, text: str) -> list:
        """Converts text to list of ASCII codes."""
        return [ord(c) for c in text]

    def ascii_to_text(self, ascii_list: list) -> str:
        """Converts list of ASCII codes to text."""
        return ''.join(chr(i) for i in ascii_list)

    # ---------------------------------------------------------
    # 🔄 UNICODE CONVERSION
    # ---------------------------------------------------------
    def text_to_unicode(self, text: str) -> list:
        """Converts text to list of Unicode code points."""
        return [f"U+{ord(c):04X}" for c in text]

    def unicode_to_text(self, unicode_list: list) -> str:
        """Converts list of Unicode code points to text."""
        return ''.join(chr(int(u.replace('U+', ''), 16)) for u in unicode_list)

    # ---------------------------------------------------------
    # 🔏 ROT13 ENCODING
    # ---------------------------------------------------------
    def rot13_encode(self, text: str) -> str:
        """Encodes/decodes text using ROT13 cipher."""
        return codecs.encode(text, 'rot_13')

    # ---------------------------------------------------------
    # 🧬 MORSE CODE ENCODING / DECODING
    # ---------------------------------------------------------
    def morse_encode(self, text: str) -> str:
        """Encodes text into Morse code."""
        morse = {
            'A': '.-', 'B': '-...', 'C': '-.-.', 'D': '-..',
            'E': '.', 'F': '..-.', 'G': '--.', 'H': '....',
            'I': '..', 'J': '.---', 'K': '-.-', 'L': '.-..',
            'M': '--', 'N': '-.', 'O': '---', 'P': '.--.',
            'Q': '--.-', 'R': '.-.', 'S': '...', 'T': '-',
            'U': '..-', 'V': '...-', 'W': '.--', 'X': '-..-',
            'Y': '-.--', 'Z': '--..', '1': '.----', '2': '..---',
            '3': '...--', '4': '....-', '5': '.....', '6': '-....',
            '7': '--...', '8': '---..', '9': '----.', '0': '-----',
            ' ': '/'
        }
        return ' '.join(morse.get(c.upper(), '?') for c in text)

    def morse_decode(self, code: str) -> str:
        """Decodes Morse code into text."""
        morse = {
            '.-': 'A', '-...': 'B', '-.-.': 'C', '-..': 'D',
            '.': 'E', '..-.': 'F', '--.': 'G', '....': 'H',
            '..': 'I', '.---': 'J', '-.-': 'K', '.-..': 'L',
            '--': 'M', '-.': 'N', '---': 'O', '.--.': 'P',
            '--.-': 'Q', '.-.': 'R', '...': 'S', '-': 'T',
            '..-': 'U', '...-': 'V', '.--': 'W', '-..-': 'X',
            '-.--': 'Y', '--..': 'Z', '.----': '1', '..---': '2',
            '...--': '3', '....-': '4', '.....': '5', '-....': '6',
            '--...': '7', '---..': '8', '----.': '9', '-----': '0',
            '/': ' '
        }
        return ''.join(morse.get(c, '?') for c in code.split())

#_________________________________Number_System_Coverter_______________________
class Number_System:
    def __init__(self):
        print("🔢 Number System module initialized — supports Binary, Octal, Decimal, and Hexadecimal conversions.")

    # ---------------------------------------------------------
    # 🧮 BASIC CONVERSIONS
    # ---------------------------------------------------------
    def dec_to_bin(self, num: int) -> str:
        """Converts Decimal → Binary."""
        return bin(num)[2:]

    def dec_to_oct(self, num: int) -> str:
        """Converts Decimal → Octal."""
        return oct(num)[2:]

    def dec_to_hex(self, num: int) -> str:
        """Converts Decimal → Hexadecimal."""
        return hex(num)[2:].upper()

    def bin_to_dec(self, num: str) -> int:
        """Converts Binary → Decimal."""
        return int(num, 2)

    def oct_to_dec(self, num: str) -> int:
        """Converts Octal → Decimal."""
        return int(num, 8)

    def hex_to_dec(self, num: str) -> int:
        """Converts Hexadecimal → Decimal."""
        return int(num, 16)

    # ---------------------------------------------------------
    # 🔄 CROSS CONVERSIONS (BINARY, OCTAL, HEXADECIMAL)
    # ---------------------------------------------------------
    def bin_to_oct(self, num: str) -> str:
        """Converts Binary → Octal."""
        return oct(int(num, 2))[2:]

    def bin_to_hex(self, num: str) -> str:
        """Converts Binary → Hexadecimal."""
        return hex(int(num, 2))[2:].upper()

    def oct_to_bin(self, num: str) -> str:
        """Converts Octal → Binary."""
        return bin(int(num, 8))[2:]

    def oct_to_hex(self, num: str) -> str:
        """Converts Octal → Hexadecimal."""
        return hex(int(num, 8))[2:].upper()

    def hex_to_bin(self, num: str) -> str:
        """Converts Hexadecimal → Binary."""
        return bin(int(num, 16))[2:]

    def hex_to_oct(self, num: str) -> str:
        """Converts Hexadecimal → Octal."""
        return oct(int(num, 16))[2:]

    # ---------------------------------------------------------
    # 🧰 UNIVERSAL CONVERTER
    # ---------------------------------------------------------
    def convert(self, value: str, from_base: int, to_base: int) -> str:
        """
        Universal number converter.
        Supports bases between 2 and 36.
        Example: convert("1010", 2, 16) → 'A'
        """
        decimal_value = int(value, from_base)
        digits = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        if to_base == 10:
            return str(decimal_value)
        result = ""
        while decimal_value > 0:
            result = digits[decimal_value % to_base] + result
            decimal_value //= to_base
        return result or "0"

    # ---------------------------------------------------------
    # 🧪 VALIDATION
    # ---------------------------------------------------------
    def validate(self, value: str, base: int) -> bool:
        """Validates if a number string belongs to a base system."""
        valid_chars = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ"[:base]
        value = value.upper()
        return all(ch in valid_chars for ch in value)

    # ---------------------------------------------------------
    # 📦 BATCH CONVERSION
    # ---------------------------------------------------------
    def convert_list(self, values: list, from_base: int, to_base: int) -> list:
        """Converts a list of numbers between any two bases."""
        return [self.convert(v, from_base, to_base) for v in values]

    # ---------------------------------------------------------
    # 🧾 FORMAT DISPLAY
    # ---------------------------------------------------------
    def format_all(self, num: int) -> dict:
        """Returns a dictionary of all equivalent representations."""
        return {
            "Decimal": num,
            "Binary": self.dec_to_bin(num),
            "Octal": self.dec_to_oct(num),
            "Hexadecimal": self.dec_to_hex(num)
        }

    # ---------------------------------------------------------
    # 🧠 AUTO DETECT BASE
    # ---------------------------------------------------------
    def detect_base(self, value: str) -> str:
        """
        Detects number system based on prefix or pattern.
        """
        value = value.strip().lower()
        if value.startswith("0b"):
            return "Binary"
        elif value.startswith("0o"):
            return "Octal"
        elif value.startswith("0x"):
            return "Hexadecimal"
        elif all(c in "01" for c in value):
            return "Binary (No Prefix)"
        else:
            return "Decimal or Custom Base"

#_________________________Complex Numbers_______________________________

class ComplexNumber:
    def __init__(self, real=0.0, imag=0.0):
        self.real = real
        self.imag = imag

    def __str__(self):
        if self.imag >= 0:
            return f"{self.real} + {self.imag}i"
        else:
            return f"{self.real} - {abs(self.imag)}i"

    # =========================
    #   Basic Operations
    # =========================
    def add(self, other):
        return ComplexNumber(self.real + other.real, self.imag + other.imag)

    def subtract(self, other):
        return ComplexNumber(self.real - other.real, self.imag - other.imag)

    def multiply(self, other):
        r = self.real * other.real - self.imag * other.imag
        i = self.real * other.imag + self.imag * other.real
        return ComplexNumber(r, i)

    def divide(self, other):
        denominator = other.real**2 + other.imag**2
        if denominator == 0:
            raise ZeroDivisionError("Cannot divide by zero complex number")
        r = (self.real * other.real + self.imag * other.imag) / denominator
        i = (self.imag * other.real - self.real * other.imag) / denominator
        return ComplexNumber(r, i)

    # =========================
    #   Properties
    # =========================
    def conjugate(self):
        return ComplexNumber(self.real, -self.imag)

    def magnitude(self):
        return math.sqrt(self.real**2 + self.imag**2)

    def phase(self):
        return math.atan2(self.imag, self.real)

    def to_polar(self):
        return (self.magnitude(), self.phase())

    @staticmethod
    def from_polar(r, theta):
        return ComplexNumber(r * math.cos(theta), r * math.sin(theta))

    # =========================
    #   Utility Functions
    # =========================
    def equals(self, other):
        return self.real == other.real and self.imag == other.imag

    def is_real(self):
        return self.imag == 0

    def is_imaginary(self):
        return self.real == 0 and self.imag != 0

    # =========================
    #   Plotting Function
    # =========================
    def plot(self):
        plt.figure(figsize=(5, 5))
        plt.axhline(0, color='gray', linewidth=0.8)
        plt.axvline(0, color='gray', linewidth=0.8)
        plt.scatter(self.real, self.imag, color='red', s=100, label=str(self))
        plt.quiver(0, 0, self.real, self.imag, angles='xy', scale_units='xy', scale=1, color='blue')
        plt.title(f"Complex Number: {self}")
        plt.xlabel("Real Axis")
        plt.ylabel("Imaginary Axis")
        plt.grid(True)
        plt.legend()
        plt.xlim(-max(5, abs(self.real)+1), max(5, abs(self.real)+1))
        plt.ylim(-max(5, abs(self.imag)+1), max(5, abs(self.imag)+1))
        plt.show()



#_____________________________________________GUI_______________________________________________
#____________________________________NeuraPy GUI (Enhanced & Complete)________________________________________
import customtkinter as ctk
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog
import threading, time

# --------------------------------------------------------------------------
# Initialization
# --------------------------------------------------------------------------
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

# --------------------------------------------------------------------------
# COMPONENT WRAPPER
# --------------------------------------------------------------------------
class Component:
    """Wraps a widget and forwards .config() calls."""
    def __init__(self, widget):
        self.widget = widget

    def config(self, **kwargs):
        """Safely configure underlying widget."""
        for method in ("configure", "config"):
            if hasattr(self.widget, method):
                try:
                    getattr(self.widget, method)(**kwargs)
                    break
                except Exception:
                    continue


# --------------------------------------------------------------------------
# CONTAINER CLASS
# --------------------------------------------------------------------------
class Container:
    """A flexible container frame that supports both grid and pack layouts."""
    def __init__(self, master, **style):
        bg = style.get("bg", None)
        self.frame = ctk.CTkFrame(master, fg_color=bg)
        self.children = []

    def add(self, component, row=None, column=None, pack=False, padx=6, pady=6, sticky="nsew", **kwargs):
        """Add a component to this container."""
        w = component.widget if isinstance(component, Component) else component
        if pack:
            w.pack(padx=padx, pady=pady, **kwargs)
        else:
            row = len(self.children) if row is None else row
            column = 0 if column is None else column
            w.grid(row=row, column=column, padx=padx, pady=pady, sticky=sticky, **kwargs)
            self.frame.grid_rowconfigure(row, weight=1)
            self.frame.grid_columnconfigure(column, weight=1)
        self.children.append(w)
        return w

    def pack(self, **kwargs): self.frame.pack(**kwargs)
    def grid(self, **kwargs): self.frame.grid(**kwargs)
    def place(self, **kwargs): self.frame.place(**kwargs)


# --------------------------------------------------------------------------
# FLOOD GAUGE
# --------------------------------------------------------------------------
class FloodGauge(Component):
    """Custom vertical flood-style progress gauge."""
    def __init__(self, master, width=60, height=200, max_value=100, fg="#1E90FF", bg="#f5f5f5"):
        canvas = tk.Canvas(master, width=width, height=height, bg=bg, highlightthickness=1)
        super().__init__(canvas)
        self.width, self.height, self.max_value, self.value, self.fg, self.bg = width, height, max_value, 0, fg, bg
        self._draw()

    def _draw(self):
        c = self.widget
        c.delete("all")
        fill_height = (self.value / self.max_value) * (self.height - 10)
        y1 = self.height - 5 - fill_height
        c.create_rectangle(5, y1, self.width - 5, self.height - 5, fill=self.fg, outline="")
        c.create_rectangle(5, 5, self.width - 5, self.height - 5, outline="#999", width=1)
        c.create_text(self.width / 2, 12, text=f"{int(self.value / self.max_value * 100)}%", font=("Arial", 9, "bold"))

    def set(self, value):
        """Set the gauge value."""
        self.value = max(0, min(self.max_value, value))
        self._draw()

    def increment(self, step=1):
        """Increment gauge by step."""
        self.set(self.value + step)


# --------------------------------------------------------------------------
# MAIN GUI CLASS
# --------------------------------------------------------------------------
class GUI:
    """Simplified, powerful GUI builder for NeuraPy."""
    def __init__(self, title="NeuraPy GUI", size=(1000, 700), appearance="System", theme="blue"):
        ctk.set_appearance_mode(appearance)
        ctk.set_default_color_theme(theme)
        self.root = ctk.CTk()
        self.root.geometry(f"{size[0]}x{size[1]}")
        self.root.title(title)
        self.global_style = {}

        # Layout frames
        self.main = ctk.CTkFrame(self.root)
        self.main.pack(fill="both", expand=True)
        self.left = ctk.CTkFrame(self.main, width=200)
        self.center = ctk.CTkFrame(self.main)
        self.right = ctk.CTkFrame(self.main, width=240)
        self.left.pack(side="left", fill="y")
        self.center.pack(side="left", fill="both", expand=True)
        self.right.pack(side="right", fill="y")

        self.containers = {
            "left": Container(self.left),
            "center": Container(self.center),
            "right": Container(self.right),
            "root": Container(self.root)
        }

    # ----------------------------------------------------------------------
    # Basic Widgets
    # ----------------------------------------------------------------------
    def button(self, text="Button", command=None, **style):
        return Component(ctk.CTkButton(self.center, text=text, command=command, **{**self.global_style, **style}))

    def label(self, text="Label", **style):
        return Component(ctk.CTkLabel(self.center, text=text, **{**self.global_style, **style}))

    def textbox(self, placeholder="", multiline=False, **style):
        widget = ctk.CTkTextbox(self.center, **style) if multiline else \
                 ctk.CTkEntry(self.center, placeholder_text=placeholder, **style)
        return Component(widget)

    def slider(self, from_=0, to=100, value=0, orientation="horizontal", command=None, **style):
        if orientation == "horizontal":
            s = ctk.CTkSlider(self.center, from_=from_, to=to, **style)
            s.set(value)
            if command:
                s.configure(command=command)
            return Component(s)
        frame = ctk.CTkFrame(self.center)
        s = ttk.Scale(frame, from_=from_, to=to, orient="vertical")
        s.set(value)
        s.pack(fill="y", expand=True)
        return Component(s)

    def combo(self, values, default=None, **style):
        cb = ctk.CTkComboBox(self.center, values=values, **style)
        if default:
            cb.set(default)
        return Component(cb)

    def checkbox(self, text="Check", default=False, command=None, **style):
        var = tk.BooleanVar(value=default)
        chk = ctk.CTkCheckBox(self.center, text=text, variable=var, command=command, **style)
        chk.var = var
        return Component(chk)

    def radio(self, options, default=None, command=None, horizontal=False):
        var = tk.StringVar(value=default or options[0])
        frame = ctk.CTkFrame(self.center)
        for opt in options:
            rb = ctk.CTkRadioButton(frame, text=opt, variable=var, value=opt, command=command)
            rb.pack(side="left" if horizontal else "top", padx=4, pady=2, anchor="w")
        frame.var = var
        return Component(frame)

    def table(self, columns, height=8, **style):
        frame = ctk.CTkFrame(self.center, **style)
        tree = ttk.Treeview(frame, columns=columns, show="headings", height=height)
        vsb = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=vsb.set)
        for c in columns:
            tree.heading(c, text=c)
            tree.column(c, anchor="center")
        tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        frame.tree = tree
        return Component(frame)

    def flood_gauge(self, max_value=100, **style):
        frame = ctk.CTkFrame(self.center)
        gauge = FloodGauge(frame, max_value=max_value, **style)
        gauge.widget.pack()
        return gauge

    def loader(self, text="Loading...", duration=3):
        frame = ctk.CTkFrame(self.center)
        lbl = ctk.CTkLabel(frame, text=text)
        pb = ttk.Progressbar(frame, mode="indeterminate")
        lbl.pack(pady=4)
        pb.pack(pady=4)

        def spin():
            pb.start(10)
            time.sleep(duration)
            pb.stop()
        threading.Thread(target=spin, daemon=True).start()
        return Component(frame)

    # ----------------------------------------------------------------------
    # Dialogs and Message Boxes
    # ----------------------------------------------------------------------
    def alert(self, title="Alert", message="Something happened!"):
        messagebox.showinfo(title, message)

    def error(self, title="Error", message="An error occurred!"):
        messagebox.showerror(title, message)

    def ask_yes_no(self, title="Confirm", message="Are you sure?"):
        return messagebox.askyesno(title, message)

    def ask_input(self, title="Input", prompt="Enter value:"):
        return simpledialog.askstring(title, prompt)

    # ----------------------------------------------------------------------
    # File Dialogs
    # ----------------------------------------------------------------------
    def open_file(self, filetypes=(("All files", "*.*"),)):
        return filedialog.askopenfilename(filetypes=filetypes)

    def save_file(self, filetypes=(("All files", "*.*"),)):
        return filedialog.asksaveasfilename(filetypes=filetypes)

    def select_folder(self):
        return filedialog.askdirectory()

    # ----------------------------------------------------------------------
    # Utility Functions
    # ----------------------------------------------------------------------
    def toast(self, message, duration=2):
        """Temporary popup message."""
        popup = ctk.CTkToplevel(self.root)
        popup.overrideredirect(True)
        popup.geometry("+500+500")
        lbl = ctk.CTkLabel(popup, text=message, fg_color="#333", text_color="white", corner_radius=8)
        lbl.pack(padx=10, pady=10)
        self.root.after(duration * 1000, popup.destroy)

    def tooltip(self, widget, text):
        """Show tooltip when hovering over a widget."""
        tip = tk.Toplevel(widget)
        tip.withdraw()
        tip.overrideredirect(True)
        label = tk.Label(tip, text=text, bg="#333", fg="white", relief="solid", borderwidth=1)
        label.pack(ipadx=4)

        def enter(e):
            tip.geometry(f"+{e.x_root+10}+{e.y_root+10}")
            tip.deiconify()

        def leave(e):
            tip.withdraw()

        widget.bind("<Enter>", enter)
        widget.bind("<Leave>", leave)

    def set_style(self, **style):
        """Set global styling for all widgets."""
        self.global_style.update(style)

    # ----------------------------------------------------------------------
    # Window Controls
    # ----------------------------------------------------------------------
    def set_title(self, title):
        self.root.title(title)

    def resize(self, width, height):
        self.root.geometry(f"{width}x{height}")

    def fullscreen(self, enable=True):
        self.root.attributes("-fullscreen", enable)

    def close(self):
        self.root.destroy()

    # ----------------------------------------------------------------------
    # Event Binding System
    # ----------------------------------------------------------------------
    def bind_key(self, key, func):
        """Bind a keyboard key or combination globally to the window."""
        self.root.bind(key, func)

    def bind_mouse(self, event, func):
        """Bind a mouse event globally (e.g., '<Button-1>', '<Motion>')."""
        self.root.bind(event, func)

    def bind_to_widget(self, widget, event, func):
        """Bind an event to a specific widget."""
        target = widget.widget if isinstance(widget, Component) else widget
        target.bind(event, func)

    # ----------------------------------------------------------------------
    # Access & Run
    # ----------------------------------------------------------------------
    def container(self, name):
        return self.containers.get(name)

    def show(self):
        self.root.mainloop()

#_____________________________________Assembly Languae____________________________________
#______________________Full assembly language support____________________


from typing import Tuple, List, Dict, Optional, Union,Any
import mmap
import ctypes
import os
import platform
import sys
import tempfile
import struct
import logging

# --- Optional libs (robust imports) ---
KEYSTONE_AVAILABLE = False
CAPSTONE_AVAILABLE = False
UC_AVAILABLE = False

Ks = None
KsError = None
KS_ARCH_X86 = KS_ARCH_ARM = KS_ARCH_ARM64 = KS_ARCH_MIPS = KS_ARCH_PPC = KS_ARCH_SPARC = KS_ARCH_SYSTEMZ = None
KS_ARCH_HEXAGON = None
KS_MODE_16 = KS_MODE_32 = KS_MODE_64 = KS_MODE_ARM = KS_MODE_THUMB = None
KS_OPT_SYNTAX_INTEL = KS_OPT_SYNTAX_ATT = None

try:
    import keystone
    # bring commonly-used names, but tolerate missing attrs
    Ks = getattr(keystone, "Ks", None)
    KsError = getattr(keystone, "KsError", Exception)
    KS_ARCH_X86 = getattr(keystone, "KS_ARCH_X86", None)
    KS_ARCH_ARM = getattr(keystone, "KS_ARCH_ARM", None)
    KS_ARCH_ARM64 = getattr(keystone, "KS_ARCH_ARM64", None)
    KS_ARCH_MIPS = getattr(keystone, "KS_ARCH_MIPS", None)
    KS_ARCH_PPC = getattr(keystone, "KS_ARCH_PPC", None)
    KS_ARCH_SPARC = getattr(keystone, "KS_ARCH_SPARC", None)
    KS_ARCH_SYSTEMZ = getattr(keystone, "KS_ARCH_SYSTEMZ", None)
    KS_ARCH_HEXAGON = getattr(keystone, "KS_ARCH_HEXAGON", None)
    KS_MODE_16 = getattr(keystone, "KS_MODE_16", None)
    KS_MODE_32 = getattr(keystone, "KS_MODE_32", None)
    KS_MODE_64 = getattr(keystone, "KS_MODE_64", None)
    KS_MODE_ARM = getattr(keystone, "KS_MODE_ARM", None)
    KS_MODE_THUMB = getattr(keystone, "KS_MODE_THUMB", None)
    KS_OPT_SYNTAX_INTEL = getattr(keystone, "KS_OPT_SYNTAX_INTEL", None)
    KS_OPT_SYNTAX_ATT = getattr(keystone, "KS_OPT_SYNTAX_ATT", None)
    if Ks is not None:
        KEYSTONE_AVAILABLE = True
except Exception:
    # leave variables as None
    pass

try:
    import capstone
    Cs = getattr(capstone, "Cs", None)
    CS_ARCH_X86 = getattr(capstone, "CS_ARCH_X86", None)
    CS_ARCH_ARM = getattr(capstone, "CS_ARCH_ARM", None)
    CS_ARCH_ARM64 = getattr(capstone, "CS_ARCH_ARM64", None)
    CS_ARCH_MIPS = getattr(capstone, "CS_ARCH_MIPS", None)
    CS_ARCH_PPC = getattr(capstone, "CS_ARCH_PPC", None)
    CS_MODE_16 = getattr(capstone, "CS_MODE_16", None)
    CS_MODE_32 = getattr(capstone, "CS_MODE_32", None)
    CS_MODE_64 = getattr(capstone, "CS_MODE_64", None)
    CS_MODE_THUMB = getattr(capstone, "CS_MODE_THUMB", None)
    if Cs is not None:
        CAPSTONE_AVAILABLE = True
except Exception:
    Cs = None

try:
    import unicorn
    Uc = getattr(unicorn, "Uc", None)
    UC_ARCH_X86 = getattr(unicorn, "UC_ARCH_X86", None)
    UC_ARCH_ARM = getattr(unicorn, "UC_ARCH_ARM", None)
    UC_ARCH_ARM64 = getattr(unicorn, "UC_ARCH_ARM64", None)
    UC_ARCH_MIPS = getattr(unicorn, "UC_ARCH_MIPS", None)
    UC_MODE_16 = getattr(unicorn, "UC_MODE_16", None)
    UC_MODE_32 = getattr(unicorn, "UC_MODE_32", None)
    UC_MODE_64 = getattr(unicorn, "UC_MODE_64", None)
    UC_MODE_ARM = getattr(unicorn, "UC_MODE_ARM", None)
    UC_MODE_THUMB = getattr(unicorn, "UC_MODE_THUMB", None)
    # const modules
    try:
        import unicorn.x86_const as uc_x86_const
    except Exception:
        uc_x86_const = None
    try:
        import unicorn.arm_const as uc_arm_const
    except Exception:
        uc_arm_const = None
    try:
        import unicorn.arm64_const as uc_arm64_const
    except Exception:
        uc_arm64_const = None
    if Uc is not None:
        UC_AVAILABLE = True
except Exception:
    Uc = None
    uc_x86_const = None
    uc_arm_const = None
    uc_arm64_const = None

# For type clarity
BytesLike = Union[bytes, bytearray, str]  # str when hex string is allowed

# Setup logging
logger = logging.getLogger("Assembly")
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)


class AssemblyError(Exception):
    pass


class Assembly:
    """
    Assembly helper using Keystone (assembler), Capstone (disassembler),
    Unicorn (emulator), and native mmap/ctypes for direct execution.
    """

    def __init__(self, arch: str = "x86", mode: Union[int, str] = 64, syntax: str = "intel"):
        self.arch = arch.lower()
        self.mode = mode
        self.syntax = syntax.lower() if isinstance(syntax, str) else "intel"
        self.ks = None
        self._init_keystone()
        # emulator (Unicorn) instance placeholder
        self.uc: Optional[Any] = None
        self.emu_mapped: Dict[int, Union[int, tuple]] = {}  # addr -> size (UC) or (tmpfile, size) simulated
        self.emu_hooks = []
        logger.info(f"Assembly initialized ({self.arch}, {self.mode}, {self.syntax})")

    # ---------------- Keyston initialization ----------------
    def _init_keystone(self):
        if not KEYSTONE_AVAILABLE or Ks is None:
            self.ks = None
            logger.warning("Keystone engine not available. Assembly will be disabled.")
            return

        arch_map = {
            "x86": KS_ARCH_X86,
            "arm": KS_ARCH_ARM,
            "arm64": KS_ARCH_ARM64,
            "mips": KS_ARCH_MIPS,
            "ppc": KS_ARCH_PPC,
            "sparc": KS_ARCH_SPARC,
            "systemz": KS_ARCH_SYSTEMZ,
            "hexagon": KS_ARCH_HEXAGON,
        }
        mode_map = {
            16: KS_MODE_16,
            32: KS_MODE_32,
            64: KS_MODE_64,
            "arm": KS_MODE_ARM,
            "thumb": KS_MODE_THUMB,
        }

        arch_val = arch_map.get(self.arch)
        if arch_val is None:
            raise AssemblyError(f"Keystone does not support architecture '{self.arch}' in this wrapper.")
        mode_val = mode_map.get(self.mode)
        if mode_val is None:
            mode_val = KS_MODE_64 if self.mode == 64 else KS_MODE_32

        try:
            self.ks = Ks(arch_val, mode_val)
            # try setting syntax via attribute or option
            if self.arch == "x86" and KS_OPT_SYNTAX_INTEL is not None and KS_OPT_SYNTAX_ATT is not None:
                try:
                    # prefer option API if available
                    if hasattr(self.ks, "option"):
                        desired = KS_OPT_SYNTAX_INTEL if self.syntax == "intel" else KS_OPT_SYNTAX_ATT
                        try:
                            # some ks versions require integer key Ks.OPT_SYNTAX or attribute 'OPT_SYNTAX'
                            opt_key = getattr(Ks, "OPT_SYNTAX", None)
                            if opt_key is not None:
                                self.ks.option(opt_key, desired)
                            else:
                                # try attribute assignment
                                self.ks.syntax = desired
                        except Exception:
                            # fallback to attribute
                            try:
                                self.ks.syntax = desired
                            except Exception:
                                pass
                except Exception:
                    pass
        except Exception as e:
            # map to AssemblyError
            raise AssemblyError(f"Keystone initialization failed: {e}")

    # ---------------- Assembly ----------------
    def assemble(self, code: str, addr: int = 0x1000, as_hex: bool = True) -> Tuple[Union[str, bytes], int]:
        """
        Assemble the provided assembly text to bytes or hex string.
        Returns (data, count) where data is bytes or hex string depending on as_hex.
        """
        if not self.ks:
            raise AssemblyError("Keystone not available. Install keystone-engine to assemble.")
        try:
            encoding, count = self.ks.asm(code, addr)
            data = bytes(encoding)
            return (data.hex(), count) if as_hex else (data, count)
        except Exception as e:
            raise AssemblyError(f"Assembly failed: {e}")

    def batch_assemble(self, blocks: List[Dict], base_addr: int = 0x1000, as_hex: bool = True) -> List[Dict]:
        """
        Assemble multiple blocks.
        blocks: List of dicts: [{'code': 'mov eax, 1', 'addr': 0x1000}, ...]
        Returns list with {'addr', 'count', 'len', 'data'(hex or bytes)}
        """
        results = []
        for b in blocks:
            code = b.get("code", "")
            addr = b.get("addr", base_addr)
            data, count = self.assemble(code, addr, as_hex=False)
            results.append({
                "addr": addr,
                "count": count,
                "len": len(data),
                "data": data.hex() if as_hex else data
            })
        return results

    # ---------------- Disassembly ----------------
    def disassemble(self, code_bytes: BytesLike, addr: int = 0x1000) -> List[Tuple[int, str, str]]:
        """
        Disassemble bytes -> list of (addr, mnemonic, op_str)
        Accepts bytes or hex string.
        """
        if isinstance(code_bytes, str):
            s = code_bytes.strip()
            if s == "":
                return []
            # guess hex if only hex chars
            if all(c in "0123456789abcdefABCDEF" for c in s):
                code_bytes = bytes.fromhex(s)
            else:
                code_bytes = s.encode("latin1")
        if not CAPSTONE_AVAILABLE or Cs is None:
            raise AssemblyError("Capstone not available. Install capstone to disassemble.")

        arch_map = {
            "x86": CS_ARCH_X86,
            "arm": CS_ARCH_ARM,
            "arm64": CS_ARCH_ARM64,
            "mips": CS_ARCH_MIPS,
            "ppc": CS_ARCH_PPC,
        }
        mode_map = {
            16: CS_MODE_16,
            32: CS_MODE_32,
            64: CS_MODE_64,
            "thumb": CS_MODE_THUMB,
        }

        cs_arch = arch_map.get(self.arch)
        if cs_arch is None:
            raise AssemblyError(f"Disassembly: unsupported arch {self.arch}")

        cs_mode = mode_map.get(self.mode, CS_MODE_64 if self.mode == 64 else CS_MODE_32)
        cs = Cs(cs_arch, cs_mode)
        out = []
        for i in cs.disasm(code_bytes, addr):
            out.append((i.address, i.mnemonic, i.op_str))
        return out

    # ---------------- Native execution (dangerous) ----------------
    def execute_native(self, code_bytes: BytesLike, as_hex: bool = True):
        """
        Execute raw machine code within the current process memory.
        WARNING: Extremely dangerous. Use only on trusted code and in controlled environment.
        Works on POSIX via mmap and on Windows via VirtualAlloc.
        """
        if isinstance(code_bytes, str) and as_hex:
            code_bytes = bytes.fromhex(code_bytes.strip())
        elif isinstance(code_bytes, str) and not as_hex:
            code_bytes = code_bytes.encode('latin1')

        size = len(code_bytes)
        if size == 0:
            raise AssemblyError("No code to execute.")

        if os.name == "nt":
            # Windows: use VirtualAlloc to get executable RWX memory
            PAGE_EXECUTE_READWRITE = 0x40
            MEM_COMMIT = 0x1000
            MEM_RESERVE = 0x2000
            kernel32 = ctypes.windll.kernel32
            kernel32.VirtualAlloc.restype = ctypes.c_void_p
            addr = kernel32.VirtualAlloc(None, ctypes.c_size_t(size), MEM_COMMIT | MEM_RESERVE, PAGE_EXECUTE_READWRITE)
            if not addr:
                raise AssemblyError("VirtualAlloc failed.")
            # copy bytes to addr
            ctypes.memmove(addr, code_bytes, size)
            functype = ctypes.CFUNCTYPE(ctypes.c_void_p)
            func = functype(addr)
            logger.warning("Executing native code on Windows — this may crash your process.")
            try:
                res = func()
                return res
            finally:
                kernel32.VirtualFree(ctypes.c_void_p(addr), 0, 0x8000)  # MEM_RELEASE
        else:
            # POSIX: use mmap with PROT_EXEC if available
            PROT_READ = getattr(mmap, "PROT_READ", 1)
            PROT_WRITE = getattr(mmap, "PROT_WRITE", 2)
            PROT_EXEC = getattr(mmap, "PROT_EXEC", 4)
            prot = PROT_READ | PROT_WRITE | PROT_EXEC
            anon_flag = getattr(mmap, "MAP_ANONYMOUS", getattr(mmap, "MAP_ANON", 0))
            flags = mmap.MAP_PRIVATE | anon_flag if anon_flag != 0 else mmap.MAP_PRIVATE
            # create mapping
            mm = mmap.mmap(-1, size, prot=prot, flags=flags)
            try:
                mm.write(code_bytes)
                mm.seek(0)
                address = ctypes.addressof(ctypes.c_char.from_buffer(mm))
                functype = ctypes.CFUNCTYPE(ctypes.c_void_p)
                func = functype(address)
                logger.warning("Executing native code — this may crash your process. Proceed only if you know what you are doing.")
                try:
                    res = func()
                    return res
                finally:
                    pass
            finally:
                mm.close()

    # ---------------- Memory management for emulator ----------------
    def memory_alloc(self, size: int, addr: Optional[int] = None) -> int:
        """
        Allocate space inside the emulator address space. Returns mapped address or handle.
        If Unicorn is not available, simulate using a temporary file-backed mapping and return a handle.
        """
        if UC_AVAILABLE and self.uc is not None:
            # choose an address if not provided
            if addr is None:
                addr = 0x1000000 + (len(self.emu_mapped) * 0x10000)
            page_size = 0x1000
            alloc_size = ((size + page_size - 1) // page_size) * page_size
            self.uc.mem_map(addr, alloc_size)
            self.emu_mapped[addr] = alloc_size
            logger.info(f"Unicorn mapped {alloc_size} bytes at 0x{addr:X}")
            return addr
        else:
            tmp = tempfile.TemporaryFile()
            tmp.truncate(size)
            handle = id(tmp)
            # store object to keep it alive + size
            self.emu_mapped[handle] = (tmp, size)
            logger.info(f"Simulated allocation: handle {handle}, size {size}")
            return handle

    def memory_free(self, addr):
        if UC_AVAILABLE and self.uc is not None and isinstance(addr, int) and addr in self.emu_mapped:
            size = self.emu_mapped.pop(addr)
            self.uc.mem_unmap(addr, size)
            logger.info(f"Unmapped emulator memory at 0x{addr:X} size {size}")
            return True
        else:
            entry = self.emu_mapped.pop(addr, None)
            if entry:
                if isinstance(entry, tuple):
                    entry[0].close()
                logger.info(f"Freed simulated allocation {addr}")
                return True
        return False

    def memory_write(self, addr: int, data: BytesLike, offset: int = 0):
        """
        Write data to mapped memory. For simulated mappings, address is the handle returned by memory_alloc.
        Supports optional offset for simulated mappings.
        """
        if isinstance(data, str):
            s = data.strip()
            if all(c in "0123456789abcdefABCDEF" for c in s) and s != "":
                data = bytes.fromhex(s)
            else:
                data = data.encode('latin1')

        if UC_AVAILABLE and self.uc is not None and isinstance(addr, int) and addr in self.emu_mapped:
            self.uc.mem_write(addr + offset, data)
            logger.info(f"Wrote {len(data)} bytes to 0x{addr+offset:X} in emulator")
            return True
        else:
            entry = self.emu_mapped.get(addr)
            if entry and isinstance(entry, tuple):
                tmp, size = entry
                if offset < 0 or offset + len(data) > size:
                    raise AssemblyError("Write exceeds simulated mapping bounds.")
                tmp.seek(offset)
                tmp.write(data)
                tmp.flush()
                logger.info(f"Wrote {len(data)} bytes to simulated handle {addr} at offset {offset}")
                return True
        raise AssemblyError("Emulator not running and no suitable simulated mapping found.")

    def memory_read(self, addr: int, size: int, offset: int = 0) -> bytes:
        """
        Read from mapped memory. For simulated mappings, address is the handle returned by memory_alloc.
        """
        if UC_AVAILABLE and self.uc is not None and isinstance(addr, int) and addr in self.emu_mapped:
            return self.uc.mem_read(addr + offset, size)
        else:
            entry = self.emu_mapped.get(addr)
            if entry and isinstance(entry, tuple):
                tmp, sz = entry
                if offset < 0 or offset + size > sz:
                    raise AssemblyError("Read exceeds simulated mapping bounds.")
                tmp.seek(offset)
                return tmp.read(size)
        raise AssemblyError("Emulator not running and no suitable simulated mapping found.")

    # ---------------- Emulator (Unicorn) ----------------
    def emu_init(self):
        if not UC_AVAILABLE or Uc is None:
            raise AssemblyError("Unicorn engine not available. Install 'unicorn' package to emulate.")
        arch_map = {
            "x86": UC_ARCH_X86,
            "arm": UC_ARCH_ARM,
            "arm64": UC_ARCH_ARM64,
            "mips": UC_ARCH_MIPS,
        }
        mode_map = {
            16: UC_MODE_16,
            32: UC_MODE_32,
            64: UC_MODE_64,
            "arm": UC_MODE_ARM,
            "thumb": UC_MODE_THUMB,
        }
        uc_arch = arch_map.get(self.arch)
        if uc_arch is None:
            raise AssemblyError(f"Emulation: unsupported architecture {self.arch}")
        uc_mode = mode_map.get(self.mode, UC_MODE_64 if self.mode == 64 else UC_MODE_32)
        self.uc = Uc(uc_arch, uc_mode)
        self.emu_mapped = {}
        self.emu_hooks = []
        logger.info("Unicorn emulator initialized")

    def emu_map_and_write(self, addr: int, code: BytesLike):
        if isinstance(code, str):
            s = code.strip()
            if all(c in "0123456789abcdefABCDEF" for c in s) and s != "":
                code = bytes.fromhex(s)
            else:
                code = code.encode('latin1')
        if self.uc is None:
            raise AssemblyError("Unicorn emulator not initialized. Call emu_init() first.")
        size = len(code)
        page_size = 0x1000
        alloc_size = ((size + page_size - 1) // page_size) * page_size
        self.uc.mem_map(addr, alloc_size)
        self.uc.mem_write(addr, code)
        self.emu_mapped[addr] = alloc_size
        logger.info(f"Mapped and wrote {size} bytes to 0x{addr:X}")

    def emu_run(self, start_addr: int, size: int, timeout: int = 0, regs_init: Optional[Dict[str, int]] = None):
        """
        Run code in emulator. Returns dict with final registers snapshot.
        regs_init: mapping of architecture register names -> values
        timeout: microseconds
        """
        if not UC_AVAILABLE or self.uc is None:
            raise AssemblyError("Unicorn not available or not initialized.")
        if regs_init:
            self._emu_set_regs_from_map(regs_init)
        try:
            end_addr = start_addr + size
            # call with keywords first (modern unicorn), otherwise try positional
            try:
                if timeout:
                    self.uc.emu_start(start_addr, end_addr, timeout=timeout, count=0)
                else:
                    self.uc.emu_start(start_addr, end_addr)
            except TypeError:
                # try positional fallback (timeout, count)
                if timeout:
                    self.uc.emu_start(start_addr, end_addr, timeout, 0)
                else:
                    self.uc.emu_start(start_addr, end_addr)
        except Exception as e:
            logger.warning(f"Emulation stopped with exception: {e}")
        finally:
            return self.regs_snapshot()

    # ---------------- Registers (emulator) ----------------
    def _get_reg_const(self, name: str) -> Optional[int]:
        """Map common register name strings to Unicorn constants depending on arch."""
        if not isinstance(name, str) or name == "":
            return None
        n = name.lower()
        if self.arch == "x86" and uc_x86_const is not None:
            reg_map = {
                "rax": getattr(uc_x86_const, "UC_X86_REG_RAX", None),
                "rbx": getattr(uc_x86_const, "UC_X86_REG_RBX", None),
                "rcx": getattr(uc_x86_const, "UC_X86_REG_RCX", None),
                "rdx": getattr(uc_x86_const, "UC_X86_REG_RDX", None),
                "rsp": getattr(uc_x86_const, "UC_X86_REG_RSP", None),
                "rbp": getattr(uc_x86_const, "UC_X86_REG_RBP", None),
                "rsi": getattr(uc_x86_const, "UC_X86_REG_RSI", None),
                "rdi": getattr(uc_x86_const, "UC_X86_REG_RDI", None),
                "rip": getattr(uc_x86_const, "UC_X86_REG_RIP", None),
                "eax": getattr(uc_x86_const, "UC_X86_REG_EAX", None),
                "ebx": getattr(uc_x86_const, "UC_X86_REG_EBX", None),
                "ecx": getattr(uc_x86_const, "UC_X86_REG_ECX", None),
                "edx": getattr(uc_x86_const, "UC_X86_REG_EDX", None),
                "esp": getattr(uc_x86_const, "UC_X86_REG_ESP", None),
                "ebp": getattr(uc_x86_const, "UC_X86_REG_EBP", None),
                "esi": getattr(uc_x86_const, "UC_X86_REG_ESI", None),
                "edi": getattr(uc_x86_const, "UC_X86_REG_EDI", None),
            }
            return reg_map.get(n)
        elif self.arch == "arm" and uc_arm_const is not None:
            # r0..r12
            if n.startswith("r") and len(n) > 1 and n[1:].isdigit():
                idx = int(n[1:])
                return getattr(uc_arm_const, f"UC_ARM_REG_R{idx}", None)
            map_simple = {
                "sp": getattr(uc_arm_const, "UC_ARM_REG_SP", None),
                "lr": getattr(uc_arm_const, "UC_ARM_REG_LR", None),
                "pc": getattr(uc_arm_const, "UC_ARM_REG_PC", None),
            }
            if n in map_simple:
                return map_simple[n]
            try:
                return getattr(uc_arm_const, n.upper(), None)
            except Exception:
                return None
        elif self.arch == "arm64" and uc_arm64_const is not None:
            if n.startswith("x") and len(n) > 1 and n[1:].isdigit():
                idx = int(n[1:])
                return getattr(uc_arm64_const, f"UC_ARM64_REG_X{idx}", None)
            if n == "sp":
                return getattr(uc_arm64_const, "UC_ARM64_REG_SP", None)
            if n == "pc":
                return getattr(uc_arm64_const, "UC_ARM64_REG_PC", None)
            return None
        else:
            return None

    def _emu_set_regs_from_map(self, regs: Dict[str, int]):
        """Set multiple registers in the running emulator."""
        for name, val in regs.items():
            rc = self._get_reg_const(name)
            if rc is None:
                logger.warning(f"Unknown register '{name}' for arch {self.arch}")
                continue
            self.uc.reg_write(rc, val)

    def regs_snapshot(self) -> Dict[str, int]:
        """Return a snapshot of commonly useful registers depending on arch."""
        snapshot = {}
        if not UC_AVAILABLE or self.uc is None:
            raise AssemblyError("Unicorn emulator not available/running.")
        if self.arch == "x86" and uc_x86_const is not None:
            regs = ["RAX", "RBX", "RCX", "RDX", "RSP", "RBP", "RSI", "RDI", "RIP"]
            for r in regs:
                const = getattr(uc_x86_const, f"UC_X86_REG_{r}", None)
                if const is not None:
                    snapshot[r.lower()] = self.uc.reg_read(const)
        elif self.arch == "arm" and uc_arm_const is not None:
            for i in range(13):
                const = getattr(uc_arm_const, f"UC_ARM_REG_R{i}", None)
                if const is not None:
                    snapshot[f"r{i}"] = self.uc.reg_read(const)
            if getattr(uc_arm_const, "UC_ARM_REG_SP", None):
                snapshot["sp"] = self.uc.reg_read(uc_arm_const.UC_ARM_REG_SP)
            if getattr(uc_arm_const, "UC_ARM_REG_LR", None):
                snapshot["lr"] = self.uc.reg_read(uc_arm_const.UC_ARM_REG_LR)
            if getattr(uc_arm_const, "UC_ARM_REG_PC", None):
                snapshot["pc"] = self.uc.reg_read(uc_arm_const.UC_ARM_REG_PC)
        elif self.arch == "arm64" and uc_arm64_const is not None:
            for i in range(31):
                const = getattr(uc_arm64_const, f"UC_ARM64_REG_X{i}", None)
                if const is not None:
                    snapshot[f"x{i}"] = self.uc.reg_read(const)
            if getattr(uc_arm64_const, "UC_ARM64_REG_SP", None):
                snapshot["sp"] = self.uc.reg_read(uc_arm64_const.UC_ARM64_REG_SP)
            if getattr(uc_arm64_const, "UC_ARM64_REG_PC", None):
                snapshot["pc"] = self.uc.reg_read(uc_arm64_const.UC_ARM64_REG_PC)
        else:
            logger.warning("regs_snapshot: limited register set for arch " + self.arch)
        return snapshot

    def reg_read(self, name: str) -> int:
        """Read a single register value using emulator."""
        if not UC_AVAILABLE or self.uc is None:
            raise AssemblyError("Unicorn emulator not available/running.")
        rc = self._get_reg_const(name)
        if rc is None:
            raise AssemblyError(f"Unknown register '{name}' for arch {self.arch}")
        return self.uc.reg_read(rc)

    def reg_write(self, name: str, value: int):
        """Write a single register value using emulator."""
        if not UC_AVAILABLE or self.uc is None:
            raise AssemblyError("Unicorn emulator not available/running.")
        rc = self._get_reg_const(name)
        if rc is None:
            raise AssemblyError(f"Unknown register '{name}' for arch {self.arch}")
        self.uc.reg_write(rc, value)
        return True

    # ---------------- Hardware IO (simulated where not allowed) ----------------
    def write_port(self, port: int, value: int):
        """Simulated port I/O. Real port I/O needs platform-specific drivers and privileges."""
        logger.info(f"Simulated write to port 0x{port:X} with value 0x{value:X}")
        return True

    def read_port(self, port: int) -> int:
        logger.info(f"Simulated read from port 0x{port:X} -> returning 0")
        return 0

    # ---------------- Utilities ----------------
    @staticmethod
    def hex_dump(byte_data: BytesLike, width: int = 16) -> str:
        if isinstance(byte_data, str):
            s = byte_data.strip()
            if s == "":
                data = b""
            elif all(c in "0123456789abcdefABCDEF" for c in s):
                data = bytes.fromhex(s)
            else:
                data = s.encode('latin1')
        else:
            data = bytes(byte_data)
        lines = []
        for i in range(0, len(data), width):
            chunk = data[i:i + width]
            hex_chunk = ' '.join(f"{b:02X}" for b in chunk)
            ascii_chunk = ''.join(chr(b) if 32 <= b <= 126 else '.' for b in chunk)
            lines.append(f"{i:08X}: {hex_chunk:<{width*3}} {ascii_chunk}")
        return "\n".join(lines)

    def save_binary(self, filename: str, data: BytesLike):
        if isinstance(data, str):
            s = data.strip()
            if s != "" and all(c in "0123456789abcdefABCDEF" for c in s):
                data = bytes.fromhex(s)
            else:
                data = data.encode('latin1')
        with open(filename, "wb") as f:
            f.write(data)
        logger.info(f"Wrote {len(data)} bytes to {filename}")
        return filename

    def load_binary(self, filename: str) -> bytes:
        with open(filename, "rb") as f:
            data = f.read()
        logger.info(f"Loaded {len(data)} bytes from {filename}")
        return data

    # ---------------- High-level helper ----------------
    def emulate_code(self, asm_code: str, start_addr: int = 0x1000, regs_init: Optional[Dict[str, int]] = None) -> Dict[str, int]:
        """
        Assemble and emulate a small code snippet and return register snapshot.
        """
        code_bytes, cnt = self.assemble(asm_code, addr=start_addr, as_hex=False)
        if isinstance(code_bytes, str):
            code_bytes = bytes.fromhex(code_bytes)
        if not UC_AVAILABLE:
            raise AssemblyError("Unicorn not available. Cannot emulate code.")
        if self.uc is None:
            self.emu_init()
        self.emu_map_and_write(start_addr, code_bytes)
        snapshot = self.emu_run(start_addr, len(code_bytes), timeout=0, regs_init=regs_init)
        return snapshot

    

#____________________________________Web Scrapper__________________________________
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import os
import json
import re

class Web_Scrapper:
    """
    🌐 Web_Scrapper: A universal and powerful web scraping class.
    
    Features:
    - Fetch HTML content from URLs
    - Extract tags, text, scripts, styles, tables, forms
    - Extract links, images, emails, and phone numbers
    - Download and save media
    - Handle headers, user agents, proxies, and timeouts
    - Save and load scraped data as JSON
    - Recursive link crawling
    """

    def __init__(self, user_agent: str = None, timeout: int = 10, proxies: dict = None):
        self.session = requests.Session()
        self.timeout = timeout
        self.headers = {
            "User-Agent": user_agent or 
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0 Safari/537.36"
        }
        self.proxies = proxies
        self.last_url = None
        self.soup = None
        self.content = None

    # ------------------------------------------------------------
    # 🔹 Fetching & Parsing
    # ------------------------------------------------------------
    def fetch(self, url: str, parse: bool = True):
        """Fetch HTML content from a URL. Optionally parse it."""
        try:
            response = self.session.get(url, headers=self.headers, timeout=self.timeout, proxies=self.proxies)
            response.raise_for_status()
            self.last_url = url
            self.content = response.text
            if parse:
                self.soup = BeautifulSoup(self.content, "lxml")
            return True
        except requests.RequestException as e:
            print(f"[❌] Error fetching URL: {e}")
            return False

    def get_html(self):
        """Return the raw HTML of the last fetched page."""
        return self.content

    def parse_html(self, html: str):
        """Parse raw HTML text."""
        self.content = html
        self.soup = BeautifulSoup(html, "lxml")

    # ------------------------------------------------------------
    # 🔹 Extractors
    # ------------------------------------------------------------
    def get_title(self):
        return self.soup.title.string if self.soup and self.soup.title else None

    def get_text(self, selector=None):
        if not self.soup: return None
        if selector:
            elements = self.soup.select(selector)
            return [el.get_text(strip=True) for el in elements]
        return self.soup.get_text(separator="\n", strip=True)

    def get_links(self):
        """Return all hyperlinks on the page."""
        if not self.soup: return []
        links = [urljoin(self.last_url, a["href"]) for a in self.soup.find_all("a", href=True)]
        return list(set(links))

    def get_images(self):
        """Return all image URLs on the page."""
        if not self.soup: return []
        imgs = [urljoin(self.last_url, img["src"]) for img in self.soup.find_all("img", src=True)]
        return list(set(imgs))

    def get_scripts(self):
        """Return all script sources."""
        if not self.soup: return []
        return [urljoin(self.last_url, s.get("src")) for s in self.soup.find_all("script", src=True)]

    def get_styles(self):
        """Return all CSS links."""
        if not self.soup: return []
        return [urljoin(self.last_url, l.get("href")) for l in self.soup.find_all("link", rel="stylesheet")]

    def get_tables(self):
        """Return all tables as list of lists."""
        if not self.soup: return []
        tables = []
        for table in self.soup.find_all("table"):
            rows = []
            for tr in table.find_all("tr"):
                row = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                rows.append(row)
            tables.append(rows)
        return tables

    def get_forms(self):
        """Return all form action URLs."""
        if not self.soup: return []
        return [urljoin(self.last_url, f.get("action")) for f in self.soup.find_all("form", action=True)]

    def get_meta_tags(self):
        """Return meta tag data as a dictionary."""
        if not self.soup: return {}
        metas = {}
        for tag in self.soup.find_all("meta"):
            name = tag.get("name") or tag.get("property")
            content = tag.get("content")
            if name and content:
                metas[name] = content
        return metas

    def get_emails(self):
        """Extract emails from the page."""
        if not self.content: return []
        return list(set(re.findall(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", self.content)))

    def get_phone_numbers(self):
        """Extract phone numbers (basic patterns)."""
        if not self.content: return []
        pattern = r"\+?\d[\d\-\s]{7,}\d"
        return list(set(re.findall(pattern, self.content)))

    def find(self, tag, **kwargs):
        if not self.soup: return None
        return self.soup.find(tag, **kwargs)

    def find_all(self, tag, **kwargs):
        if not self.soup: return []
        return self.soup.find_all(tag, **kwargs)

    # ------------------------------------------------------------
    # 🔹 Download Utilities
    # ------------------------------------------------------------
    def download_file(self, url: str, save_path: str = None):
        try:
            r = self.session.get(url, headers=self.headers, timeout=self.timeout, stream=True)
            r.raise_for_status()
            filename = save_path or os.path.basename(urlparse(url).path)
            with open(filename, "wb") as f:
                for chunk in r.iter_content(1024):
                    f.write(chunk)
            print(f"[✅] Downloaded: {filename}")
            return filename
        except Exception as e:
            print(f"[❌] Failed to download {url}: {e}")
            return None

    def download_all_images(self, folder="images"):
        if not os.path.exists(folder):
            os.makedirs(folder)
        for img_url in self.get_images():
            filename = os.path.join(folder, os.path.basename(urlparse(img_url).path))
            self.download_file(img_url, filename)

    # ------------------------------------------------------------
    # 🔹 Data Handling
    # ------------------------------------------------------------
    def save_json(self, data, filename="scraped_data.json"):
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        print(f"[💾] Data saved to {filename}")

    def load_json(self, filename="scraped_data.json"):
        if not os.path.exists(filename):
            print("[⚠️] File not found!")
            return None
        with open(filename, "r", encoding="utf-8") as f:
            return json.load(f)

    # ------------------------------------------------------------
    # 🔹 Utility
    # ------------------------------------------------------------
    def search(self, pattern: str):
        if not self.content: return []
        return re.findall(pattern, self.content, re.IGNORECASE)

    def clear_cache(self):
        self.content = None
        self.soup = None
        self.last_url = None

    def crawl_links(self, depth: int = 1, filter_fn=None):
        """
        Recursively crawl links up to a specified depth.
        Optionally filter URLs using filter_fn(url) -> bool
        """
        visited = set()
        to_visit = [(self.last_url, 0)]
        all_links = []

        while to_visit:
            url, d = to_visit.pop(0)
            if url in visited or d > depth:
                continue
            if self.fetch(url):
                links = self.get_links()
                if filter_fn:
                    links = [l for l in links if filter_fn(l)]
                all_links.extend(links)
                to_visit.extend([(l, d+1) for l in links])
                visited.add(url)
        return list(set(all_links))

    def __repr__(self):
        return f"<Web_Scrapper(url='{self.last_url}')>"


#_____________________________________Sensors_______________________________________
import psutil
import sounddevice as sd
import numpy as np
import screen_brightness_control as sbc
import platform
import time
import cv2

class PhysicalSensors:
    def __init__(self):
        self.system = platform.system()
        print(f"🛰️ PhysicalSensors initialized on {self.system}")

    # ---------------- SOUND SENSOR ----------------
    def sound_level(self, duration=1):
        """Measure average microphone sound level (RMS)."""
        try:
            data = sd.rec(int(duration * 44100), samplerate=44100, channels=1, dtype='float64')
            sd.wait()
            rms = np.sqrt(np.mean(np.square(data)))
            return round(rms * 1000, 3)
        except Exception as e:
            return f"Sound sensor not accessible: {e}"

    # ---------------- LIGHT SENSOR ----------------
    def brightness(self):
        """Get display brightness (cross-platform)."""
        try:
            brightness = sbc.get_brightness(display=0)
            return brightness[0] if brightness else "Unknown"
        except Exception as e:
            return f"Brightness sensor not accessible: {e}"

    # ---------------- TEMPERATURE SENSOR ----------------
    def cpu_temperature(self):
        """Get CPU temperature if available."""
        try:
            temps = psutil.sensors_temperatures()
            if not temps:
                return "No temperature sensors found"
            # Return first available CPU temp
            for name, entries in temps.items():
                for entry in entries:
                    if "cpu" in name.lower() or "core" in name.lower():
                        label = entry.label if entry.label else name
                        return f"{label}: {entry.current:.1f}°C"
            return "CPU temperature sensor not found"
        except Exception as e:
            return f"Temperature sensor error: {e}"

    # ---------------- BATTERY SENSOR ----------------
    def battery_status(self):
        """Get battery percentage and charging status."""
        try:
            battery = psutil.sensors_battery()
            if not battery:
                return "Battery not found"
            time_left = battery.secsleft if battery.secsleft != psutil.POWER_TIME_UNLIMITED else "Unlimited"
            return {
                "percentage": battery.percent,
                "charging": battery.power_plugged,
                "time_left": time_left
            }
        except Exception as e:
            return f"Battery sensor error: {e}"

    # ---------------- SYSTEM INFO ----------------
    def system_info(self):
        """Return basic system info."""
        return {
            "os": self.system,
            "cpu_usage": psutil.cpu_percent(interval=1),
            "memory_usage": psutil.virtual_memory().percent
        }

    # ---------------- OPTIONAL MOTION DETECTION ----------------
    def motion_detected(self, camera_index=0, threshold=300000):
        """Detect movement using webcam frame difference (optional)."""
        try:
            cap = cv2.VideoCapture(camera_index)
            if not cap.isOpened():
                return "Camera not accessible"

            ret, frame1 = cap.read()
            ret, frame2 = cap.read()
            motion = False

            while ret:
                diff = cv2.absdiff(frame1, frame2)
                gray = cv2.cvtColor(diff, cv2.COLOR_BGR2GRAY)
                blur = cv2.GaussianBlur(gray, (5,5), 0)
                _, thresh = cv2.threshold(blur, 20, 255, cv2.THRESH_BINARY)
                motion_value = np.sum(thresh)

                if motion_value > threshold:
                    motion = True
                    break

                frame1 = frame2
                ret, frame2 = cap.read()

            cap.release()
            cv2.destroyAllWindows()
            return "Movement detected" if motion else "No significant motion"
        except Exception as e:
            return f"Motion detection error: {e}"



#_____________________________Quantum_______________________________
import numpy as np
import math
import cmath
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import reduce
import itertools
import random
import time
import scipy
# Constants
HBAR = 1.054571817e-34
ELECTRON_MASS = 9.10938356e-31
E_CHARGE = 1.602176634e-19
PI = math.pi


# Helper utilities
def kron(*mats):
    """Kronecker product of multiple matrices."""
    return reduce(np.kron, mats)

def normalize(vec):
    """Normalize complex vector."""
    norm = np.linalg.norm(vec)
    return vec / norm if norm > 0 else vec

def random_complex_state(n):
    """Generate random normalized quantum state of n qubits."""
    state = np.random.rand(2**n) + 1j*np.random.rand(2**n)
    return normalize(state)


# _______________________________Quantum Calculations (1.2.4)_______________________________ #
import numpy as np
import math, cmath, random, time
from functools import reduce
from concurrent.futures import ThreadPoolExecutor, as_completed
import itertools
import scipy.linalg

# Constants
HBAR = 1.054571817e-34
ELECTRON_MASS = 9.10938356e-31
E_CHARGE = 1.602176634e-19
PI = math.pi

# Utility helpers
def kron(*mats): return reduce(np.kron, mats)
def normalize(vec): return vec / np.linalg.norm(vec) if np.linalg.norm(vec) != 0 else vec
def dagger(M): return np.conjugate(M.T)
def random_state(n): return normalize(np.random.rand(2**n) + 1j*np.random.rand(2**n))

# _________________________________ Quantum Calculations _________________________________ #
class QuantumCalculation:
    """
    Handles theoretical physics and chemistry quantum calculations:
    - Potential wells
    - Quantum harmonic oscillators
    - Energy levels
    - Tunneling
    - Uncertainty, interference, probability
    - Wavepacket evolution
    """

    def __init__(self):
        self.history = {}

    # ---------- Fundamental Mechanics ----------
    def wkb_tunneling(self, E, V0, a, m=ELECTRON_MASS):
        if E >= V0: return 1.0
        kappa = np.sqrt(2*m*(V0 - E)) / HBAR
        T = np.exp(-2*kappa*a)
        return float(T)

    def infinite_square_well_energy(self, n, L, m=ELECTRON_MASS):
        return (n**2 * PI**2 * HBAR**2) / (2*m*L**2)

    def harmonic_oscillator_energy(self, n, omega):
        return HBAR * omega * (n + 0.5)

    def particle_in_box_wavefunc(self, n, L, x):
        return np.sqrt(2/L) * np.sin(n * PI * x / L)

    def uncertainty_relation(self, sigma_x, sigma_p):
        return sigma_x * sigma_p >= HBAR / 2

    def commutator(self, A, B):
        return A @ B - B @ A

    def expectation_value(self, psi, operator):
        return np.vdot(psi, operator @ psi).real

    def interference_pattern(self, amp1, amp2):
        return np.abs(amp1 + amp2)**2

    # ---------- Advanced Numerical ----------
    def eigenvalues_and_vectors(self, H):
        vals, vecs = np.linalg.eigh(H)
        return {'energies': vals, 'states': vecs}

    def potential_energy_curve(self, potential_func, x_range):
        return np.array([potential_func(x) for x in x_range])

    def wavepacket(self, x, x0, p0, sigma):
        pref = (1/(sigma * np.sqrt(2*PI)))
        return pref * np.exp(-(x-x0)**2/(2*sigma**2)) * np.exp(1j*p0*(x-x0)/HBAR)

    def propagate_wavepacket(self, psi, H, t):
        U = scipy.linalg.expm(-1j * H * t / HBAR)
        return U @ psi

    def probability_density(self, psi):
        return np.abs(psi)**2

    def expectation_potential(self, psi, V):
        return np.real(np.vdot(psi, V * psi))

    def density_matrix(self, psi):
        return np.outer(psi, np.conjugate(psi))

    def entanglement_entropy(self, rho):
        vals = np.linalg.eigvalsh(rho)
        vals = np.clip(vals, 1e-12, 1)
        return -np.sum(vals*np.log2(vals))

    # ---------- Quantum Chemistry ----------
    def hamiltonian_hydrogen_like(self, Z, n):
        """Simplified hydrogenic energy levels: E = -13.6 * Z² / n² eV"""
        return -13.6 * (Z**2) / (n**2)

    def molecular_overlap(self, psi1, psi2):
        return np.abs(np.vdot(psi1, psi2))**2

    def coulomb_potential(self, q1, q2, r):
        return (q1*q2) / (4 * PI * 8.854e-12 * r)

    def run_parallel(self, funcs):
        results = {}
        with ThreadPoolExecutor(max_workers=len(funcs)) as ex:
            futures = {ex.submit(func): name for name, func in funcs.items()}
            for f in as_completed(futures):
                results[futures[f]] = f.result()
        return results


# _________________________________ Quantum Computation _________________________________ #
class QuantumComputing:
    """
    Handles algorithmic and computational quantum logic operations:
    - Gate simulation
    - QFT, Grover, Teleportation
    - Measurement, density matrix, entropy
    - Variational circuits, entanglement detection
    """

    def __init__(self):
        self.I = np.eye(2, complex)
        self.X = np.array([[0,1],[1,0]], complex)
        self.Y = np.array([[0,-1j],[1j,0]], complex)
        self.Z = np.array([[1,0],[0,-1]], complex)
        self.H = (1/np.sqrt(2))*np.array([[1,1],[1,-1]], complex)
        self.S = np.array([[1,0],[0,1j]], complex)
        self.T = np.array([[1,0],[0, np.exp(1j*PI/4)]], complex)
        self.gates = {'I':self.I, 'X':self.X, 'Y':self.Y, 'Z':self.Z, 'H':self.H}

    def create_state(self, n):
        s = np.zeros(2**n, complex); s[0]=1; return s

    def apply_gate(self, state, gate, target, n):
        ops=[self.I]*n; ops[target]=gate
        return kron(*ops) @ state

    def controlled_gate(self, state, control, target, gate, n):
        P0=np.array([[1,0],[0,0]]); P1=np.array([[0,0],[0,1]])
        op0=[self.I]*n; op1=[self.I]*n
        op0[control]=P0; op1[control]=P1; op1[target]=gate
        U=kron(*op0)+kron(*op1)
        return U@state

    def qft(self, state, n):
        N=2**n
        F=np.array([[np.exp(2j*PI*j*k/N) for k in range(N)] for j in range(N)]) / np.sqrt(N)
        return F @ state

    def grover(self, n, marked):
        N=2**n; s=np.ones(N)/np.sqrt(N); s[marked]*=-1
        return 2*s*np.vdot(np.ones(N)/np.sqrt(N), s)-s

    def teleportation(self, psi):
        bell=(1/np.sqrt(2))*np.array([1,0,0,1], complex)
        state=kron(psi,bell)
        state=self.controlled_gate(state,0,1,self.X,3)
        state=self.apply_gate(state,self.H,0,3)
        return normalize(state)

    def measure(self, state):
        probs=np.abs(state)**2
        index=np.random.choice(len(state), p=probs/np.sum(probs))
        return format(index,'b').zfill(int(np.log2(len(state)))), probs[index]

    def density_matrix(self, state): return np.outer(state, state.conj())

    def partial_trace(self, rho, keep, dims):
        """Partial trace for subsystem reductions."""
        n = len(dims)
        keep_dims = np.prod([dims[i] for i in keep])
        trace_dims = np.prod([dims[i] for i in range(n) if i not in keep])
        rho = rho.reshape(dims + dims)
        traced = np.trace(rho, axis1=0, axis2=n)
        return traced.reshape((keep_dims, keep_dims))

    def entanglement_entropy(self, state, dims=[2,2]):
        rho = self.density_matrix(state)
        rhoA = self.partial_trace(rho, [0], dims)
        vals = np.real_if_close(np.linalg.eigvals(rhoA))
        vals = np.clip(vals, 1e-12, 1)
        return -np.sum(vals*np.log2(vals))

    def variational_ansatz(self, params):
        """Simple variational ansatz for demonstration."""
        theta, phi = params
        return np.array([math.cos(theta/2), np.exp(1j*phi)*math.sin(theta/2)], complex)

    def cost_function(self, psi, H):
        return np.real(np.vdot(psi, H @ psi))

    def variational_optimize(self, H):
        best = (None, float('inf'))
        for theta in np.linspace(0,2*PI,50):
            for phi in np.linspace(0,2*PI,50):
                psi=self.variational_ansatz((theta,phi))
                e=self.cost_function(psi,H)
                if e<best[1]: best=(psi,e)
        return best

    def decoherence(self, rho, p):
        d=len(rho)
        return (1-p)*rho + p*np.eye(d)/d

    def benchmark(self):
        start=time.time()
        tasks={'grover':lambda:self.grover(3,5),'qft':lambda:self.qft(self.create_state(3),3)}
        results={}
        with ThreadPoolExecutor(max_workers=2) as ex:
            futures={ex.submit(func):name for name,func in tasks.items()}
            for f in as_completed(futures):
                results[futures[f]]=f.result()
        return {"time":time.time()-start,"results":results}


#___________________________Statistics__________________________________
"""
statistics_toolkit.py

Statistics (Hybrid Edition v3.0)

A single-file, production-oriented Statistics class that provides an extensive
suite of statistical methods and plotting utilities. The implementation is
built on a lightweight core (NumPy + Matplotlib) and automatically enables
advanced functionality when optional libraries are present (SciPy, Pandas,
Statsmodels, Seaborn).

Design goals:
- Run in minimal environments (NumPy + Matplotlib) while gracefully enabling
  pro features when optional libs are installed.
- Accept common input types (list, np.ndarray, pd.Series, pd.DataFrame).
- Provide broad statistical coverage (means, dispersion, tests, regression,
  distributions, time-series) and versatile plotting (single/multi-frame).

Author: Generated for Ibrahim Shahid
"""


import math
from typing import Optional, Sequence, Tuple, List, Dict, Any, Union
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# Optional libraries (auto-detected)
try:
    from scipy import stats
    from scipy.stats import gaussian_kde
except Exception:
    stats = None
    gaussian_kde = None

try:
    import statsmodels.api as sm
    import statsmodels.tsa.api as tsa
    from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
except Exception:
    sm = None
    tsa = None
    plot_acf = None
    plot_pacf = None

try:
    import seaborn as sns
except Exception:
    sns = None


ArrayLike = Union[Sequence[float], np.ndarray, pd.Series]


class Statistics:
    """Hybrid Statistics class with core + pro features.

    Core dependencies: numpy, matplotlib
    Optional (pro) dependencies: scipy, pandas, statsmodels, seaborn

    Instantiate with a 1-D dataset (list, numpy array, pandas Series). For
    multi-column inputs (DataFrame), several functions accept DataFrame
    directly (pairwise plots / covariance / correlation matrix).
    """

    def __init__(self, data: Optional[ArrayLike] = None, name: Optional[str] = None):
        self.name = name or 'series'
        self.data: Optional[pd.Series] = None
        if data is not None:
            self.set_data(data)

    # ------------------------------ Input handling ------------------------------
    @staticmethod
    def _to_series(data: ArrayLike) -> pd.Series:
        if isinstance(data, pd.Series):
            return data.dropna().reset_index(drop=True)
        if isinstance(data, pd.DataFrame):
            raise ValueError('Expected 1-D data; received DataFrame. Use column or pass DataFrame to multi-variate methods.')
        arr = np.asarray(data)
        if arr.ndim > 1:
            arr = arr.flatten()
        series = pd.Series(arr).replace([np.inf, -np.inf], np.nan).dropna().reset_index(drop=True)
        return series

    def set_data(self, data: ArrayLike):
        self.data = self._to_series(data)
        return self

    def as_array(self) -> np.ndarray:
        if self.data is None:
            raise ValueError('No data set')
        return self.data.to_numpy()

    # ------------------------------ Central tendency ------------------------------
    def arithmetic_mean(self) -> float:
        return float(self.data.mean())

    def geometric_mean(self) -> float:
        x = self.as_array()
        if np.any(x <= 0):
            raise ValueError('Geometric mean requires all positive values')
        return float(np.exp(np.mean(np.log(x))))

    def harmonic_mean(self) -> float:
        x = self.as_array()
        if np.any(x == 0):
            raise ValueError('Harmonic mean undefined for zero values')
        return float(len(x) / np.sum(1.0 / x))

    def rms(self) -> float:
        x = self.as_array()
        return float(np.sqrt(np.mean(x ** 2)))

    def weighted_mean(self, weights: Sequence[float]) -> float:
        x = self.as_array()
        w = np.asarray(weights, dtype=float)
        if len(w) != len(x):
            raise ValueError('weights must have same length as data')
        return float(np.sum(w * x) / np.sum(w))

    def trimmed_mean(self, proportion: float = 0.1) -> float:
        x = np.sort(self.as_array())
        n = len(x)
        k = int(math.floor(proportion * n))
        if k * 2 >= n:
            raise ValueError('proportion too large')
        trimmed = x[k: n - k]
        return float(np.mean(trimmed))

    def median(self) -> float:
        return float(self.data.median())

    def mode(self) -> List[float]:
        m = self.data.mode()
        return list(m.values) if not m.empty else []

    def midrange(self) -> float:
        x = self.as_array()
        return float((np.min(x) + np.max(x)) / 2.0)

    # ------------------------------ Dispersion ------------------------------
    def variance(self, ddof: int = 1) -> float:
        return float(self.data.var(ddof=ddof))

    def population_variance(self) -> float:
        return float(self.data.var(ddof=0))

    def std(self, ddof: int = 1) -> float:
        return float(self.data.std(ddof=ddof))

    def population_std(self) -> float:
        return float(self.data.std(ddof=0))

    def coefficient_of_variation(self) -> float:
        mean = self.arithmetic_mean()
        if mean == 0:
            return float('nan')
        return float(self.std() / mean)

    def range(self) -> float:
        x = self.as_array()
        return float(np.max(x) - np.min(x))

    def iqr(self) -> float:
        return float(self.data.quantile(0.75) - self.data.quantile(0.25))

    def quartile_range(self, lower: float = 0.25, upper: float = 0.75) -> float:
        return float(self.data.quantile(upper) - self.data.quantile(lower))

    def mad(self) -> float:
        x = self.as_array()
        med = np.median(x)
        return float(np.median(np.abs(x - med)))

    # ------------------------------ Shape & outliers ------------------------------
    def skewness(self) -> float:
        if stats is not None:
            return float(stats.skew(self.as_array(), nan_policy='omit'))
        return float(self.data.skew())

    def kurtosis(self) -> float:
        if stats is not None:
            return float(stats.kurtosis(self.as_array(), fisher=True, nan_policy='omit'))
        return float(self.data.kurt())

    def z_scores(self) -> np.ndarray:
        x = self.as_array()
        mu = np.mean(x)
        sd = np.std(x, ddof=1)
        return (x - mu) / sd

    def outliers_iqr(self, factor: float = 1.5) -> np.ndarray:
        q1 = self.data.quantile(0.25)
        q3 = self.data.quantile(0.75)
        iqr = q3 - q1
        low = q1 - factor * iqr
        high = q3 + factor * iqr
        return self.data[(self.data < low) | (self.data > high)].to_numpy()

    def outliers_zscore(self, threshold: float = 3.0) -> np.ndarray:
        zs = self.z_scores()
        return self.as_array()[np.abs(zs) > threshold]

    # ------------------------------ Percentiles & quantiles ------------------------------
    def percentile(self, p: float) -> float:
        return float(np.nanpercentile(self.as_array(), p))

    def quantiles(self, q: Sequence[float] = (0.25, 0.5, 0.75)) -> Dict[str, float]:
        qv = np.quantile(self.as_array(), q)
        return {f'q{int(p*100)}': float(v) for p, v in zip(q, qv)}

    # ------------------------------ Correlation & association ------------------------------
    @staticmethod
    def pearsonr(x: Sequence[float], y: Sequence[float]) -> Tuple[float, Optional[float]]:
        if stats is not None:
            r, p = stats.pearsonr(np.asarray(x), np.asarray(y))
            return float(r), float(p)
        r = np.corrcoef(np.asarray(x), np.asarray(y))[0, 1]
        return float(r), None

    @staticmethod
    def spearmanr(x: Sequence[float], y: Sequence[float]) -> Tuple[float, Optional[float]]:
        if stats is not None:
            r, p = stats.spearmanr(np.asarray(x), np.asarray(y))
            return float(r), float(p)
        xr = pd.Series(x).rank().values
        yr = pd.Series(y).rank().values
        r = np.corrcoef(xr, yr)[0, 1]
        return float(r), None

    @staticmethod
    def kendalltau(x: Sequence[float], y: Sequence[float]) -> Tuple[float, Optional[float]]:
        if stats is not None:
            r, p = stats.kendalltau(np.asarray(x), np.asarray(y))
            return float(r), float(p)
        return float('nan'), None

    # ------------------------------ Distribution & tests ------------------------------
    def fit_normal(self) -> Dict[str, float]:
        x = self.as_array()
        mu = float(np.mean(x))
        sigma = float(np.std(x, ddof=1))
        return {'mu': mu, 'sigma': sigma}

    def ks_test_normal(self) -> Dict[str, float]:
        if stats is None:
            raise RuntimeError('scipy required for ks_test_normal')
        mu, sigma = self.fit_normal()['mu'], self.fit_normal()['sigma']
        d, p = stats.kstest(self.as_array(), 'norm', args=(mu, sigma))
        return {'d_stat': float(d), 'p_value': float(p)}

    def shapiro_test(self) -> Dict[str, float]:
        if stats is None:
            raise RuntimeError('scipy required for shapiro_test')
        stat, p = stats.shapiro(self.as_array())
        return {'stat': float(stat), 'p_value': float(p)}

    def t_test_1sample(self, popmean: float = 0.0) -> Dict[str, Any]:
        x = self.as_array()
        if stats is not None:
            tstat, p = stats.ttest_1samp(x, popmean)
            return {'t_stat': float(tstat), 'p_value': float(p)}
        # fallback
        n = len(x)
        mu = float(np.mean(x))
        se = float(np.std(x, ddof=1) / math.sqrt(n))
        tstat = (mu - popmean) / se
        return {'t_stat': float(tstat), 'p_value': None}

    def t_test_ind(self, other: Sequence[float], equal_var: bool = True) -> Dict[str, Any]:
        if stats is not None:
            tstat, p = stats.ttest_ind(self.as_array(), np.asarray(other), equal_var=equal_var, nan_policy='omit')
            return {'t_stat': float(tstat), 'p_value': float(p)}
        # fallback: basic implementation
        a = self.as_array()
        b = np.asarray(other)
        na, nb = len(a), len(b)
        ma, mb = a.mean(), b.mean()
        sa2, sb2 = a.var(ddof=1), b.var(ddof=1)
        if equal_var:
            sp2 = ((na - 1) * sa2 + (nb - 1) * sb2) / (na + nb - 2)
            se = math.sqrt(sp2 * (1 / na + 1 / nb))
            df = na + nb - 2
        else:
            se = math.sqrt(sa2 / na + sb2 / nb)
            df = (sa2 / na + sb2 / nb) ** 2 / ((sa2 ** 2) / (na ** 2 * (na - 1)) + (sb2 ** 2) / (nb ** 2 * (nb - 1)))
        t_stat = (ma - mb) / se
        return {'t_stat': float(t_stat), 'df': float(df), 'p_value': None}

    def anova_oneway(self, groups: List[Sequence[float]]) -> Dict[str, float]:
        if stats is None:
            raise RuntimeError('scipy required for anova_oneway')
        fstat, p = stats.f_oneway(*groups)
        return {'f_stat': float(fstat), 'p_value': float(p)}

    # ------------------------------ Regression & modeling ------------------------------
    @staticmethod
    def ols(x: Sequence[float], y: Sequence[float], add_constant: bool = True) -> Dict[str, Any]:
        X = np.asarray(x)
        Y = np.asarray(y)
        if add_constant:
            Xmat = np.column_stack((np.ones(len(X)), X))
        else:
            Xmat = X.reshape(-1, 1)
        beta = np.linalg.lstsq(Xmat, Y, rcond=None)[0]
        yhat = Xmat @ beta
        residuals = Y - yhat
        sse = (residuals ** 2).sum()
        s2 = sse / (len(Y) - Xmat.shape[1])
        cov = s2 * np.linalg.pinv(Xmat.T @ Xmat)
        se = np.sqrt(np.diag(cov))
        return {'beta': beta, 'yhat': yhat, 'residuals': residuals, 'se': se}

    @staticmethod
    def polyfit(x: Sequence[float], y: Sequence[float], deg: int = 2) -> Dict[str, Any]:
        coeffs = np.polyfit(np.asarray(x), np.asarray(y), deg)
        p = np.poly1d(coeffs)
        yhat = p(np.asarray(x))
        residuals = np.asarray(y) - yhat
        return {'coeffs': coeffs, 'yhat': yhat, 'residuals': residuals}

    # ------------------------------ Time series helpers ------------------------------
    @staticmethod
    def rolling_mean(series: Sequence[float], window: int = 5, center: bool = False) -> pd.Series:
        s = pd.Series(series).dropna()
        return s.rolling(window=window, center=center).mean()

    @staticmethod
    def ewma(series: Sequence[float], span: int = 12) -> pd.Series:
        s = pd.Series(series).dropna()
        return s.ewm(span=span, adjust=False).mean()

    @staticmethod
    def acf(series: Sequence[float], nlags: int = 40) -> Tuple[np.ndarray, np.ndarray]:
        s = pd.Series(series).dropna()
        n = len(s)
        acf_vals = np.array([s.autocorr(lag=i) for i in range(min(n - 1, nlags + 1))])
        lags = np.arange(len(acf_vals))
        return lags, acf_vals

    @staticmethod
    def pacf(series: Sequence[float], nlags: int = 40) -> Tuple[np.ndarray, np.ndarray]:
        if plot_pacf is not None and tsa is not None:
            vals = tsa.stattools.pacf(np.asarray(series), nlags=nlags)
            lags = np.arange(len(vals))
            return lags, vals
        return np.arange(nlags + 1), np.zeros(nlags + 1)

    @staticmethod
    def seasonal_decompose(series: Sequence[float], period: Optional[int] = None) -> Dict[str, pd.Series]:
        if tsa is None:
            raise RuntimeError('statsmodels required for seasonal_decompose')
        res = tsa.seasonal_decompose(pd.Series(series).dropna(), period=period, model='additive', extrapolate_trend='freq')
        return {'trend': res.trend, 'seasonal': res.seasonal, 'resid': res.resid}

    # ------------------------------ Spectral analysis ------------------------------
    @staticmethod
    def fourier(series: Sequence[float]) -> Dict[str, np.ndarray]:
        x = np.asarray(series)
        n = len(x)
        freqs = np.fft.rfftfreq(n)
        fft_vals = np.fft.rfft(x)
        psd = (np.abs(fft_vals) ** 2) / n
        return {'freqs': freqs, 'fft': fft_vals, 'psd': psd}

    # ------------------------------ Plotting utilities ------------------------------
    @staticmethod
    def _ensure_ax(ax: Optional[plt.Axes] = None) -> Tuple[plt.Figure, plt.Axes]:
        if ax is None:
            fig, ax = plt.subplots()
            return fig, ax
        return ax.figure, ax

    def plot_histogram(self, bins: int = 30, ax: Optional[plt.Axes] = None, kde: bool = False, **kwargs) -> plt.Axes:
        fig, ax = self._ensure_ax(ax)
        if sns is not None and kde:
            sns.histplot(self.data, bins=bins, kde=True, ax=ax, **kwargs)
        else:
            ax.hist(self.as_array(), bins=bins, **kwargs)
            if kde and gaussian_kde is not None:
                x = np.linspace(np.min(self.as_array()), np.max(self.as_array()), 200)
                kde_vals = gaussian_kde(self.as_array())(x)
                ax.plot(x, kde_vals * len(self.as_array()) * (x[1] - x[0]) * bins / 10)
        ax.set_title(f'Histogram - {self.name}')
        return ax

    def plot_box(self, ax: Optional[plt.Axes] = None, **kwargs) -> plt.Axes:
        fig, ax = self._ensure_ax(ax)
        ax.boxplot(self.as_array(), **kwargs)
        ax.set_title(f'Boxplot - {self.name}')
        return ax

    def plot_violin(self, ax: Optional[plt.Axes] = None, **kwargs) -> plt.Axes:
        if sns is None:
            return self.plot_box(ax=ax, **kwargs)
        fig, ax = self._ensure_ax(ax)
        sns.violinplot(x=self.data, ax=ax, **kwargs)
        ax.set_title(f'Violin - {self.name}')
        return ax

    def plot_kde(self, ax: Optional[plt.Axes] = None, **kwargs) -> plt.Axes:
        fig, ax = self._ensure_ax(ax)
        if sns is not None:
            sns.kdeplot(self.data, ax=ax, **kwargs)
        elif gaussian_kde is not None:
            x = np.linspace(np.min(self.as_array()), np.max(self.as_array()), 200)
            ax.plot(x, gaussian_kde(self.as_array())(x))
        else:
            raise RuntimeError('seaborn or scipy required for KDE plot')
        ax.set_title(f'KDE - {self.name}')
        return ax

    def plot_scatter(self, x: Sequence[float], y: Sequence[float], ax: Optional[plt.Axes] = None, label: Optional[str] = None, **kwargs) -> plt.Axes:
        fig, ax = self._ensure_ax(ax)
        ax.scatter(np.asarray(x), np.asarray(y), label=label, **kwargs)
        if label:
            ax.legend()
        ax.set_title('Scatter')
        return ax

    @staticmethod
    def plot_pairplot(df: pd.DataFrame, diag_kind: str = 'hist', **kwargs) -> plt.Figure:
        if sns is None:
            raise RuntimeError('seaborn required for pairplot')
        g = sns.pairplot(df, diag_kind=diag_kind, **kwargs)
        return g.fig

    def plot_time_series(self, times: Optional[Sequence[Any]] = None, ax: Optional[plt.Axes] = None, label: Optional[str] = None) -> plt.Axes:
        fig, ax = self._ensure_ax(ax)
        if times is None:
            ax.plot(self.as_array(), label=label)
        else:
            ax.plot(times, self.as_array(), label=label)
        if label:
            ax.legend()
        ax.set_title(f'Time Series - {self.name}')
        return ax

    @staticmethod
    def plot_multi(plots: List[Tuple[str, Any]], layout: Tuple[int, int] = (2, 2), figsize: Tuple[int, int] = (12, 8)) -> plt.Figure:
        """Plot a list of (title, plotting_function_or_tuple) into a grid.

        Each item in `plots` can be either:
        - ('title', callable(ax) -> None) where the callable draws into provided ax, or
        - ('title', ('hist', data, kwargs)) shorthand for built-in plot types.
        """
        rows, cols = layout
        fig, axes = plt.subplots(rows, cols, figsize=figsize)
        axes_flat = np.array(axes).flatten()
        for i, (title, plot_spec) in enumerate(plots):
            if i >= len(axes_flat):
                break
            ax = axes_flat[i]
            if callable(plot_spec):
                plot_spec(ax)
            elif isinstance(plot_spec, tuple):
                ptype = plot_spec[0]
                if ptype == 'hist':
                    ax.hist(np.asarray(plot_spec[1]), **(plot_spec[2] if len(plot_spec) > 2 else {}))
                elif ptype == 'line':
                    ax.plot(np.asarray(plot_spec[1]))
                elif ptype == 'scatter':
                    ax.scatter(np.asarray(plot_spec[1]), np.asarray(plot_spec[2]))
                elif ptype == 'box':
                    ax.boxplot(np.asarray(plot_spec[1]))
                else:
                    ax.text(0.5, 0.5, 'Unknown plot type', ha='center')
            else:
                ax.text(0.5, 0.5, 'Invalid plot spec', ha='center')
            ax.set_title(title)
        for j in range(i + 1, len(axes_flat)):
            axes_flat[j].set_visible(False)
        fig.tight_layout()
        return fig

    def plot_acf_pacf(self, nlags: int = 40) -> plt.Figure:
        if plot_acf is not None and plot_pacf is not None:
            fig, axes = plt.subplots(2, 1, figsize=(10, 6))
            plot_acf(self.as_array(), ax=axes[0], lags=nlags)
            plot_pacf(self.as_array(), ax=axes[1], lags=nlags)
            fig.tight_layout()
            return fig
        else:
            lags, acf_vals = self.acf(self.as_array(), nlags=nlags)
            lags_p, pacf_vals = self.pacf(self.as_array(), nlags=nlags)
            fig, axes = plt.subplots(2, 1, figsize=(10, 6))
            axes[0].bar(lags, acf_vals)
            axes[1].bar(lags_p, pacf_vals)
            axes[0].set_title('ACF')
            axes[1].set_title('PACF')
            fig.tight_layout()
            return fig
    def show(self):
        plt.show()
        
    def plot_spectrum(self) -> plt.Figure:
        sp = self.fourier(self.as_array())
        fig, ax = plt.subplots()
        ax.plot(sp['freqs'], sp['psd'])
        ax.set_title('Power Spectral Density')
        return fig


# # ------------------------------ Example usage  ------------------------------
# if __name__ == '__main__':
#     # This block demonstrates typical usage. When used as a module, import the
#     # Statistics class and instantiate with your data.
#     import numpy as _np

#     rng = _np.random.default_rng(0)
#     data = rng.normal(loc=5.0, scale=2.0, size=500)
#     S = Statistics(data, name='example_series')

#     # Basic stats
#     print('Mean:', S.arithmetic_mean())
#     print('Geometric mean ok? ->', end=' ')
#     try:
#         print(S.geometric_mean())
#     except Exception as e:
#         print('error:', e)

#     # Plots
#     fig = S.plot_histogram(kde=True)
#     fig.savefig('example_hist.png')

#     # Multi-plot example
#     plots = [
#         ('Histogram', ('hist', data, {'bins': 30})),
#         ('Box', ('box', data)),
#         ('KDE', lambda ax: S.plot_kde(ax=ax)),
#         ('Time', ('line', data))
#     ]
#     fig2 = Statistics.plot_multi(plots, layout=(2, 2))
#     fig2.savefig('example_multi.png')

#     print('Saved example plots: example_hist.png, example_multi.png')



#________________________Translator_____________________________
import asyncio
from googletrans import Translator as GoogleTranslator, LANGUAGES

class Translator:
    def __init__(self):
        self.translator = GoogleTranslator()

    async def _translate_async(self, input_text, output_language):
        # Detect language
        detected = await self.translator.detect(input_text)
        detected_lang_code = detected.lang
        detected_lang_name = LANGUAGES.get(detected_lang_code, "Unknown").capitalize()

        # Find target language code
        output_lang_code = None
        for code, name in LANGUAGES.items():
            if name.lower() == output_language.lower():
                output_lang_code = code
                break
        if not output_lang_code:
            return {"Error": f"Output language '{output_language}' not supported."}

        # Translate text
        translation = await self.translator.translate(input_text, src=detected_lang_code, dest=output_lang_code)
        return {
            "Detected Language": detected_lang_name,
            "Target Language": output_language.capitalize(),
            "Input Text": input_text,
            "Translated Text": translation.text
        }

    def translate(self, input_text="Hello", output_language="Urdu"):
        """
        Synchronous wrapper for users.
        They can call this normally without async knowledge.
        """
        try:
            loop = asyncio.get_running_loop()
            return loop.run_until_complete(self._translate_async(input_text, output_language))
        except RuntimeError:
            return asyncio.run(self._translate_async(input_text, output_language))
        
        
        
#____________________________________Encryption__________________________________________________
# encryption_suite.py
# Requires Python 3.8+. Optional packages: pycryptodome, bcrypt, argon2-cffi

import base64
import hashlib
import hmac
import secrets
from typing import Tuple

# Optional imports
try:
    from Crypto.Cipher import AES, PKCS1_OAEP
    from Crypto.PublicKey import RSA
    from Crypto.Util.Padding import pad, unpad
    HAS_PYCRYPTODOME = True
except Exception:
    HAS_PYCRYPTODOME = False

try:
    import bcrypt
    HAS_BCRYPT = True
except Exception:
    HAS_BCRYPT = False

try:
    from argon2 import PasswordHasher
    HAS_ARGON2 = True
except Exception:
    HAS_ARGON2 = False


class Encryption_Decryption:
    """
    Multi-tier encryption & hashing utilities.
    Methods are named for clarity: trivial -> intermediate -> strong.
    All encrypt/decrypt methods accept bytes or str (converted to utf-8).
    """

    def __init__(self):
        pass

    # --------------------------
    # Utilities
    # --------------------------
    @staticmethod
    def _to_bytes(data):
        return data if isinstance(data, (bytes, bytearray)) else str(data).encode('utf-8')

    @staticmethod
    def secure_random_bytes(n: int) -> bytes:
        return secrets.token_bytes(n)

    # --------------------------
    # Trivial / Educational (easy to crack)
    # --------------------------
    def caesar_encrypt(self, plaintext: str, shift: int = 3) -> str:
        b = self._to_bytes(plaintext)
        return bytes(( (c + shift) % 256 ) for c in b).hex()

    def caesar_decrypt(self, hex_cipher: str, shift: int = 3) -> str:
        data = bytes.fromhex(hex_cipher)
        return bytes(( (c - shift) % 256 ) for c in data).decode('utf-8', errors='ignore')

    def xor_encrypt(self, plaintext: str, key: bytes) -> str:
        b = self._to_bytes(plaintext)
        k = key if isinstance(key, (bytes,bytearray)) else self._to_bytes(key)
        out = bytes([b[i] ^ k[i % len(k)] for i in range(len(b))])
        return base64.b64encode(out).decode('ascii')

    def xor_decrypt(self, b64_cipher: str, key: bytes) -> str:
        data = base64.b64decode(b64_cipher)
        k = key if isinstance(key, (bytes,bytearray)) else self._to_bytes(key)
        out = bytes([data[i] ^ k[i % len(k)] for i in range(len(data))])
        return out.decode('utf-8', errors='ignore')

    def rot13(self, text: str) -> str:
        return text.translate(str.maketrans(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz",
            "NOPQRSTUVWXYZABCDEFGHIJKLMnopqrstuvwxyzabcdefghijklm"))

    # --------------------------
    # Moderate (symmetric) - requires pycryptodome
    # --------------------------
    def aes_cbc_encrypt(self, plaintext: str, key: bytes) -> Tuple[str, str]:
        """
        AES-256-CBC with PKCS7 padding.
        Returns (iv_hex, cipher_hex).
        """
        if not HAS_PYCRYPTODOME:
            raise RuntimeError("pycryptodome required for AES functions. pip install pycryptodome")
        key = hashlib.sha256(self._to_bytes(key)).digest()  # derive 32 bytes key
        iv = self.secure_random_bytes(16)
        cipher = AES.new(key, AES.MODE_CBC, iv)
        ct = cipher.encrypt(pad(self._to_bytes(plaintext), AES.block_size))
        return iv.hex(), ct.hex()

    def aes_cbc_decrypt(self, iv_hex: str, cipher_hex: str, key: bytes) -> str:
        if not HAS_PYCRYPTODOME:
            raise RuntimeError("pycryptodome required for AES functions. pip install pycryptodome")
        key = hashlib.sha256(self._to_bytes(key)).digest()
        iv = bytes.fromhex(iv_hex)
        ct = bytes.fromhex(cipher_hex)
        cipher = AES.new(key, AES.MODE_CBC, iv)
        pt = unpad(cipher.decrypt(ct), AES.block_size)
        return pt.decode('utf-8', errors='ignore')

    def aes_gcm_encrypt(self, plaintext: str, key: bytes) -> Tuple[str, str, str]:
        """
        AES-256-GCM. Returns (nonce_hex, cipher_hex, tag_hex).
        """
        if not HAS_PYCRYPTODOME:
            raise RuntimeError("pycryptodome required for AES functions. pip install pycryptodome")
        key = hashlib.sha256(self._to_bytes(key)).digest()
        nonce = self.secure_random_bytes(12)
        cipher = AES.new(key, AES.MODE_GCM, nonce=nonce)
        ct, tag = cipher.encrypt_and_digest(self._to_bytes(plaintext))
        return nonce.hex(), ct.hex(), tag.hex()

    def aes_gcm_decrypt(self, nonce_hex: str, cipher_hex: str, tag_hex: str, key: bytes) -> str:
        if not HAS_PYCRYPTODOME:
            raise RuntimeError("pycryptodome required for AES functions. pip install pycryptodome")
        key = hashlib.sha256(self._to_bytes(key)).digest()
        nonce = bytes.fromhex(nonce_hex)
        ct = bytes.fromhex(cipher_hex)
        tag = bytes.fromhex(tag_hex)
        cipher = AES.new(key, AES.MODE_GCM, nonce=nonce)
        pt = cipher.decrypt_and_verify(ct, tag)
        return pt.decode('utf-8', errors='ignore')

    # --------------------------
    # Asymmetric (RSA) - requires pycryptodome
    # --------------------------
    def rsa_generate_keys(self, bits: int = 2048) -> Tuple[str, str]:
        if not HAS_PYCRYPTODOME:
            raise RuntimeError("pycryptodome required for RSA. pip install pycryptodome")
        key = RSA.generate(bits)
        private = key.export_key().decode('utf-8')
        public = key.publickey().export_key().decode('utf-8')
        return private, public

    def rsa_encrypt(self, plaintext: str, public_key_pem: str) -> str:
        if not HAS_PYCRYPTODOME:
            raise RuntimeError("pycryptodome required for RSA. pip install pycryptodome")
        pub = RSA.import_key(public_key_pem)
        cipher = PKCS1_OAEP.new(pub)
        ct = cipher.encrypt(self._to_bytes(plaintext))
        return base64.b64encode(ct).decode('ascii')

    def rsa_decrypt(self, b64_cipher: str, private_key_pem: str) -> str:
        if not HAS_PYCRYPTODOME:
            raise RuntimeError("pycryptodome required for RSA. pip install pycryptodome")
        priv = RSA.import_key(private_key_pem)
        cipher = PKCS1_OAEP.new(priv)
        pt = cipher.decrypt(base64.b64decode(b64_cipher))
        return pt.decode('utf-8', errors='ignore')

    # --------------------------
    # Hashing (one-way) - fast -> slow adaptive
    # --------------------------
    def hash_md5(self, data: str) -> str:
        return hashlib.md5(self._to_bytes(data)).hexdigest()

    def hash_sha1(self, data: str) -> str:
        return hashlib.sha1(self._to_bytes(data)).hexdigest()

    def hash_sha256(self, data: str) -> str:
        return hashlib.sha256(self._to_bytes(data)).hexdigest()

    def hash_sha3_512(self, data: str) -> str:
        return hashlib.sha3_512(self._to_bytes(data)).hexdigest()

    def hmac_sha256(self, key: str, message: str) -> str:
        return hmac.new(self._to_bytes(key), self._to_bytes(message), hashlib.sha256).hexdigest()

    # PBKDF2
    def pbkdf2_hash(self, password: str, salt: bytes = None, iterations: int = 200_000) -> Tuple[str, str, int]:
        salt = salt or self.secure_random_bytes(16)
        dk = hashlib.pbkdf2_hmac('sha256', self._to_bytes(password), salt, iterations)
        return dk.hex(), salt.hex(), iterations

    # scrypt (built-in)
    def scrypt_hash(self, password: str, salt: bytes = None, n: int = 2**14, r: int = 8, p: int = 1) -> Tuple[str, str]:
        salt = salt or self.secure_random_bytes(16)
        dk = hashlib.scrypt(self._to_bytes(password), salt=salt, n=n, r=r, p=p, dklen=64)
        return dk.hex(), salt.hex()

    # bcrypt (external)
    def bcrypt_hash(self, password: str, rounds: int = 12) -> str:
        if not HAS_BCRYPT:
            raise RuntimeError("bcrypt required. pip install bcrypt")
        return bcrypt.hashpw(self._to_bytes(password), bcrypt.gensalt(rounds)).decode('utf-8')

    def bcrypt_check(self, password: str, hashed: str) -> bool:
        if not HAS_BCRYPT:
            raise RuntimeError("bcrypt required. pip install bcrypt")
        return bcrypt.checkpw(self._to_bytes(password), hashed.encode('utf-8'))

    # argon2 (external)
    def argon2_hash(self, password: str) -> str:
        if not HAS_ARGON2:
            raise RuntimeError("argon2-cffi required. pip install argon2-cffi")
        ph = PasswordHasher()
        return ph.hash(password)

    def argon2_verify(self, password: str, hashed: str) -> bool:
        if not HAS_ARGON2:
            raise RuntimeError("argon2-cffi required. pip install argon2-cffi")
        ph = PasswordHasher()
        try:
            return ph.verify(hashed, password)
        except Exception:
            return False

    # --------------------------
    # Signed HMAC + Verify
    # --------------------------
    def sign_hmac(self, key: str, message: str) -> str:
        return hmac.new(self._to_bytes(key), self._to_bytes(message), hashlib.sha256).hexdigest()

    def verify_hmac(self, key: str, message: str, signature_hex: str) -> bool:
        expected = self.sign_hmac(key, message)
        return hmac.compare_digest(expected, signature_hex)

# --------------------------
# Quick usage examples
# --------------------------
# if __name__ == "__main__":
#     e = Encryption_Decryption()

#     # Trivial
#     c = e.caesar_encrypt("hello", 5)
#     assert e.caesar_decrypt(c, 5) == "hello"

#     # XOR
#     k = b"key"
#     x = e.xor_encrypt("secret", k)
#     assert e.xor_decrypt(x, k) == "secret"

#     # AES (if available)
#     if HAS_PYCRYPTODOME:
#         iv, ct = e.aes_cbc_encrypt("top secret", b"my password")
#         assert "secret" in e.aes_cbc_decrypt(iv, ct, b"my password")  # example

#     # RSA (if available)
#     if HAS_PYCRYPTODOME:
#         priv, pub = e.rsa_generate_keys(2048)
#         ct = e.rsa_encrypt("msg", pub)
#         assert e.rsa_decrypt(ct, priv) == "msg"

#     # Hashes
#     print("sha256:", e.hash_sha256("data"))

#     # bcrypt/argon2 (if installed)
#     if HAS_BCRYPT:
#         h = e.bcrypt_hash("hunter2")
#         assert e.bcrypt_check("hunter2", h)

#     if HAS_ARGON2:
#         ah = e.argon2_hash("hunter2")
#         assert e.argon2_verify("hunter2", ah)


#____________________________________OS____________________________________________________
import os
import platform
import subprocess
import shutil
from pathlib import Path
import psutil
import socket
import sys

class OS:
    def __init__(self):
        self.name = platform.system().lower()

    # Detect operating system
    def get_os(self):
        return self.name

    # Shutdown device
    def shutdown(self):
        print(f"Detected OS: {self.name}")
        try:
            if "windows" in self.name:
                os.system("shutdown /s /t 1")
            elif "linux" in self.name or "darwin" in self.name:
                os.system("sudo shutdown now")
            else:
                print("Unsupported OS for shutdown.")
        except Exception as e:
            print(f"Error: {e}")

    # Create folder
    def create_folder(self, path):
        try:
            os.makedirs(path, exist_ok=True)
            print(f"Folder created: {path}")
        except Exception as e:
            print(f"Error creating folder: {e}")

    # Create file
    def create_file(self, path, content=""):
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"File created: {path}")
        except Exception as e:
            print(f"Error creating file: {e}")

    # Delete folder
    def delete_folder(self, path):
        try:
            if os.path.exists(path):
                shutil.rmtree(path)
                print(f"Folder deleted: {path}")
            else:
                print("Folder does not exist.")
        except PermissionError:
            print("Admin privileges required to delete this folder.")
        except Exception as e:
            print(f"Error deleting folder: {e}")

    # Delete file
    def delete_file(self, path):
        try:
            if os.path.exists(path):
                os.remove(path)
                print(f"File deleted: {path}")
            else:
                print("File does not exist.")
        except PermissionError:
            print("Admin privileges required to delete this file.")
        except Exception as e:
            print(f"Error deleting file: {e}")

    # Check if path exists
    def path_exists(self, path):
        return Path(path).exists()

    # Run system command (cross-platform)
    def run_command(self, command):
        try:
            result = subprocess.run(command, shell=True, text=True, capture_output=True)
            if result.stdout.strip():
                print("Command Output:\n", result.stdout.strip())
            if result.stderr.strip():
                print("Errors:\n", result.stderr.strip())
        except Exception as e:
            print(f"Error executing command: {e}")

    # Open file or application
    def open_file(self, path):
        try:
            if not os.path.exists(path):
                print("Path does not exist.")
                return
            if "windows" in self.name:
                os.startfile(path)
            elif "darwin" in self.name:
                subprocess.run(["open", path])
            elif "linux" in self.name:
                subprocess.run(["xdg-open", path])
            else:
                print("Unsupported OS for file opening.")
        except Exception as e:
            print(f"Error opening file: {e}")

    # Fetch manufacturer (cross-platform)
    def get_manufacturer(self):
        try:
            if "windows" in self.name:
                # Try PowerShell first
                cmd = ['powershell', '-Command', "(Get-CimInstance Win32_ComputerSystem).Manufacturer"]
                result = subprocess.run(cmd, capture_output=True, text=True)
                if result.stdout.strip():
                    return result.stdout.strip()
                # Fallback to WMI if available
                try:
                    import wmi
                    c = wmi.WMI()
                    for system in c.Win32_ComputerSystem():
                        return system.Manufacturer
                except:
                    return None
            elif "linux" in self.name:
                if os.path.exists("/sys/class/dmi/id/sys_vendor"):
                    with open("/sys/class/dmi/id/sys_vendor") as f:
                        return f.read().strip()
            elif "darwin" in self.name:
                cmd = ["system_profiler", "SPHardwareDataType"]
                result = subprocess.run(cmd, capture_output=True, text=True)
                for line in result.stdout.splitlines():
                    if "Manufacturer" in line:
                        return line.split(":")[1].strip()
            return None
        except Exception:
            return None

    # Device information
    def device_info(self):
        try:
            info = {
                "os": platform.system(),
                "os_version": platform.version(),
                "os_release": platform.release(),
                "architecture": platform.machine(),
                "processor": platform.processor(),
                "cpu_usage_percent": psutil.cpu_percent(interval=1),
                "cpu_cores": psutil.cpu_count(logical=True),
                "memory_total_gb": round(psutil.virtual_memory().total / (1024**3), 2),
                "memory_used_gb": round(psutil.virtual_memory().used / (1024**3), 2),
                "memory_percent": psutil.virtual_memory().percent,
                "disk_total_gb": round(psutil.disk_usage('/').total / (1024**3), 2),
                "disk_used_gb": round(psutil.disk_usage('/').used / (1024**3), 2),
                "disk_percent": psutil.disk_usage('/').percent,
                "hostname": socket.gethostname(),
                "ip_address": socket.gethostbyname(socket.gethostname()),
                "battery_percent": psutil.sensors_battery().percent if psutil.sensors_battery() else None,
                "is_plugged": psutil.sensors_battery().power_plugged if psutil.sensors_battery() else None,
                "user": os.getenv("USERNAME") or os.getenv("USER"),
                "manufacturer": self.get_manufacturer()
            }
            return info
        except Exception as e:
            print(f"Error retrieving device info: {e}")
            return {}
